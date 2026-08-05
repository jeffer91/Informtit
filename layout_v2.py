from __future__ import annotations

from typing import Any

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import PageBreak, Paragraph, Spacer, Table, TableStyle

import institutional_export as base


def _font(run: Any, size: float, bold: bool = False) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold


def _margins(cell: Any, top: int = 30, start: int = 40, bottom: int = 30, end: int = 40) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _row_height(row: Any, inches: float, exact: bool = True) -> None:
    row.height = Inches(inches)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY if exact else WD_ROW_HEIGHT_RULE.AT_LEAST


def _cell_text(cell: Any, text: str, size: float = 8, bold: bool = False, left: bool = False) -> None:
    p = base.clear_cell(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if left else WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    _font(p.add_run(text), size, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _margins(cell)


def _clear_xml(cell: Any) -> None:
    for child in list(cell._tc):
        if child.tag != qn("w:tcPr"):
            cell._tc.remove(child)


def _right_header_table(cell: Any, report: dict[str, Any], width: float) -> None:
    _clear_xml(cell)
    nested = cell.add_table(rows=2, cols=1)
    nested.style = "Table Grid"
    nested.alignment = WD_TABLE_ALIGNMENT.CENTER
    nested.autofit = False
    for row in nested.rows:
        base.set_width(row.cells[0], width)
    _row_height(nested.rows[0], 0.31)
    _row_height(nested.rows[1], 0.22)
    _cell_text(nested.cell(0, 0), f"Código:\n{report.get('code', '')}", 6.5)
    _cell_text(nested.cell(1, 0), f"Versión: {report.get('version', '1.0')}", 6.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _margins(cell, 0, 0, 0, 0)


def setup_header(document: Any, report: dict[str, Any]) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1.55)
    section.bottom_margin = Inches(0.67)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.14)
    section.footer_distance = Inches(0.22)

    header = section.header
    header.is_linked_to_previous = False
    header.paragraphs[0].text = ""
    header.paragraphs[0].paragraph_format.space_after = Pt(0)

    table = header.add_table(rows=2, cols=3, width=Inches(7.15))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (1.65, 3.93, 1.57)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            base.set_width(cell, widths[index])
            _margins(cell, 18, 22, 18, 22)
    _row_height(table.rows[0], 0.53)
    _row_height(table.rows[1], 0.50)

    logo_cell = table.cell(0, 0)
    p = base.clear_cell(logo_cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    logo = base.image_path(base.image_for(report, base.LOGO))
    if logo:
        p.add_run().add_picture(str(logo), width=Inches(1.20))
    logo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    _cell_text(table.cell(0, 1), "Unidad Titulación y Eficiencia Terminal", 8.0)
    _right_header_table(table.cell(0, 2), report, widths[2] - 0.02)
    _cell_text(
        table.cell(1, 0),
        f"Fecha de Elaboración:\n{base.format_date(report.get('elaboration_date'))}",
        6.5,
    )
    _cell_text(table.cell(1, 1), base.header_title(report), 6.5, True)

    page_cell = table.cell(1, 2)
    p = base.clear_cell(page_cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _font(p.add_run("Página "), 7.0)
    base.field(p, "PAGE")
    _font(p.add_run(" de "), 7.0)
    base.field(p, "NUMPAGES")
    page_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _margins(page_cell)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _font(p.add_run("Página "), 8.0)
    base.field(p, "PAGE")
    _font(p.add_run(" de "), 8.0)
    base.field(p, "NUMPAGES")


def _signature_top(cell: Any, report: dict[str, Any], label: str, section_name: str) -> None:
    p = base.clear_cell(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    _font(p.add_run(label), 8.0)
    image = base.image_path(base.image_for(report, section_name))
    if image:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(str(image), width=Inches(1.55))
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _margins(cell, 35, 35, 20, 35)


def _label_value(cell: Any, label: str, value: str, line_break: bool = False) -> None:
    p = base.clear_cell(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    _font(p.add_run(label), 7.8, True)
    if line_break:
        p.add_run("\n")
    _font(p.add_run(value), 7.8)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _margins(cell, 20, 55, 20, 35)


def cover_docx(document: Any, report: dict[str, Any]) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Inches(1.55)
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(p.add_run("Informe Final Del Proceso De Titulación."), 18, True)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(p.add_run(str(report.get("period", ""))), 18, True)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Inches(2.85)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(p.add_run(f"Modalidad {base.modality(report)}"), 18, True)

    table = document.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            base.set_width(cell, 2.38)
    _row_height(table.rows[0], 1.60)
    _row_height(table.rows[1], 0.30)
    _row_height(table.rows[2], 0.58)

    people = [
        ("ELABORADO POR:", base.SIG_PREPARED, str(report.get("prepared_by") or ""), str(report.get("prepared_role") or "")),
        ("REVISADO POR:", base.SIG_REVIEWED, str(report.get("reviewed_by") or ""), str(report.get("reviewed_role") or "")),
        ("APROBADO POR:", base.SIG_APPROVED, str(report.get("approved_by") or ""), str(report.get("approved_role") or "")),
    ]
    for index, (label, section_name, name, role) in enumerate(people):
        _signature_top(table.cell(0, index), report, label, section_name)
        _label_value(table.cell(1, index), "NOMBRE: ", name)
        _label_value(table.cell(2, index), "CARGO:", role, True)
    document.add_page_break()


def _wrap(canvas: Any, text: str, width: float, font: str, size: float, limit: int) -> list[str]:
    lines: list[str] = []
    for explicit in str(text).split("\n"):
        current = ""
        for word in explicit.split():
            candidate = f"{current} {word}".strip()
            if canvas.stringWidth(candidate, font, size) <= width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = word
        lines.append(current)
    return lines[:limit]


def _box_text(canvas: Any, text: str, x: float, y: float, width: float, height: float, size: float, bold: bool = False, limit: int = 3) -> None:
    font = "Helvetica-Bold" if bold else "Helvetica"
    lines = _wrap(canvas, text, width - 8, font, size, limit)
    leading = size + 1.5
    block = max(leading, len(lines) * leading)
    baseline = y + (height + block) / 2 - leading
    canvas.setFont(font, size)
    for index, line in enumerate(lines):
        canvas.drawCentredString(x + width / 2, baseline - index * leading, line)


def draw_header(canvas: Any, report: dict[str, Any], page: int, pages: int) -> None:
    page_width, page_height = A4
    x = 1.2 * cm
    top = page_height - 0.65 * cm
    top_row = 1.18 * cm
    bottom_row = 1.15 * cm
    bottom = top - top_row - bottom_row
    total_width = page_width - 2.4 * cm
    left = 4.35 * cm
    right = 4.10 * cm
    middle = total_width - left - right
    x_middle = x + left
    x_right = x_middle + middle

    canvas.saveState()
    canvas.setLineWidth(0.75)
    canvas.rect(x, bottom, total_width, top_row + bottom_row)
    canvas.line(x, top - top_row, x + total_width, top - top_row)
    canvas.line(x_middle, bottom, x_middle, top)
    canvas.line(x_right, bottom, x_right, top)
    version_line = top - top_row * 0.66
    canvas.line(x_right, version_line, x + total_width, version_line)

    logo = base.image_path(base.image_for(report, base.LOGO))
    if logo:
        canvas.drawImage(
            str(logo),
            x + 0.18 * cm,
            top - top_row + 0.12 * cm,
            width=left - 0.36 * cm,
            height=top_row - 0.24 * cm,
            preserveAspectRatio=True,
            anchor="c",
            mask="auto",
        )

    _box_text(canvas, "Unidad Titulación y Eficiencia Terminal", x_middle, top - top_row, middle, top_row, 8.0, False, 2)
    _box_text(canvas, f"Código:\n{report.get('code', '')}", x_right, version_line, right, top - version_line, 6.7, False, 2)
    _box_text(canvas, f"Versión: {report.get('version', '1.0')}", x_right, top - top_row, right, version_line - (top - top_row), 6.7, False, 1)
    _box_text(canvas, f"Fecha de Elaboración:\n{base.format_date(report.get('elaboration_date'))}", x, bottom, left, bottom_row, 6.6, False, 2)
    _box_text(canvas, base.header_title(report), x_middle, bottom, middle, bottom_row, 6.6, True, 3)
    _box_text(canvas, f"Página {page} de {pages}", x_right, bottom, right, bottom_row, 7.1, False, 1)

    canvas.setFont("Helvetica", 8)
    canvas.drawRightString(page_width - 1.25 * cm, 0.62 * cm, f"Página {page} de {pages}")
    canvas.restoreState()


def _signature_top_flowables(report: dict[str, Any], label: str, section_name: str, styles: Any) -> list[Any]:
    style = ParagraphStyle(
        f"SignatureLabel{section_name}",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8.5,
        leading=10,
        alignment=TA_CENTER,
        spaceAfter=4,
    )
    items: list[Any] = [Paragraph(label, style), Spacer(1, 0.15 * cm)]
    image = base.image_path(base.image_for(report, section_name))
    items.append(base.fit_image(image, 4.6 * cm, 1.75 * cm) if image else Spacer(1, 1.75 * cm))
    return items


def _name(name: str, styles: Any) -> Paragraph:
    style = ParagraphStyle("SignatureNameV2", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.2, leading=9.2, alignment=TA_LEFT)
    return Paragraph(f"<b>NOMBRE:</b> {name}", style)


def _role(role: str, styles: Any) -> Paragraph:
    style = ParagraphStyle("SignatureRoleV2", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.2, leading=9.2, alignment=TA_LEFT)
    return Paragraph(f"<b>CARGO:</b><br/>{role}", style)


def cover_pdf(report: dict[str, Any], styles: Any) -> list[Any]:
    title = ParagraphStyle("CoverTitleV2", parent=styles["Title"], alignment=TA_CENTER, fontName="Helvetica-Bold", fontSize=17, leading=20, spaceAfter=2)
    period = ParagraphStyle("CoverPeriodV2", parent=title, fontSize=16, leading=19, spaceAfter=2)
    mode = ParagraphStyle("CoverModeV2", parent=title, fontSize=16, leading=19, spaceAfter=0)
    people = [
        ("ELABORADO POR:", base.SIG_PREPARED, str(report.get("prepared_by") or ""), str(report.get("prepared_role") or "")),
        ("REVISADO POR:", base.SIG_REVIEWED, str(report.get("reviewed_by") or ""), str(report.get("reviewed_role") or "")),
        ("APROBADO POR:", base.SIG_APPROVED, str(report.get("approved_by") or ""), str(report.get("approved_role") or "")),
    ]
    data = [
        [_signature_top_flowables(report, label, section_name, styles) for label, section_name, _, _ in people],
        [_name(name, styles) for _, _, name, _ in people],
        [_role(role, styles) for _, _, _, role in people],
    ]
    table = Table(data, colWidths=[5.55 * cm] * 3, rowHeights=[3.25 * cm, 0.65 * cm, 1.15 * cm])
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.75, colors.black),
        ("VALIGN", (0, 0), (-1, 0), "TOP"),
        ("VALIGN", (0, 1), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, 0), 4),
        ("RIGHTPADDING", (0, 0), (-1, 0), 4),
        ("TOPPADDING", (0, 0), (-1, 0), 4),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 4),
        ("LEFTPADDING", (0, 1), (-1, -1), 5),
        ("RIGHTPADDING", (0, 1), (-1, -1), 4),
        ("TOPPADDING", (0, 1), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 1), (-1, -1), 2),
    ]))
    return [
        Spacer(1, 4.1 * cm),
        Paragraph("Informe Final Del Proceso De Titulación.", title),
        Paragraph(str(report.get("period", "")), period),
        Paragraph(f"Modalidad {base.modality(report)}", mode),
        Spacer(1, 7.15 * cm),
        table,
        PageBreak(),
    ]


def install() -> None:
    base.setup_header = setup_header
    base.cover_docx = cover_docx
    base.draw_header = draw_header
    base.cover_pdf = cover_pdf
