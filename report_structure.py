from __future__ import annotations

import html
import re
import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import PageBreak, Paragraph, Spacer, Table, TableStyle
from reportlab.platypus.doctemplate import SimpleDocTemplate
from reportlab.platypus.tableofcontents import TableOfContents

import app as core
import institutional_export as base
import process_export
from analytics import summary
from db import connection, utcnow
from optional_content import is_present
from process_service import get_projects, get_schedules
from roster_service import REQUIREMENTS, get_report_roster


INTRODUCTION_PARAGRAPHS = (
    "El presente Informe Final del Proceso de Titulación tiene como propósito presentar, consolidar y analizar la información correspondiente al período académico {period}, en la modalidad {modality}, bajo la coordinación de la Unidad de Titulación y Eficiencia Terminal del Instituto Tecnológico Superior Quito Metropolitano. Su elaboración permite documentar el desarrollo del proceso, los resultados obtenidos y las principales situaciones identificadas durante las diferentes etapas de titulación.",
    "El proceso de titulación constituye una etapa fundamental en la formación académica de los estudiantes, debido a que permite verificar la integración de los conocimientos, habilidades y competencias adquiridas durante su carrera. En este contexto, el informe recoge información relacionada con {scope}, considerando la participación de los estudiantes, el cumplimiento de las actividades planificadas, la aplicación de las evaluaciones correspondientes y la consolidación de los resultados finales.",
    "La información presentada fue obtenida de los registros académicos institucionales, las plataformas utilizadas durante el proceso, las matrices de seguimiento y los documentos de respaldo generados por las diferentes áreas responsables. {cutoff_sentence} Los datos fueron revisados, organizados y consolidados procurando mantener correspondencia entre los estudiantes registrados, las calificaciones obtenidas, las evaluaciones rendidas y los estados académicos asignados.",
    "Para facilitar su interpretación, los resultados se presentan de manera diferenciada por carrera y etapa de evaluación. {results_sentence} Esta organización permite conservar la trazabilidad de la información y verificar la evolución de cada estudiante desde su registro inicial hasta la definición de su estado final.",
    "El informe también permite identificar el comportamiento general del proceso de titulación, los niveles de aprobación alcanzados, la incidencia de las evaluaciones supletorias cuando corresponda, las diferencias existentes entre carreras y las situaciones que requieren seguimiento académico o administrativo. De esta manera, los resultados no se limitan a la presentación de calificaciones, sino que constituyen una fuente de información para evaluar el cumplimiento de la planificación y reconocer oportunidades de mejora.",
    "En este sentido, el documento se establece como un instrumento técnico de seguimiento, control y evaluación institucional. La información consolidada servirá como apoyo para la toma de decisiones, la planificación de futuros períodos, el fortalecimiento del acompañamiento estudiantil y la formulación de acciones orientadas a mejorar la calidad, transparencia, eficiencia y trazabilidad del proceso de titulación.",
)

POST_RESULT_KEYS = {"analisis_estrategico", "conclusiones", "recomendaciones"}
SKIP_SECTION_KEYS = {"cronograma"}


def ensure_structure_schema() -> None:
    with connection() as conn:
        columns = {row[1] for row in conn.execute("PRAGMA table_info(reports)").fetchall()}
        if "cutoff_date" not in columns:
            conn.execute("ALTER TABLE reports ADD COLUMN cutoff_date TEXT DEFAULT ''")


def decimal(value: float | int | None) -> str:
    if value is None:
        return "—"
    return f"{float(value):.2f}".replace(".", ",")


def percentage(value: float | int | None) -> str:
    return f"{decimal(value)} %" if value is not None else "—"


def _format_date(value: Any) -> str:
    return base.format_date(value) if value else ""


def _has_grade(student: dict[str, Any]) -> bool:
    return any(
        student.get(key) is not None
        for key in (
            "ordinary_theory",
            "ordinary_practical",
            "supplementary_theory",
            "supplementary_practical",
            "source_total_course",
        )
    )


def _career_has_notes(career: dict[str, Any]) -> bool:
    return any(_has_grade(student) for student in career.get("students", []))


def _content_flags(report: dict[str, Any], report_id: int) -> dict[str, bool]:
    return {
        "complexive": any(
            _career_has_notes(career) for career in report.get("careers", [])
        ),
        "projects": bool(get_projects(report_id).get("projects")),
        "complexive_schedule": is_present(report_id, "schedule_complexive"),
        "thesis_schedule": is_present(report_id, "schedule_thesis"),
    }


def _scope(flags: dict[str, bool]) -> str:
    if flags["complexive"] and flags["projects"]:
        return "los núcleos estructurantes, el Examen Complexivo y el Trabajo de Titulación"
    if flags["complexive"]:
        return "los núcleos estructurantes y el Examen Complexivo, en sus etapas ordinaria y supletoria"
    if flags["projects"]:
        return "el desarrollo, evaluación y defensa de los trabajos de titulación"
    return "el cumplimiento de los requisitos y la planificación institucional del proceso de titulación"


def _results_sentence(flags: dict[str, bool]) -> str:
    if flags["complexive"] and flags["projects"]:
        return (
            "Se consideran los resultados del Examen Complexivo en sus fases ordinaria y "
            "supletoria, así como los resultados independientes del Trabajo de Titulación."
        )
    if flags["complexive"]:
        return (
            "Se consideran los resultados obtenidos en la fase ordinaria, los estudiantes que "
            "requirieron procesos supletorios, quienes alcanzaron la aprobación mediante estos "
            "mecanismos, los casos de reprobación final y aquellos que no completaron la evaluación."
        )
    if flags["projects"]:
        return (
            "Se consideran el trabajo escrito, la evaluación práctica, la defensa oral y la "
            "calificación final de los estudiantes registrados en esta opción de titulación."
        )
    return (
        "Se consideran la nómina registrada, el cumplimiento de requisitos y las actividades "
        "institucionales efectivamente incorporadas al informe."
    )


def introduction(report: dict[str, Any], report_id: int) -> list[str]:
    flags = _content_flags(report, report_id)
    cutoff = _format_date(report.get("cutoff_date"))
    values = {
        "period": str(report.get("period") or "").strip(),
        "modality": base.modality(report),
        "scope": _scope(flags),
        "cutoff_sentence": (
            f"La información fue consolidada con fecha de corte {cutoff}." if cutoff else ""
        ),
        "results_sentence": _results_sentence(flags),
    }
    return [" ".join(paragraph.format(**values).split()) for paragraph in INTRODUCTION_PARAGRAPHS]


def requirement_analysis(report_id: int) -> dict[str, Any] | None:
    students = get_report_roster(report_id).get("students", [])
    if not students:
        return None

    active: list[tuple[str, str]] = []
    for key, label in REQUIREMENTS:
        if any(str(student.get(key) or "").strip() for student in students):
            active.append((key, label))
    if not active:
        return None

    def classify(student: dict[str, Any]) -> str:
        values = [str(student.get(key) or "").strip().upper() for key, _ in active]
        if any(value == "NO CUMPLE" for value in values):
            return "pending"
        if any(not value for value in values):
            return "incomplete"
        return "complete" if all(value == "CUMPLE" for value in values) else "incomplete"

    total = len(students)
    classifications = [classify(student) for student in students]
    complete = classifications.count("complete")
    pending = classifications.count("pending")
    incomplete = classifications.count("incomplete")

    requirement_rows: list[dict[str, Any]] = []
    for key, label in active:
        values = [str(student.get(key) or "").strip().upper() for student in students]
        complies = sum(value == "CUMPLE" for value in values)
        requirement_rows.append(
            {
                "key": key,
                "label": label,
                "complies": complies,
                "does_not_comply": sum(value == "NO CUMPLE" for value in values),
                "blank": sum(not value for value in values),
                "percentage": round(complies / total * 100, 2),
            }
        )

    grouped: dict[str, list[dict[str, Any]]] = {}
    for student in students:
        grouped.setdefault(str(student.get("career_name") or "Sin carrera"), []).append(student)

    career_rows: list[dict[str, Any]] = []
    for career_name, career_students in sorted(grouped.items()):
        states = [classify(student) for student in career_students]
        registered = len(career_students)
        career_complete = states.count("complete")
        career_rows.append(
            {
                "career": career_name,
                "registered": registered,
                "complete": career_complete,
                "pending": states.count("pending"),
                "incomplete": states.count("incomplete"),
                "percentage": round(career_complete / registered * 100, 2),
            }
        )

    lowest = min(requirement_rows, key=lambda row: row["percentage"])
    most_pending = max(career_rows, key=lambda row: row["pending"] + row["incomplete"])
    overall = round(complete / total * 100, 2)
    narrative = (
        f"De los {total} estudiantes registrados, {complete} cumplieron integralmente los "
        f"requisitos considerados, equivalente al {percentage(overall)}. Se identificaron "
        f"{pending} estudiantes con al menos un requisito marcado como NO CUMPLE y {incomplete} "
        f"con información incompleta. El requisito con menor nivel de cumplimiento fue "
        f"{lowest['label']}, con {lowest['complies']} registros conformes de {total} "
        f"({percentage(lowest['percentage'])}). La mayor cantidad de casos pendientes o "
        f"incompletos se concentró en {most_pending['career']}, con "
        f"{most_pending['pending'] + most_pending['incomplete']} registros que requieren revisión."
    )

    return {
        "total": total,
        "complete": complete,
        "pending": pending,
        "incomplete": incomplete,
        "percentage": overall,
        "requirements": requirement_rows,
        "careers": career_rows,
        "narrative": narrative,
    }


def _valid_sections(report: dict[str, Any], post: bool) -> list[dict[str, Any]]:
    sections: list[dict[str, Any]] = []
    for section in report.get("sections", []):
        key = str(section.get("section_key") or "")
        content = str(section.get("content") or "").strip()
        if key in SKIP_SECTION_KEYS or not content:
            continue
        if (key in POST_RESULT_KEYS) == post:
            sections.append(section)
    return sections


def _set_update_fields(document: Document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def _add_toc(document: Document) -> None:
    title = document.add_paragraph("ÍNDICE", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "El índice se actualizará automáticamente al abrir el documento."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, placeholder, end])
    document.add_page_break()


def _docx_table(
    document: Document,
    headers: list[str],
    rows: list[list[Any]],
    widths: list[float],
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        base.set_width(cell, widths[index])
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(8)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            base.set_width(cells[index], widths[index])
            cells[index].text = str(value)
            for paragraph in cells[index].paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
                )
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8)
    document.add_paragraph()


def _add_docx_sections(document: Document, report: dict[str, Any], post: bool) -> None:
    for section in _valid_sections(report, post):
        document.add_heading(section["title"], level=1)
        for text in str(section.get("content") or "").split("\n"):
            if text.strip():
                document.add_paragraph(text.strip())
        if not post and section.get("section_key") == "metodologia":
            infographic = base.image_for(report, base.INFOGRAPHIC)
            if infographic:
                document.add_heading("Infografía del proceso de Examen Complexivo", level=2)
                base.add_docx_image(document, infographic, 6.3, False)


def _add_docx_requirements(document: Document, data: dict[str, Any]) -> None:
    document.add_heading("ANÁLISIS DEL CUMPLIMIENTO DE REQUISITOS DE TITULACIÓN", level=1)
    document.add_paragraph(
        f"Se analizaron {data['total']} estudiantes. {data['complete']} presentaron cumplimiento "
        f"integral, {data['pending']} registraron al menos un requisito pendiente y "
        f"{data['incomplete']} presentaron información incompleta. El cumplimiento integral fue "
        f"del {percentage(data['percentage'])}."
    )
    _docx_table(
        document,
        ["Requisito", "Cumple", "No cumple", "Sin información", "Cumplimiento"],
        [
            [
                row["label"],
                row["complies"],
                row["does_not_comply"],
                row["blank"],
                percentage(row["percentage"]),
            ]
            for row in data["requirements"]
        ],
        [2.55, 0.8, 0.85, 1.05, 1.05],
    )
    document.add_heading("Cumplimiento por carrera", level=2)
    _docx_table(
        document,
        ["Carrera", "Registrados", "Completos", "Pendientes", "Sin información", "Cumplimiento"],
        [
            [
                row["career"],
                row["registered"],
                row["complete"],
                row["pending"],
                row["incomplete"],
                percentage(row["percentage"]),
            ]
            for row in data["careers"]
        ],
        [2.45, 0.75, 0.75, 0.75, 0.9, 0.95],
    )
    document.add_paragraph(data["narrative"])


def _add_docx_complexive(document: Document, report: dict[str, Any]) -> None:
    careers = [career for career in report.get("careers", []) if _career_has_notes(career)]
    if not careers:
        return
    document.add_heading("RESULTADOS DEL EXAMEN COMPLEXIVO", level=1)
    for career in careers:
        document.add_page_break()
        document.add_heading(career["name"], level=2)
        diagram = base.image_for(report, base.NUCLEI, int(career["id"]))
        if diagram:
            base.add_docx_image(document, diagram, 4.7, False)
        base.legacy._add_docx_phase(
            document, career, "ordinario", "Resultados de la evaluación ordinaria"
        )
        supplementary = summary(career["students"], "supletorio")
        if supplementary["rows"]:
            base.legacy._add_docx_phase(
                document, career, "supletorio", "Resultados de la evaluación supletoria"
            )
        else:
            document.add_heading("Resultados de la evaluación supletoria", level=3)
            document.add_paragraph(
                "No se registraron estudiantes habilitados para la evaluación supletoria en esta carrera."
            )
        base.legacy._add_docx_phase(
            document, career, "consolidado", "Resultado consolidado"
        )
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temporary:
            chart = Path(temporary.name)
        try:
            base.legacy.create_chart(career, chart)
            document.add_picture(str(chart), width=Inches(6.4))
        finally:
            chart.unlink(missing_ok=True)


def _additional_images(report: dict[str, Any]) -> list[dict[str, Any]]:
    general = [
        image
        for image in report.get("general_images", [])
        if image.get("section") not in base.RESERVED and base.image_path(image)
    ]
    career_images = [
        image
        for career in report.get("careers", [])
        for image in career.get("images", [])
        if image.get("section") != base.NUCLEI and base.image_path(image)
    ]
    return general + career_images


def build_docx(report_id: int) -> Path:
    ensure_structure_schema()
    base.EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    report = base.legacy.load_report_data(report_id)
    output = base.EXPORT_DIR / f"informtit_{report_id}.docx"
    document = Document()
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(11)
    _set_update_fields(document)
    base.setup_header(document, report)
    base.cover_docx(document, report)
    _add_toc(document)

    document.add_heading("INTRODUCCIÓN", level=1)
    for paragraph in introduction(report, report_id):
        document.add_paragraph(paragraph)

    _add_docx_sections(document, report, post=False)

    requirements = requirement_analysis(report_id)
    if requirements:
        _add_docx_requirements(document, requirements)

    schedules = get_schedules(report_id)
    if is_present(report_id, "schedule_complexive") and schedules.get("complexive"):
        process_export._add_docx_schedule(
            document,
            "Cronograma de Núcleos y Examen Complexivo",
            schedules["complexive"],
            False,
        )
    if is_present(report_id, "schedule_thesis") and schedules.get("thesis"):
        process_export._add_docx_schedule(
            document,
            "Cronograma del Trabajo de Titulación",
            schedules["thesis"],
            True,
        )

    _add_docx_complexive(document, report)
    if get_projects(report_id).get("projects"):
        process_export._add_docx_projects(document, report_id)

    _add_docx_sections(document, report, post=True)

    images = _additional_images(report)
    if images:
        document.add_heading("ANEXOS", level=1)
        for image in images:
            base.add_docx_image(document, image)

    document.save(output)
    return output


class TocDocTemplate(SimpleDocTemplate):
    def afterFlowable(self, flowable: Any) -> None:
        if not isinstance(flowable, Paragraph):
            return
        levels = {"Heading1": 0, "Heading2": 1, "Heading3": 2}
        if flowable.style.name not in levels:
            return
        level = levels[flowable.style.name]
        text = flowable.getPlainText()
        key = f"toc_{self.seq.nextf('heading')}"
        self.canv.bookmarkPage(key)
        self.canv.addOutlineEntry(text, key, level=level, closed=False)
        self.notify("TOCEntry", (level, text, self.page, key))


def _pdf_table(headers: list[Any], rows: list[list[Any]], widths: list[float]) -> Table:
    table = Table([headers] + rows, repeatRows=1, colWidths=widths)
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#244a73")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (1, 1), (-1, -1), "CENTER"),
            ]
        )
    )
    return table


def _add_pdf_sections(story: list[Any], report: dict[str, Any], styles: Any, post: bool) -> None:
    for section in _valid_sections(report, post):
        story.append(Paragraph(section["title"], styles["Heading1"]))
        for text in str(section.get("content") or "").split("\n"):
            if text.strip():
                story.append(Paragraph(text.strip(), styles["BodyJustified"]))
        if not post and section.get("section_key") == "metodologia":
            image = base.image_path(base.image_for(report, base.INFOGRAPHIC))
            if image:
                story += [
                    Paragraph("Infografía del proceso de Examen Complexivo", styles["Heading2"]),
                    base.fit_image(image, 16.5 * cm, 19.5 * cm),
                    Spacer(1, 0.3 * cm),
                ]


def _add_pdf_requirements(story: list[Any], data: dict[str, Any], styles: Any) -> None:
    story.append(Paragraph("ANÁLISIS DEL CUMPLIMIENTO DE REQUISITOS DE TITULACIÓN", styles["Heading1"]))
    story.append(
        Paragraph(
            f"Se analizaron {data['total']} estudiantes. {data['complete']} presentaron cumplimiento "
            f"integral, {data['pending']} registraron al menos un requisito pendiente y "
            f"{data['incomplete']} presentaron información incompleta. El cumplimiento integral fue "
            f"del {percentage(data['percentage'])}.",
            styles["BodyJustified"],
        )
    )
    story += [
        Spacer(1, 0.2 * cm),
        _pdf_table(
            ["Requisito", "Cumple", "No cumple", "Sin información", "Cumplimiento"],
            [
                [
                    Paragraph(html.escape(row["label"]), styles["BodyText"]),
                    row["complies"],
                    row["does_not_comply"],
                    row["blank"],
                    percentage(row["percentage"]),
                ]
                for row in data["requirements"]
            ],
            [6.4 * cm, 2.1 * cm, 2.2 * cm, 2.4 * cm, 2.5 * cm],
        ),
        Spacer(1, 0.35 * cm),
        Paragraph("Cumplimiento por carrera", styles["Heading2"]),
        _pdf_table(
            ["Carrera", "Registrados", "Completos", "Pendientes", "Sin información", "Cumplimiento"],
            [
                [
                    Paragraph(html.escape(row["career"]), styles["BodyText"]),
                    row["registered"],
                    row["complete"],
                    row["pending"],
                    row["incomplete"],
                    percentage(row["percentage"]),
                ]
                for row in data["careers"]
            ],
            [5.7 * cm, 2.0 * cm, 2.0 * cm, 2.0 * cm, 2.3 * cm, 2.5 * cm],
        ),
        Spacer(1, 0.3 * cm),
        Paragraph(data["narrative"], styles["BodyJustified"]),
    ]


def _add_pdf_complexive(
    story: list[Any],
    report: dict[str, Any],
    styles: Any,
    temp_paths: list[Path],
) -> None:
    careers = [career for career in report.get("careers", []) if _career_has_notes(career)]
    if not careers:
        return
    story.append(Paragraph("RESULTADOS DEL EXAMEN COMPLEXIVO", styles["Heading1"]))
    for career in careers:
        story += [PageBreak(), Paragraph(career["name"], styles["Heading2"])]
        diagram = base.image_path(base.image_for(report, base.NUCLEI, int(career["id"])))
        if diagram:
            story += [base.fit_image(diagram, 11.5 * cm, 8.5 * cm), Spacer(1, 0.3 * cm)]
        for phase, heading in (
            ("ordinario", "Resultados de la evaluación ordinaria"),
            ("supletorio", "Resultados de la evaluación supletoria"),
            ("consolidado", "Resultado consolidado"),
        ):
            phase_data = summary(career["students"], phase)
            story.append(Paragraph(heading, styles["Heading3"]))
            if phase == "supletorio" and not phase_data["rows"]:
                story.append(
                    Paragraph(
                        "No se registraron estudiantes habilitados para la evaluación supletoria en esta carrera.",
                        styles["BodyJustified"],
                    )
                )
                continue
            analysis = career["analyses"].get(phase, {})
            table_data = base.legacy._pdf_table_data(phase_data, phase)
            story += [
                Paragraph(
                    analysis.get("text_before")
                    or base.legacy._default_before(career["name"], phase, phase_data),
                    styles["BodyJustified"],
                ),
                _pdf_table(
                    table_data[0],
                    table_data[1:],
                    base.legacy._pdf_col_widths(phase),
                ),
                Spacer(1, 0.2 * cm),
                Paragraph(
                    analysis.get("text_after") or base.legacy._default_after(phase_data),
                    styles["BodyJustified"],
                ),
            ]
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temporary:
            chart = Path(temporary.name)
        base.legacy.create_chart(career, chart)
        temp_paths.append(chart)
        story += [base.fit_image(chart, 16 * cm, 9 * cm), Spacer(1, 0.3 * cm)]


def build_pdf(report_id: int) -> Path:
    ensure_structure_schema()
    base.EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    report = base.legacy.load_report_data(report_id)
    output = base.EXPORT_DIR / f"informtit_{report_id}.pdf"
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            "BodyJustified",
            parent=styles["BodyText"],
            alignment=TA_JUSTIFY,
            leading=15,
            spaceAfter=7,
        )
    )
    for heading in (styles["Heading1"], styles["Heading2"], styles["Heading3"]):
        heading.keepWithNext = True

    story: list[Any] = base.cover_pdf(report, styles)
    story.append(Paragraph("ÍNDICE", styles["Title"]))
    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle("TOC1", fontName="Helvetica-Bold", fontSize=10, leading=14, spaceBefore=4),
        ParagraphStyle("TOC2", fontName="Helvetica", fontSize=9, leading=12, leftIndent=14),
        ParagraphStyle("TOC3", fontName="Helvetica", fontSize=8, leading=11, leftIndent=28),
    ]
    story += [toc, PageBreak(), Paragraph("INTRODUCCIÓN", styles["Heading1"])]
    for paragraph in introduction(report, report_id):
        story.append(Paragraph(paragraph, styles["BodyJustified"]))

    _add_pdf_sections(story, report, styles, post=False)

    requirements = requirement_analysis(report_id)
    if requirements:
        _add_pdf_requirements(story, requirements, styles)

    schedules = get_schedules(report_id)
    if is_present(report_id, "schedule_complexive") and schedules.get("complexive"):
        story += process_export._pdf_schedule(
            "Cronograma de Núcleos y Examen Complexivo",
            schedules["complexive"],
            False,
            styles,
        )
    if is_present(report_id, "schedule_thesis") and schedules.get("thesis"):
        story += process_export._pdf_schedule(
            "Cronograma del Trabajo de Titulación",
            schedules["thesis"],
            True,
            styles,
        )

    temp_paths: list[Path] = []
    _add_pdf_complexive(story, report, styles, temp_paths)
    if get_projects(report_id).get("projects"):
        story += process_export._pdf_projects(report_id, styles)

    _add_pdf_sections(story, report, styles, post=True)

    images = _additional_images(report)
    if images:
        story += [PageBreak(), Paragraph("ANEXOS", styles["Heading1"])]
        for image in images:
            path = base.image_path(image)
            if path:
                story += [base.fit_image(path, 16.5 * cm, 20 * cm), Spacer(1, 0.3 * cm)]

    document = TocDocTemplate(
        str(output),
        pagesize=A4,
        rightMargin=1.45 * cm,
        leftMargin=1.45 * cm,
        topMargin=3.4 * cm,
        bottomMargin=1.35 * cm,
        title=report["name"],
    )
    try:
        document.multiBuild(
            story,
            canvasmaker=lambda *args, **kwargs: base.NumberedCanvas(
                *args, report=report, **kwargs
            ),
        )
    finally:
        for path in temp_paths:
            path.unlink(missing_ok=True)
    return output


def _install_cutoff_route() -> None:
    if getattr(core.InformtitHandler, "_informtit_cutoff_route", False):
        return
    original = core.InformtitHandler._handle_api_write

    def handle(self: Any, method: str, path: str, payload: dict[str, Any]) -> None:
        match = re.fullmatch(r"/api/reports/(\d+)", path)
        if method == "PUT" and match and "cutoff_date" in payload:
            report_id = int(match.group(1))
            remaining = dict(payload)
            cutoff_date = str(remaining.pop("cutoff_date", "") or "").strip()
            with connection() as conn:
                conn.execute(
                    "UPDATE reports SET cutoff_date=?, updated_at=? WHERE id=?",
                    (cutoff_date, utcnow(), report_id),
                )
            if remaining:
                original(self, method, path, remaining)
            else:
                self._send_json({"ok": True})
            return
        original(self, method, path, payload)

    core.InformtitHandler._handle_api_write = handle
    core.InformtitHandler._informtit_cutoff_route = True


def install() -> None:
    ensure_structure_schema()
    _install_cutoff_route()
    core.build_docx = build_docx
    core.build_pdf = build_pdf
