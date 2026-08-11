from __future__ import annotations

from typing import Any

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

import report_enhancements
import report_quality
import report_structure


def _shade(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _repeat_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def docx_table(document: Any, headers: list[str], rows: list[list[Any]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    _repeat_header(table.rows[0])

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        report_quality.base.set_width(cell, widths[index])
        cell.text = report_enhancements.public_text(header)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _shade(cell, "244A73")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(3)
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(8.2)
                run.font.color.rgb = RGBColor(255, 255, 255)

    for row_index, values in enumerate(rows, start=1):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            report_quality.base.set_width(cells[index], widths[index])
            cells[index].text = report_enhancements.public_text(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2 == 0:
                _shade(cells[index], "F4F7FA")
            for paragraph in cells[index].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8)
    document.add_paragraph()


def install() -> None:
    report_enhancements._docx_table_pretty = docx_table
    report_structure._docx_table = docx_table
