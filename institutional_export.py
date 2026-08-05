from __future__ import annotations

import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas as pdfcanvas
from reportlab.platypus import Image as RLImage, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

import report_service as legacy
from analytics import summary

EXPORT_DIR = legacy.EXPORT_DIR
UPLOAD_DIR = legacy.UPLOAD_DIR
LOGO = "logo_institucional"
SIG_PREPARED = "firma_elaborado"
SIG_REVIEWED = "firma_revisado"
SIG_APPROVED = "firma_aprobado"
INFOGRAPHIC = "infografia_complexivo"
NUCLEI = "diagrama_nucleos"
RESERVED = {LOGO, SIG_PREPARED, SIG_REVIEWED, SIG_APPROVED, INFOGRAPHIC}
MONTHS = ["", "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]


def modality(report: dict[str, Any]) -> str:
    return "Online" if report.get("modality") == "en_linea" else "Presencial"


def header_title(report: dict[str, Any]) -> str:
    return f"Informe Final Del Proceso De Titulación. {report.get('period', '')} - Modalidad {modality(report)}"


def format_date(value: Any) -> str:
    text = str(value or "").strip()
    parts = text.split("-")
    if len(parts) == 3 and all(part.isdigit() for part in parts):
        year, month, day = map(int, parts)
        if 1 <= month <= 12:
            return f"{day}-{MONTHS[month]}-{year}"
    return text


def image_path(image: dict[str, Any] | None) -> Path | None:
    if not image:
        return None
    path = UPLOAD_DIR / str(image.get("filename") or "")
    return path if path.exists() else None


def image_for(report: dict[str, Any], section: str, career_id: int | None = None) -> dict[str, Any] | None:
    if career_id is None:
        images = report.get("general_images", [])
    else:
        images = next((c.get("images", []) for c in report.get("careers", []) if int(c["id"]) == int(career_id)), [])
    matches = [item for item in images if item.get("section") == section]
    return matches[-1] if matches else None


# ---------------- DOCX ----------------
def field(paragraph: Any, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = instruction
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = "1"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def clear_cell(cell: Any) -> Any:
    cell.text = ""
    return cell.paragraphs[0]


def cell_text(cell: Any, text: str, size: float = 8, bold: bool = False) -> None:
    p = clear_cell(cell); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = bold; r.font.name = "Arial"; r.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_width(cell: Any, inches: float) -> None:
    cell.width = Inches(inches)
    tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW"); cell._tc.get_or_add_tcPr().append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440))); tc_w.set(qn("w:type"), "dxa")


def setup_header(document: Document, report: dict[str, Any]) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1.35); section.bottom_margin = Inches(.7)
    section.left_margin = Inches(.7); section.right_margin = Inches(.7)
    section.header_distance = Inches(.18); section.footer_distance = Inches(.25)
    header = section.header; header.is_linked_to_previous = False
    table = header.add_table(rows=2, cols=3, width=Inches(7)); table.style = "Table Grid"; table.autofit = False
    widths = (1.75, 3.75, 1.5)
    for row in table.rows:
        for i, cell in enumerate(row.cells): set_width(cell, widths[i])
    logo = image_path(image_for(report, LOGO))
    p = clear_cell(table.cell(0, 0)); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo: p.add_run().add_picture(str(logo), width=Inches(1.55))
    else: p.add_run("LOGO INSTITUCIONAL\nNO CARGADO").bold = True
    cell_text(table.cell(0, 1), "Unidad Titulación y Eficiencia Terminal", 9)
    p = clear_cell(table.cell(0, 2)); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Código:\n").bold = True; p.add_run(str(report.get("code") or "")); p.add_run("\nVersión: ").bold = True; p.add_run(str(report.get("version") or "1.0"))
    for r in p.runs: r.font.name = "Arial"; r.font.size = Pt(7.5)
    cell_text(table.cell(1, 0), f"Fecha de Elaboración:\n{format_date(report.get('elaboration_date'))}", 7.5)
    cell_text(table.cell(1, 1), header_title(report), 7.5, True)
    p = clear_cell(table.cell(1, 2)); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Página "); field(p, "PAGE"); p.add_run(" de "); field(p, "NUMPAGES")
    for r in p.runs: r.font.name = "Arial"; r.font.size = Pt(8)
    footer = section.footer; footer.is_linked_to_previous = False
    p = footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Página "); field(p, "PAGE"); p.add_run(" de "); field(p, "NUMPAGES")


def signature_cell(cell: Any, report: dict[str, Any], label: str, section: str, name: str, role: str) -> None:
    p = clear_cell(cell); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label); r.bold = True; r.font.size = Pt(7.5)
    p = cell.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    path = image_path(image_for(report, section))
    if path: p.add_run().add_picture(str(path), width=Inches(1.45))
    else: p.add_run("FIRMA / QR\nNO CARGADO").bold = True
    p = cell.add_paragraph(); p.add_run("NOMBRE: ").bold = True; p.add_run(name)
    p = cell.add_paragraph(); p.add_run("CARGO: ").bold = True; p.add_run(role)
    for para in cell.paragraphs:
        for run in para.runs: run.font.name = "Arial"; run.font.size = Pt(7.5)


def cover_docx(document: Document, report: dict[str, Any]) -> None:
    p = document.add_paragraph(); p.paragraph_format.space_before = Inches(1.65); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Informe Final Del Proceso De Titulación."); r.bold = True; r.font.size = Pt(18)
    p = document.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{report.get('period', '')} Modalidad {modality(report)}"); r.bold = True; r.font.size = Pt(18)
    p = document.add_paragraph(); p.paragraph_format.space_after = Inches(2.5)
    table = document.add_table(rows=1, cols=3); table.style = "Table Grid"; table.autofit = False
    for cell in table.rows[0].cells: set_width(cell, 2.33)
    signature_cell(table.cell(0, 0), report, "ELABORADO POR:", SIG_PREPARED, str(report.get("prepared_by") or ""), str(report.get("prepared_role") or ""))
    signature_cell(table.cell(0, 1), report, "REVISADO POR:", SIG_REVIEWED, str(report.get("reviewed_by") or ""), str(report.get("reviewed_role") or ""))
    signature_cell(table.cell(0, 2), report, "APROBADO POR:", SIG_APPROVED, str(report.get("approved_by") or ""), str(report.get("approved_role") or ""))
    document.add_page_break()


def add_docx_image(document: Document, image: dict[str, Any] | None, width: float = 6.2, caption: bool = True) -> bool:
    path = image_path(image)
    if not path: return False
    p = document.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    if caption:
        p = document.add_paragraph(str(image.get("title") or image.get("original_name") or "Imagen")); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return True


def build_docx(report_id: int) -> Path:
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    report = legacy.load_report_data(report_id); output = EXPORT_DIR / f"informtit_{report_id}.docx"
    document = Document(); document.styles["Normal"].font.name = "Arial"; document.styles["Normal"].font.size = Pt(11)
    setup_header(document, report); cover_docx(document, report)
    document.add_heading("Contenido", level=1)
    for section in report["sections"]: document.add_paragraph(section["title"], style="List Number")
    for career in report["careers"]: document.add_paragraph(f"Resultados: {career['name']}", style="List Number")
    document.add_page_break()
    inf_added = False
    for section in report["sections"]:
        if section.get("section_key") == "metodologia" and not inf_added:
            inf = image_for(report, INFOGRAPHIC)
            if inf:
                document.add_heading("Infografía Proceso de Examen Complexivo", level=1); add_docx_image(document, inf, 6.3, False); document.add_page_break()
            inf_added = True
        document.add_heading(section["title"], level=1)
        for text in str(section.get("content") or "").split("\n"):
            if text.strip(): document.add_paragraph(text.strip())
    for image in report["general_images"]:
        if image.get("section") not in RESERVED: add_docx_image(document, image)
    for career in report["careers"]:
        document.add_page_break(); document.add_heading(career["name"], level=1)
        add_docx_image(document, image_for(report, NUCLEI, int(career["id"])), 4.7, False)
        legacy._add_docx_phase(document, career, "ordinario", "Resultados de la evaluación ordinaria")
        legacy._add_docx_phase(document, career, "supletorio", "Resultados de la evaluación supletoria")
        legacy._add_docx_phase(document, career, "consolidado", "Resultado consolidado")
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp: chart = Path(temp.name)
        legacy.create_chart(career, chart); document.add_picture(str(chart), width=Inches(6.4)); chart.unlink(missing_ok=True)
        for image in career["images"]:
            if image.get("section") != NUCLEI: add_docx_image(document, image)
    document.save(output); return output


# ---------------- PDF ----------------
def fit_image(path: Path, max_w: float, max_h: float) -> RLImage:
    image = RLImage(str(path)); scale = min(max_w / image.imageWidth, max_h / image.imageHeight)
    image.drawWidth = image.imageWidth * scale; image.drawHeight = image.imageHeight * scale; return image


def centered(canvas: pdfcanvas.Canvas, text: str, x: float, y: float, width: float, size: float = 7, bold: bool = False, max_lines: int = 3) -> None:
    font = "Helvetica-Bold" if bold else "Helvetica"; canvas.setFont(font, size)
    lines, line = [], ""
    for word in text.split():
        candidate = f"{line} {word}".strip()
        if canvas.stringWidth(candidate, font, size) <= width - 8: line = candidate
        else:
            if line: lines.append(line)
            line = word
    if line: lines.append(line)
    lines = lines[:max_lines]; leading = size + 2; start = y + (max_lines - len(lines)) * leading / 2
    for i, item in enumerate(lines): canvas.drawCentredString(x + width / 2, start - i * leading, item)


def draw_header(canvas: pdfcanvas.Canvas, report: dict[str, Any], page: int, pages: int) -> None:
    width, height = A4; x = 1.25 * cm; top = height - .75 * cm; row = .95 * cm
    total = width - 2.5 * cm; left = 4.25 * cm; right = 4.15 * cm; middle = total - left - right; bottom = top - 2 * row
    canvas.saveState(); canvas.setLineWidth(.7); canvas.rect(x, bottom, total, 2 * row); canvas.line(x, top-row, x+total, top-row); canvas.line(x+left, bottom, x+left, top); canvas.line(x+left+middle, bottom, x+left+middle, top)
    logo = image_path(image_for(report, LOGO))
    if logo: canvas.drawImage(str(logo), x+.1*cm, top-row+.08*cm, width=left-.2*cm, height=row-.16*cm, preserveAspectRatio=True, anchor="c", mask="auto")
    else: centered(canvas, "LOGO INSTITUCIONAL NO CARGADO", x, top-row+.27*cm, left, 6.5, True)
    centered(canvas, "Unidad Titulación y Eficiencia Terminal", x+left, top-row+.27*cm, middle, 8.2)
    centered(canvas, f"Código: {report.get('code','')}  Versión: {report.get('version','1.0')}", x+left+middle, top-row+.2*cm, right, 6.5)
    centered(canvas, f"Fecha de Elaboración: {format_date(report.get('elaboration_date'))}", x, bottom+.25*cm, left, 6.8, False, 2)
    centered(canvas, header_title(report), x+left, bottom+.18*cm, middle, 6.8, True)
    centered(canvas, f"Página {page} de {pages}", x+left+middle, bottom+.28*cm, right, 7.2, False, 1)
    canvas.setFont("Helvetica", 8); canvas.drawRightString(width-1.35*cm, .65*cm, f"Página {page} de {pages}"); canvas.restoreState()


class NumberedCanvas(pdfcanvas.Canvas):
    def __init__(self, *args: Any, report: dict[str, Any], **kwargs: Any):
        super().__init__(*args, **kwargs); self.states: list[dict[str, Any]] = []; self.report = report
    def showPage(self) -> None:
        self.states.append(dict(self.__dict__)); self._startPage()
    def save(self) -> None:
        pages = len(self.states)
        for page, state in enumerate(self.states, 1):
            self.__dict__.update(state); draw_header(self, self.report, page, pages); pdfcanvas.Canvas.showPage(self)
        pdfcanvas.Canvas.save(self)


def signature_items(report: dict[str, Any], label: str, section: str, name: str, role: str, styles: Any) -> list[Any]:
    style = ParagraphStyle("Sig", parent=styles["BodyText"], fontSize=7, leading=8, alignment=TA_CENTER)
    items: list[Any] = [Paragraph(label, style)]; path = image_path(image_for(report, section))
    items.append(fit_image(path, 4.2*cm, 2.2*cm) if path else Paragraph("FIRMA / QR<br/>NO CARGADO", style))
    items += [Paragraph(f"<b>NOMBRE:</b> {name}", style), Paragraph(f"<b>CARGO:</b> {role}", style)]; return items


def cover_pdf(report: dict[str, Any], styles: Any) -> list[Any]:
    title = ParagraphStyle("Cover", parent=styles["Title"], alignment=TA_CENTER, fontName="Helvetica-Bold", fontSize=17, leading=21)
    story: list[Any] = [Spacer(1, 4*cm), Paragraph("Informe Final Del Proceso De Titulación.", title), Paragraph(f"{report.get('period','')} Modalidad {modality(report)}", title), Spacer(1, 7*cm)]
    data = [[
        signature_items(report, "ELABORADO POR:", SIG_PREPARED, str(report.get("prepared_by") or ""), str(report.get("prepared_role") or ""), styles),
        signature_items(report, "REVISADO POR:", SIG_REVIEWED, str(report.get("reviewed_by") or ""), str(report.get("reviewed_role") or ""), styles),
        signature_items(report, "APROBADO POR:", SIG_APPROVED, str(report.get("approved_by") or ""), str(report.get("approved_role") or ""), styles),
    ]]
    table = Table(data, colWidths=[5.55*cm]*3); table.setStyle(TableStyle([("GRID",(0,0),(-1,-1),.7,colors.black),("VALIGN",(0,0),(-1,-1),"TOP"),("PADDING",(0,0),(-1,-1),5)]))
    story += [table, PageBreak()]; return story


def build_pdf(report_id: int) -> Path:
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    report = legacy.load_report_data(report_id); output = EXPORT_DIR / f"informtit_{report_id}.pdf"
    styles = getSampleStyleSheet(); styles.add(ParagraphStyle("BodyJustified", parent=styles["BodyText"], alignment=TA_JUSTIFY, leading=15))
    story = cover_pdf(report, styles); temp_paths: list[Path] = []; inf_added = False
    for section in report["sections"]:
        if section.get("section_key") == "metodologia" and not inf_added:
            path = image_path(image_for(report, INFOGRAPHIC))
            if path: story += [Paragraph("Infografía Proceso de Examen Complexivo", styles["Heading1"]), fit_image(path, 16.5*cm, 19.5*cm), PageBreak()]
            inf_added = True
        story.append(Paragraph(section["title"], styles["Heading1"]))
        for text in str(section.get("content") or "").split("\n"):
            if text.strip(): story += [Paragraph(text.strip(), styles["BodyJustified"]), Spacer(1,.25*cm)]
    for image in report["general_images"]:
        path = image_path(image)
        if image.get("section") not in RESERVED and path: story += [fit_image(path,16.5*cm,20*cm),Spacer(1,.3*cm)]
    for career in report["careers"]:
        story += [PageBreak(), Paragraph(career["name"], styles["Heading1"])]
        path = image_path(image_for(report, NUCLEI, int(career["id"])))
        if path: story += [fit_image(path,11.5*cm,8.5*cm),Spacer(1,.3*cm)]
        for phase, heading in [("ordinario","Resultados de la evaluación ordinaria"),("supletorio","Resultados de la evaluación supletoria"),("consolidado","Resultado consolidado")]:
            data = summary(career["students"], phase); analysis = career["analyses"].get(phase,{})
            story += [Paragraph(heading,styles["Heading2"]),Paragraph(analysis.get("text_before") or legacy._default_before(career["name"],phase,data),styles["BodyJustified"]),Spacer(1,.2*cm)]
            table = Table(legacy._pdf_table_data(data,phase),repeatRows=1,colWidths=legacy._pdf_col_widths(phase)); table.setStyle(TableStyle([("GRID",(0,0),(-1,-1),.35,colors.grey),("BACKGROUND",(0,0),(-1,0),colors.HexColor("#244a73")),("TEXTCOLOR",(0,0),(-1,0),colors.white),("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTSIZE",(0,0),(-1,-1),7),("VALIGN",(0,0),(-1,-1),"TOP")]))
            story += [table,Spacer(1,.25*cm),Paragraph(analysis.get("text_after") or legacy._default_after(data),styles["BodyJustified"]),Spacer(1,.4*cm)]
        with tempfile.NamedTemporaryFile(suffix=".png",delete=False) as temp: chart=Path(temp.name)
        legacy.create_chart(career,chart); story.append(fit_image(chart,16*cm,9*cm)); temp_paths.append(chart)
        for image in career["images"]:
            path=image_path(image)
            if image.get("section")!=NUCLEI and path: story += [fit_image(path,16.5*cm,20*cm),Spacer(1,.3*cm)]
    doc=SimpleDocTemplate(str(output),pagesize=A4,rightMargin=1.45*cm,leftMargin=1.45*cm,topMargin=3.4*cm,bottomMargin=1.35*cm,title=report["name"])
    try: doc.build(story,canvasmaker=lambda *args,**kwargs:NumberedCanvas(*args,report=report,**kwargs))
    finally:
        for path in temp_paths: path.unlink(missing_ok=True)
    return output
