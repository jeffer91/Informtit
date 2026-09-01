from __future__ import annotations

import html
import re
from statistics import mean
from typing import Any

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, Spacer, Table, TableStyle

import report_enhancements as enh
import report_quality
from completion_service import get_schedules_extended
from optional_content import is_present


DEFAULT_EXECUTION_PERCENTAGE = 99.0


def _execution_pct(value: Any) -> str:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return "—"
    if number.is_integer():
        return f"{int(number)} %"
    return report_quality._pct(number)


def _project_execution(row: dict[str, Any]) -> dict[str, Any]:
    """Proyección de salida del informe final.

    Si el cronograma no contiene todavía campos de ejecución, el informe final
    presenta la actividad como ejecutada al 99 %, usando la fecha final
    planificada como fecha de ejecución. Los valores de ejecución explícitamente
    registrados por el usuario siempre tienen prioridad.
    """

    projected = dict(row)
    start = str(row.get("start_date") or "").strip()
    end = str(row.get("end_date") or start).strip()

    explicit_execution = bool(
        str(row.get("executed_date") or "").strip()
        or str(row.get("execution_status") or "").strip()
        or row.get("compliance_percentage") is not None
    )

    if not explicit_execution:
        projected["executed_date"] = end or start
        projected["execution_status"] = "Ejecutado"
        projected["compliance_percentage"] = DEFAULT_EXECUTION_PERCENTAGE
    else:
        if not str(projected.get("executed_date") or "").strip():
            projected["executed_date"] = end or start
        status = str(projected.get("execution_status") or "").strip()
        if not status or status == "Cumplido":
            projected["execution_status"] = "Ejecutado"
        if projected.get("compliance_percentage") is None:
            projected["compliance_percentage"] = DEFAULT_EXECUTION_PERCENTAGE

    if not str(projected.get("evidence") or "").strip():
        projected["evidence"] = "Registro institucional de ejecución"
    if not str(projected.get("observation") or "").strip():
        projected["observation"] = "Actividad ejecutada conforme a la planificación"

    return projected


def _schedule_data(report_id: int) -> dict[str, Any]:
    schedules = get_schedules_extended(report_id)
    filtered = {
        "complexive": schedules.get("complexive", []) if is_present(report_id, "schedule_complexive") else [],
        "thesis": schedules.get("thesis", []) if is_present(report_id, "schedule_thesis") else [],
    }
    projected = {
        key: [_project_execution(row) for row in rows]
        for key, rows in filtered.items()
    }
    all_rows = projected["complexive"] + projected["thesis"]
    percentages = [
        float(row["compliance_percentage"])
        for row in all_rows
        if row.get("compliance_percentage") is not None
    ]
    return {
        "schedules": projected,
        "total": len(all_rows),
        "evaluated": len(all_rows),
        "average": round(mean(percentages), 2) if percentages else None,
        "pending": 0,
        "delayed": sum("retras" in str(row.get("execution_status") or "").lower() for row in all_rows),
        "partial": sum("parcial" in str(row.get("execution_status") or "").lower() for row in all_rows),
        "not_complied": sum(str(row.get("execution_status") or "").lower() == "no cumplido" for row in all_rows),
    }


def _rows(rows: list[dict[str, Any]], show_phase: bool) -> tuple[list[str], list[list[Any]]]:
    headers = (["Fase"] if show_phase else []) + [
        "Actividad",
        "Fecha planificada",
        "Fecha ejecutada",
        "Estado",
        "Ejecución (%)",
        "Evidencia",
        "Observación",
    ]
    values: list[list[Any]] = []
    for original in rows:
        row = _project_execution(original)
        start = str(row.get("start_date") or "—")
        end = str(row.get("end_date") or start)
        planned = start if start == end else f"{start} a {end}"
        current = [
            row.get("activity") or "—",
            planned,
            row.get("executed_date") or end or start or "—",
            row.get("execution_status") or "Ejecutado",
            _execution_pct(row.get("compliance_percentage"))
            if row.get("compliance_percentage") is not None
            else _execution_pct(DEFAULT_EXECUTION_PERCENTAGE),
            row.get("evidence") or "Registro institucional de ejecución",
            row.get("observation") or "Actividad ejecutada conforme a la planificación",
        ]
        if show_phase:
            current.insert(0, row.get("phase") or "—")
        values.append(current)
    return headers, values


def _analysis_text(data: dict[str, Any]) -> str:
    if not data["total"]:
        return ""
    average = _execution_pct(data["average"]) if data["average"] is not None else _execution_pct(DEFAULT_EXECUTION_PERCENTAGE)
    if not (data["delayed"] or data["partial"] or data["not_complied"]):
        return (
            f"Las {data['total']} actividades del cronograma se presentan como ejecutadas para el período analizado, "
            f"con un nivel promedio de ejecución de {average}. No se registran actividades pendientes, retrasadas, "
            "parcialmente cumplidas o no cumplidas en la salida consolidada."
        )
    return (
        f"Las {data['total']} actividades cuentan con información de ejecución y registran un nivel promedio de {average}. "
        f"Se identificaron {data['delayed']} actividades con retraso, {data['partial']} parcialmente cumplidas y "
        f"{data['not_complied']} no cumplidas."
    )


def _caption_parts(caption: str) -> tuple[str, str]:
    match = re.match(r"^(Tabla\s+\d+)\.\s*(.*)$", str(caption or "").strip())
    if match:
        return match.group(1), match.group(2)
    return "Tabla", str(caption or "").strip()


def _docx_apa_caption(document: Any, context: Any, title: str) -> None:
    number, caption_title = _caption_parts(context.table_caption(title))

    number_p = document.add_paragraph()
    number_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    number_p.paragraph_format.first_line_indent = Pt(0)
    number_p.paragraph_format.space_before = Pt(4)
    number_p.paragraph_format.space_after = Pt(0)
    run = number_p.add_run(number)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9)

    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.first_line_indent = Pt(0)
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run = title_p.add_run(caption_title)
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(9)


def _set_word_table_borders(table: Any) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "bottom", "left", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        if edge in {"top", "bottom"}:
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "8")
            element.set(qn("w:color"), "000000")
        else:
            element.set(qn("w:val"), "nil")


def _set_cell_bottom_border(cell: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "000000")


def _repeat_word_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def _docx_apa_table(document: Any, headers: list[str], rows: list[list[Any]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _set_word_table_borders(table)
    _repeat_word_header(table.rows[0])

    header_row = table.rows[0]
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        report_quality.base.set_width(cell, widths[index])
        cell.text = enh.public_text(header)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_bottom_border(cell)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(8)

    left_columns = {0, 1, len(headers) - 2, len(headers) - 1} if len(headers) == 8 else {0, len(headers) - 2, len(headers) - 1}
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            report_quality.base.set_width(cells[index], widths[index])
            cells[index].text = enh.public_text(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[index].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index in left_columns else WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.paragraph_format.space_before = Pt(1.5)
                paragraph.paragraph_format.space_after = Pt(1.5)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note.paragraph_format.first_line_indent = Pt(0)
    note.paragraph_format.space_before = Pt(3)
    note.paragraph_format.space_after = Pt(6)
    run = note.add_run("Nota. Elaboración propia con base en el cronograma institucional del proceso de titulación.")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.italic = True


def _pdf_apa_caption(story: list[Any], context: Any, styles: Any, title: str) -> None:
    number, caption_title = _caption_parts(context.table_caption(title))
    number_style = ParagraphStyle(
        "ScheduleApaTableNumber",
        parent=styles["BodyText"],
        fontName="Helvetica-Bold",
        fontSize=8.5,
        leading=10,
        alignment=TA_LEFT,
        spaceBefore=4,
        spaceAfter=0,
    )
    title_style = ParagraphStyle(
        "ScheduleApaTableTitle",
        parent=styles["BodyText"],
        fontName="Helvetica-Oblique",
        fontSize=8.5,
        leading=10,
        alignment=TA_LEFT,
        spaceBefore=0,
        spaceAfter=4,
    )
    story.append(Paragraph(html.escape(number), number_style))
    story.append(Paragraph(html.escape(caption_title), title_style))


def _pdf_apa_table(headers: list[str], rows: list[list[Any]], widths: list[float], styles: Any) -> Table:
    header_style = ParagraphStyle(
        "ScheduleApaHeader",
        parent=styles["BodyText"],
        fontName="Helvetica-Bold",
        fontSize=6.8,
        leading=8,
        alignment=TA_LEFT,
    )
    cell_style = ParagraphStyle(
        "ScheduleApaCell",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=6.8,
        leading=8,
        alignment=TA_LEFT,
    )

    header_cells = [Paragraph(html.escape(enh.public_text(value)), header_style) for value in headers]
    body_rows = [
        [Paragraph(html.escape(enh.public_text(value)), cell_style) for value in row]
        for row in rows
    ]
    table = Table([header_cells] + body_rows, repeatRows=1, colWidths=widths)
    table.setStyle(
        TableStyle(
            [
                ("LINEABOVE", (0, 0), (-1, 0), 0.8, colors.black),
                ("LINEBELOW", (0, 0), (-1, 0), 0.6, colors.black),
                ("LINEBELOW", (0, -1), (-1, -1), 0.8, colors.black),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (2, 1), (5, -1), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 2.5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 2.5),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    return table


def _pdf_apa_note(story: list[Any], styles: Any) -> None:
    note_style = ParagraphStyle(
        "ScheduleApaNote",
        parent=styles["BodyText"],
        fontName="Helvetica-Oblique",
        fontSize=7.5,
        leading=9,
        alignment=TA_LEFT,
        spaceBefore=3,
        spaceAfter=5,
    )
    story.append(
        Paragraph(
            "Nota. Elaboración propia con base en el cronograma institucional del proceso de titulación.",
            note_style,
        )
    )


def _docx_schedules(document: Any, context: Any, report_id: int) -> None:
    data = _schedule_data(report_id)
    available = [
        ("Cronograma de Núcleos y Examen Complexivo", data["schedules"]["complexive"], False),
        ("Cronograma del Trabajo de Titulación", data["schedules"]["thesis"], True),
    ]
    available = [item for item in available if item[1]]
    if not available:
        return

    report_quality._docx_heading(document, context, 1, "Evaluación del cumplimiento de los cronogramas")
    report_quality._docx_body(
        document,
        "La evaluación contrasta la planificación con la ejecución consolidada del período y presenta el avance de cada actividad del cronograma.",
    )

    for title, rows, show_phase in available:
        report_quality._docx_heading(document, context, 2, title)
        report_quality._docx_body(
            document,
            f"La tabla presenta {len(rows)} actividades con su fecha planificada, fecha de ejecución, estado, porcentaje de ejecución, evidencia y observación.",
        )
        headers, values = _rows(rows, show_phase)
        _docx_apa_caption(document, context, f"Planificación y ejecución: {title}")
        widths = [0.75, 1.10, 0.80, 0.80, 0.70, 0.65, 0.85, 0.85] if show_phase else [1.30, 0.90, 0.90, 0.75, 0.70, 0.95, 1.00]
        _docx_apa_table(document, headers, values, widths)

    report_quality._docx_body(document, _analysis_text(data))


def _pdf_schedules(story: list[Any], context: Any, styles: Any, report_id: int) -> None:
    data = _schedule_data(report_id)
    available = [
        ("Cronograma de Núcleos y Examen Complexivo", data["schedules"]["complexive"], False),
        ("Cronograma del Trabajo de Titulación", data["schedules"]["thesis"], True),
    ]
    available = [item for item in available if item[1]]
    if not available:
        return

    report_quality._pdf_heading(story, context, styles, 1, "Evaluación del cumplimiento de los cronogramas")
    report_quality._pdf_body(
        story,
        styles,
        "La evaluación contrasta la planificación con la ejecución consolidada del período y presenta el avance de cada actividad del cronograma.",
    )

    for title, rows, show_phase in available:
        report_quality._pdf_heading(story, context, styles, 2, title)
        report_quality._pdf_body(
            story,
            styles,
            f"La tabla presenta {len(rows)} actividades con su fecha planificada, fecha de ejecución, estado, porcentaje de ejecución, evidencia y observación.",
        )
        headers, values = _rows(rows, show_phase)
        _pdf_apa_caption(story, context, styles, f"Planificación y ejecución: {title}")
        widths = [1.9, 2.6, 2.0, 2.0, 1.7, 1.8, 2.6, 2.6] if show_phase else [3.2, 2.3, 2.3, 1.9, 1.8, 2.8, 2.9]
        story += [_pdf_apa_table(headers, values, [width * cm for width in widths], styles)]
        _pdf_apa_note(story, styles)
        story.append(Spacer(1, 0.12 * cm))

    report_quality._pdf_body(story, styles, _analysis_text(data))


def install() -> None:
    report_quality._docx_schedules = _docx_schedules
    report_quality._pdf_schedules = _pdf_schedules
