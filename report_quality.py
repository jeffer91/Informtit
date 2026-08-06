from __future__ import annotations

import html
import re
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Image, PageBreak, Paragraph, Spacer, Table, TableStyle
from reportlab.platypus.tableofcontents import TableOfContents

import app as core
import institutional_export as base
import nuclei_export
import report_structure
from analytics import summary
from nuclei_catalog import catalogs_for_report, create_cycle_diagram
from nuclei_service import get_nuclei
from optional_content import is_present
from parser import canonical_name_key, clean_moodle_name
from process_service import get_projects, get_schedules


METHODOLOGY = {
    "Objetivo y alcance": [
        "La metodología de núcleos estructurantes tiene como objetivo proporcionar una guía clara y organizada para el desarrollo y la evaluación de competencias en los estudiantes. Este enfoque integra conocimientos teóricos, prácticos y metodológicos y los vincula con el perfil de egreso de las carreras del Instituto Tecnológico Superior Quito Metropolitano.",
        "La Unidad de Titulación y Eficiencia Terminal articula los saberes teórico-metodológicos, profesionales y de investigación necesarios para que el estudiante demuestre el dominio de los resultados de aprendizaje alcanzados durante su formación.",
    ],
    "Funciones de la unidad": [
        "Diseñar y ejecutar entornos para la aplicación de los conocimientos científicos y técnicos adquiridos por los estudiantes desde el primer hasta el último nivel de la carrera.",
        "Garantizar que los conocimientos finales adquiridos correspondan al mínimo necesario para ejercer la profesión con solvencia.",
    ],
    "Opciones de titulación": [
        "Trabajo de Titulación.",
        "Examen Complexivo.",
    ],
    "Núcleos estructurantes en el Examen Complexivo": [
        "El Examen Complexivo permite que los estudiantes demuestren su capacidad para responder preguntas y resolver problemáticas relacionadas con su profesión mediante el uso creativo, crítico y fundamentado del conocimiento adquirido durante la carrera.",
        "Para este proceso se establecen cuatro núcleos estructurantes por carrera. Cada núcleo reúne ejes transversales del currículo derivados del objeto de estudio y vincula campos del conocimiento mediante asignaturas integradoras y ejes temáticos principales.",
    ],
    "Principios metodológicos": [
        "Integración de saberes: combinación de conocimientos teóricos, prácticos y metodológicos orientados a la resolución de problemas propios de la carrera.",
        "Ejes temáticos principales: selección de áreas clave de cada carrera en correspondencia con el perfil de egreso y el campo profesional.",
        "Metodologías activas: aplicación de aprendizaje basado en problemas, estudios de caso, simulaciones y proyectos colaborativos.",
        "Evaluación integral: utilización de actividades formativas y sumativas, proyectos integradores, casos prácticos y evaluaciones estructuradas.",
        "Retroalimentación continua: comunicación de fortalezas y aspectos de mejora durante el desarrollo de cada núcleo.",
    ],
    "Implementación": [
        "Planificación: identificación de competencias clave, selección de contenidos esenciales y diseño de actividades integradoras.",
        "Desarrollo: ejecución de talleres, actividades colaborativas, casos y simulaciones orientadas a la aplicación del conocimiento.",
        "Seguimiento: monitoreo del progreso de los estudiantes y verificación del logro de los resultados esperados.",
    ],
}

LEGAL_SUBHEADINGS = {
    "Constitución de la República del Ecuador",
    "Ley Orgánica de Educación Superior",
    "Reglamento de Régimen Académico",
    "Reglamento de la Unidad de Titulación y Eficiencia Terminal",
    "Manual de Procesos de la Unidad de Titulación y Eficiencia Terminal",
    "Aplicación del marco legal y normativo",
}

POST_KEYS = ("analisis_estrategico", "conclusiones", "recomendaciones")
GRADE_FIELDS = (
    "ordinary_theory",
    "ordinary_practical",
    "supplementary_theory",
    "supplementary_practical",
    "source_total_theory",
    "source_total_practical",
    "source_total_course",
)


@dataclass
class ExportContext:
    counters: list[int]
    figure: int = 0
    table: int = 0
    major_started: bool = False

    @classmethod
    def create(cls) -> "ExportContext":
        return cls([0, 0, 0, 0])

    def heading(self, level: int, title: str) -> str:
        index = level - 1
        self.counters[index] += 1
        for lower in range(index + 1, len(self.counters)):
            self.counters[lower] = 0
        prefix = ".".join(str(value) for value in self.counters[:level])
        return f"{prefix}. {title}"

    def figure_caption(self, title: str) -> str:
        self.figure += 1
        return f"Figura {self.figure}. {title}"

    def table_caption(self, title: str) -> str:
        self.table += 1
        return f"Tabla {self.table}. {title}"


def _section(report: dict[str, Any], key: str) -> dict[str, Any] | None:
    return next((item for item in report.get("sections", []) if item.get("section_key") == key), None)


def _has_notes(career: dict[str, Any]) -> bool:
    return any(any(student.get(field) is not None for field in GRADE_FIELDS) for student in career.get("students", []))


def _grade_count(student: dict[str, Any]) -> int:
    return sum(student.get(field) is not None for field in GRADE_FIELDS)


def _merge_student_group(items: list[dict[str, Any]]) -> dict[str, Any]:
    ordered = sorted(items, key=_grade_count, reverse=True)
    merged = dict(ordered[0])
    merged["full_name"] = clean_moodle_name(str(merged.get("full_name") or ""))
    for item in ordered[1:]:
        for field in GRADE_FIELDS:
            if merged.get(field) is None and item.get(field) is not None:
                merged[field] = item[field]
        if not merged.get("email") and item.get("email"):
            merged["email"] = item["email"]
        if not merged.get("identification") and item.get("identification"):
            merged["identification"] = item["identification"]
    return merged


def _deduplicate_career(career: dict[str, Any]) -> tuple[dict[str, Any], list[str]]:
    result = dict(career)
    groups: dict[str, list[dict[str, Any]]] = {}
    unkeyed: list[dict[str, Any]] = []
    for student in career.get("students", []):
        item = dict(student)
        item["full_name"] = clean_moodle_name(str(item.get("full_name") or ""))
        key = canonical_name_key(item["full_name"])
        if key:
            groups.setdefault(key, []).append(item)
        else:
            unkeyed.append(item)

    warnings: list[str] = []
    merged_students: list[dict[str, Any]] = []
    for items in groups.values():
        if len(items) > 1:
            warnings.append(" / ".join(item["full_name"] for item in items))
        merged_students.append(_merge_student_group(items))
    merged_students.extend(unkeyed)
    result["students"] = sorted(merged_students, key=lambda item: str(item.get("full_name") or ""))
    return result, warnings


def _report_data(report_id: int) -> dict[str, Any]:
    report = base.legacy.load_report_data(report_id)
    cleaned: list[dict[str, Any]] = []
    duplicate_warnings: dict[int, list[str]] = {}
    for career in report.get("careers", []):
        item, warnings = _deduplicate_career(career)
        cleaned.append(item)
        if warnings:
            duplicate_warnings[int(career["id"])] = warnings
    report["careers"] = cleaned
    report["duplicate_warnings"] = duplicate_warnings
    return report


def _fmt(value: Any) -> str:
    if value is None or value == "":
        return "—"
    if isinstance(value, (int, float)):
        return f"{float(value):.2f}".replace(".", ",")
    return str(value)


def _pct(value: Any) -> str:
    return f"{_fmt(value)} %"


def _plural(total: int, singular: str, plural: str | None = None) -> str:
    return singular if total == 1 else (plural or singular + "s")


def _phase_before(career_name: str, phase: str, data: dict[str, Any]) -> str:
    total = int(data["total"])
    noun = _plural(total, "estudiante")
    if phase == "consolidado":
        return (
            f"A continuación, se presentan los resultados consolidados de la carrera de {career_name}. "
            f"La tabla contiene información de {total} {noun} y resume el estado final del proceso."
        )
    label = "evaluación ordinaria" if phase == "ordinario" else "evaluación supletoria"
    return (
        f"A continuación, se presentan los resultados de la {label} de la carrera de {career_name}. "
        f"La tabla contiene información de {total} {noun} y resume el desempeño alcanzado."
    )


def _phase_after(data: dict[str, Any]) -> str:
    total = int(data["total"])
    approved = int(data["approved"])
    failed = int(data["failed"])
    not_evaluated = int(data["not_evaluated"])
    record = _plural(total, "registro")
    approved_verb = "alcanzó" if approved == 1 else "alcanzaron"
    failed_verb = "no alcanzó" if failed == 1 else "no alcanzaron"
    sentence = (
        f"De {total} {record} analizados, {approved} {approved_verb} la aprobación "
        f"({_pct(data['approved_pct'])}) y {failed} {failed_verb} la calificación mínima."
    )
    if not_evaluated:
        sentence += f" Además, {not_evaluated} {_plural(not_evaluated, 'estudiante')} no registraron una evaluación completa."
    sentence += f" El promedio registrado fue {_fmt(data['average_final'])}."
    return sentence


def _sanitize_analysis(value: str) -> str:
    text = " ".join(str(value or "").split())
    replacements = {
        "resultados de la resultado consolidado": "resultados consolidados",
        "información de 1 estudiantes": "información de 1 estudiante",
        "De los 1 registros": "De 1 registro",
        "1 alcanzaron": "1 alcanzó",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"(\d+)\.(\d{2})\s*%", lambda match: f"{match.group(1)},{match.group(2)} %", text)
    return text


# ---------------------------------------------------------------------------
# Word


def _configure_docx(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for level in range(1, 5):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style.font.bold = True
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10 if level == 1 else 7)
        style.paragraph_format.space_after = Pt(5)

    bullet = document.styles["List Bullet"]
    bullet.font.name = "Arial"
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Cm(1.25)
    bullet.paragraph_format.first_line_indent = Cm(-0.63)
    bullet.paragraph_format.space_after = Pt(3)

    if "Figure Caption" not in [style.name for style in document.styles]:
        document.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption = document.styles["Figure Caption"]
    caption.font.name = "Arial"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.space_after = Pt(4)

    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def _docx_heading(document: Document, context: ExportContext, level: int, title: str, page_break: bool = False) -> Any:
    if level == 1:
        if context.major_started:
            document.add_page_break()
        context.major_started = True
    elif page_break:
        document.add_page_break()
    paragraph = document.add_heading(context.heading(level, title), level=level)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.first_line_indent = Cm(0)
    return paragraph


def _docx_body(document: Document, text: str) -> Any:
    paragraph = document.add_paragraph(_sanitize_analysis(text))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def _docx_bullet(document: Document, text: str) -> Any:
    paragraph = document.add_paragraph(str(text).strip(), style="List Bullet")
    paragraph.paragraph_format.left_indent = Cm(1.25)
    paragraph.paragraph_format.first_line_indent = Cm(-0.63)
    return paragraph


def _docx_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text, style="Figure Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _docx_toc(document: Document) -> None:
    title = document.add_paragraph("ÍNDICE", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Cm(0)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Actualice la tabla de contenido al abrir el documento."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, placeholder, end])
    document.add_page_break()


def _docx_table(document: Document, headers: list[str], rows: list[list[Any]], widths: list[float]) -> None:
    report_structure._docx_table(document, headers, rows, widths)


def _docx_legal(document: Document, context: ExportContext, report: dict[str, Any]) -> None:
    section = _section(report, "marco_legal")
    if not section or not str(section.get("content") or "").strip():
        return
    _docx_heading(document, context, 1, "Marco legal y normativo")
    blocks = [block.strip() for block in re.split(r"\n\s*\n", str(section["content"])) if block.strip()]
    for block in blocks:
        if block in LEGAL_SUBHEADINGS:
            _docx_heading(document, context, 2, block)
        elif block.startswith("UTET-PRO-"):
            for line in block.splitlines():
                if line.strip():
                    _docx_bullet(document, line.strip())
        else:
            _docx_body(document, block)


def _docx_regulation(document: Document, context: ExportContext, report: dict[str, Any]) -> None:
    section = _section(report, "reglamento")
    if not section or not str(section.get("content") or "").strip():
        return
    _docx_heading(document, context, 1, "Reglamento del Examen Complexivo")
    for block in re.split(r"\n\s*\n", str(section["content"])):
        if block.strip():
            _docx_body(document, block.strip())


def _diagram_path(report_id: int, career_name: str) -> Path:
    base.EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    safe = re.sub(r"[^a-z0-9]+", "_", canonical_name_key(career_name)).strip("_")
    return base.EXPORT_DIR / f"nuclei_catalog_{report_id}_{safe}.png"


def _docx_methodology(document: Document, context: ExportContext, report: dict[str, Any]) -> None:
    _docx_heading(document, context, 1, "Metodología de núcleos estructurantes")
    for title, items in METHODOLOGY.items():
        _docx_heading(document, context, 2, title)
        if title in {"Funciones de la unidad", "Opciones de titulación", "Principios metodológicos", "Implementación"}:
            for item in items:
                _docx_bullet(document, item)
        else:
            for item in items:
                _docx_body(document, item)

    infographic = base.image_for(report, base.INFOGRAPHIC)
    if base.image_path(infographic):
        _docx_heading(document, context, 2, "Infografía del proceso de Examen Complexivo")
        _docx_caption(document, context.figure_caption("Proceso general de núcleos estructurantes y Examen Complexivo"))
        base.add_docx_image(document, infographic, 6.3, False)
        _docx_caption(document, "Nota. Fuente institucional cargada en Informtit.")

    catalogs = catalogs_for_report(report)
    if not catalogs:
        return
    _docx_heading(document, context, 2, "Contenido académico de los núcleos")
    for catalog in catalogs:
        _docx_heading(document, context, 3, catalog["career"], page_break=True)
        _docx_body(
            document,
            "La carrera organiza su preparación académica en cuatro núcleos estructurantes. Cada núcleo corresponde a una guía de integración curricular vinculada con el perfil de egreso.",
        )
        diagram = _diagram_path(int(report["id"]), catalog["career"])
        create_cycle_diagram(catalog, diagram)
        _docx_caption(document, context.figure_caption(f"Estructura de los núcleos de la carrera de {catalog['career']}"))
        image_paragraph = document.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.paragraph_format.first_line_indent = Cm(0)
        image_paragraph.add_run().add_picture(str(diagram), width=Inches(6.25))
        _docx_caption(document, "Nota. Elaboración propia con base en las guías de integración curricular de la carrera.")
        for nucleus in catalog.get("nuclei", []):
            _docx_heading(document, context, 4, f"Núcleo {nucleus['number']}: {nucleus['guide']}")
            subjects = nucleus.get("subjects", [])
            if subjects:
                _docx_body(document, "Esta guía articula las siguientes asignaturas:")
                for subject in subjects:
                    _docx_bullet(document, subject)
            else:
                _docx_body(
                    document,
                    "Este núcleo corresponde a la guía institucional de integración curricular registrada para la carrera. Sus contenidos específicos se mantienen en el documento académico de origen.",
                )


def _docx_requirements(document: Document, context: ExportContext, report_id: int) -> None:
    data = report_structure.requirement_analysis(report_id)
    if not data:
        return
    _docx_heading(document, context, 1, "Análisis del cumplimiento de requisitos de titulación")
    _docx_heading(document, context, 2, "Resumen general")
    _docx_body(
        document,
        f"Se analizaron {data['total']} estudiantes. {data['complete']} presentaron cumplimiento integral, "
        f"{data['pending']} registraron al menos un requisito pendiente y {data['incomplete']} presentaron "
        f"información incompleta. El cumplimiento integral fue del {_pct(data['percentage'])}.",
    )
    _docx_heading(document, context, 2, "Cumplimiento por requisito")
    _docx_caption(document, context.table_caption("Cumplimiento de los requisitos de titulación"))
    _docx_table(
        document,
        ["Requisito", "Cumple", "No cumple", "Sin información", "Cumplimiento"],
        [[row["label"], row["complies"], row["does_not_comply"], row["blank"], _pct(row["percentage"])] for row in data["requirements"]],
        [2.55, 0.8, 0.85, 1.05, 1.05],
    )
    _docx_heading(document, context, 2, "Cumplimiento por carrera")
    _docx_caption(document, context.table_caption("Cumplimiento integral por carrera"))
    _docx_table(
        document,
        ["Carrera", "Registrados", "Completos", "Pendientes", "Sin información", "Cumplimiento"],
        [[row["career"], row["registered"], row["complete"], row["pending"], row["incomplete"], _pct(row["percentage"])] for row in data["careers"]],
        [2.45, 0.75, 0.75, 0.75, 0.9, 0.95],
    )
    _docx_heading(document, context, 2, "Análisis de resultados")
    _docx_body(document, data["narrative"])


def _docx_schedule_table(document: Document, rows: list[dict[str, Any]], show_phase: bool) -> None:
    headers = ["Fase", "Actividad", "Fecha de inicio", "Fecha de fin"] if show_phase else ["Actividad", "Fecha de inicio", "Fecha de fin"]
    values: list[list[Any]] = []
    for row in rows:
        value = [row.get("activity") or "—", row.get("start_date") or "—", row.get("end_date") or "—"]
        if show_phase:
            value.insert(0, row.get("phase") or "—")
        values.append(value)
    widths = [1.45, 2.85, 1.1, 1.1] if show_phase else [3.9, 1.2, 1.2]
    _docx_table(document, headers, values, widths)


def _docx_schedules(document: Document, context: ExportContext, report_id: int) -> None:
    schedules = get_schedules(report_id)
    has_complexive = is_present(report_id, "schedule_complexive") and bool(schedules.get("complexive"))
    has_thesis = is_present(report_id, "schedule_thesis") and bool(schedules.get("thesis"))
    if not has_complexive and not has_thesis:
        return
    _docx_heading(document, context, 1, "Cronogramas")
    if has_complexive:
        _docx_heading(document, context, 2, "Cronograma de Núcleos y Examen Complexivo")
        _docx_caption(document, context.table_caption("Actividades del cronograma de Núcleos y Examen Complexivo"))
        _docx_schedule_table(document, schedules["complexive"], False)
    if has_thesis:
        _docx_heading(document, context, 2, "Cronograma del Trabajo de Titulación")
        _docx_caption(document, context.table_caption("Actividades del cronograma del Trabajo de Titulación"))
        _docx_schedule_table(document, schedules["thesis"], True)


def _docx_nucleus_results(document: Document, context: ExportContext, report_id: int) -> None:
    courses = get_nuclei(report_id).get("courses", [])
    if not courses:
        return
    _docx_heading(document, context, 1, "Resultados de los núcleos estructurantes")
    for index, course in enumerate(courses):
        _docx_heading(document, context, 2, f"{course['career_name']} – Núcleo {course['nucleus_number']}", page_break=index > 0)
        _docx_body(
            document,
            f"El curso fue impartido por {course.get('teacher_name') or 'docente pendiente de confirmar'} y contó con el seguimiento de {course.get('coordinator_name') or 'la coordinación de carrera'}. Se registraron {course.get('graded_students', 0)} estudiantes con calificaciones; {course.get('approved_count', 0)} aprobaron y {course.get('failed_count', 0)} reprobaron. El promedio general fue {_fmt(course.get('course_average'))}.",
        )
        _docx_caption(document, context.table_caption(f"Calificaciones del Núcleo {course['nucleus_number']} de {course['career_name']}"))
        nuclei_export._docx_score_table(document, course)
        _docx_heading(document, context, 3, "Promedios por actividad")
        _docx_caption(document, context.table_caption(f"Promedios de las actividades del Núcleo {course['nucleus_number']}"))
        nuclei_export._docx_averages(document, course)


def _docx_phase_table(document: Document, career: dict[str, Any], phase: str, data: dict[str, Any]) -> None:
    rows = data["rows"]
    if phase == "ordinario":
        headers = ["Estudiante", "Teórico", "Práctico", "Final", "Estado"]
        values = [[row["full_name"], _fmt(row["ordinary_theory"]), _fmt(row["ordinary_practical"]), _fmt(row["ordinary_final"]), row["ordinary_status"]] for row in rows]
        widths = [3.1, 0.8, 0.8, 0.8, 1.0]
    elif phase == "supletorio":
        headers = ["Estudiante", "Componente", "Teórico sup.", "Práctico sup.", "Final", "Estado"]
        values = [[row["full_name"], row["supplementary_component"], _fmt(row["supplementary_theory"]), _fmt(row["supplementary_practical"]), _fmt(row["supplementary_final"]), row["final_status"]] for row in rows]
        widths = [2.5, 1.15, 0.85, 0.85, 0.75, 0.9]
    else:
        headers = ["Estudiante", "Final", "Estado", "Supletorio"]
        values = [[row["full_name"], _fmt(row["final_grade"]), row["final_status"], "Sí" if row["supplementary_participant"] else "No"] for row in rows]
        widths = [3.8, 0.9, 1.15, 0.85]
    _docx_table(document, headers, values, widths)


def _docx_complexive(document: Document, context: ExportContext, report: dict[str, Any]) -> None:
    careers = [career for career in report.get("careers", []) if _has_notes(career)]
    if not careers:
        return
    _docx_heading(document, context, 1, "Resultados del Examen Complexivo")
    for index, career in enumerate(careers):
        _docx_heading(document, context, 2, str(career["name"]), page_break=index > 0)
        warnings = report.get("duplicate_warnings", {}).get(int(career["id"]), [])
        if warnings:
            _docx_body(
                document,
                "Control de calidad: se consolidaron posibles registros duplicados identificados por variaciones en el orden del nombre. Revise los siguientes casos: " + "; ".join(warnings) + ".",
            )
        for phase, heading in (("ordinario", "Resultados de la evaluación ordinaria"), ("supletorio", "Resultados de la evaluación supletoria"), ("consolidado", "Resultado consolidado")):
            data = summary(career["students"], phase)
            if phase == "supletorio" and not data["rows"]:
                continue
            _docx_heading(document, context, 3, heading)
            analysis = career.get("analyses", {}).get(phase, {})
            before = _sanitize_analysis(analysis.get("text_before") or "") or _phase_before(career["name"], phase, data)
            after = _sanitize_analysis(analysis.get("text_after") or "") or _phase_after(data)
            _docx_body(document, before)
            _docx_caption(document, context.table_caption(f"{heading} de {career['name']}"))
            _docx_phase_table(document, career, phase, data)
            _docx_body(document, after)
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temporary:
            chart = Path(temporary.name)
        try:
            base.legacy.create_chart(career, chart)
            _docx_caption(document, context.figure_caption(f"Resultados consolidados de {career['name']}"))
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.add_run().add_picture(str(chart), width=Inches(6.35))
            _docx_caption(document, "Nota. Elaboración propia con base en las calificaciones procesadas por Informtit.")
        finally:
            chart.unlink(missing_ok=True)


def _docx_projects(document: Document, context: ExportContext, report_id: int) -> None:
    data = get_projects(report_id)
    projects = data.get("projects", [])
    if not projects:
        return
    _docx_heading(document, context, 1, "Resultados del Trabajo de Titulación")
    summary_data = data["summary"]
    _docx_body(
        document,
        f"Se registraron {summary_data['total']} {_plural(int(summary_data['total']), 'estudiante')}; {summary_data['approved']} aprobaron y {summary_data['failed']} reprobaron. El promedio final fue {_fmt(summary_data['average_final'])}.",
    )
    for index, project in enumerate(projects):
        _docx_heading(document, context, 2, project["full_name"], page_break=index > 0)
        _docx_caption(document, context.table_caption(f"Información del Trabajo de Titulación de {project['full_name']}"))
        rows = [
            ["Cédula", project.get("identification") or "—", "Carrera", project.get("career_name") or "—"],
            ["Código", project.get("career_code") or "—", "Acta", project.get("act_number") or "—"],
            ["Fecha", project.get("act_date") or "—", "Trabajo escrito", _fmt(project.get("written_average"))],
            ["Defensa oral", _fmt(project.get("oral_average")), "Calificación final", _fmt(project.get("final_grade"))],
        ]
        table = document.add_table(rows=0, cols=4)
        table.style = "Table Grid"
        for row in rows:
            cells = table.add_row().cells
            for column, value in enumerate(row):
                cells[column].text = str(value)
                for run in cells[column].paragraphs[0].runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8)
                    run.bold = column % 2 == 0
        _docx_body(document, f"Vocales: {project.get('vocal_1') or '—'}; {project.get('vocal_2') or '—'}; {project.get('vocal_3') or '—'}.")
        for evaluation_type, title in (("practical", "Evaluación práctica"), ("defense", "Evaluación de la defensa")):
            scores = [row for row in project.get("scores", []) if row["evaluation_type"] == evaluation_type]
            if not scores:
                continue
            _docx_heading(document, context, 3, title)
            _docx_caption(document, context.table_caption(f"{title} de {project['full_name']}"))
            _docx_table(
                document,
                ["Criterio", "Máximo", "Primer vocal", "Segundo vocal", "Tercer vocal"],
                [[row["criterion"], _fmt(row["max_score"]), _fmt(row["vocal_1"]), _fmt(row["vocal_2"]), _fmt(row["vocal_3"])] for row in scores],
                [2.9, 0.7, 0.85, 0.85, 0.85],
            )


def _docx_post_sections(document: Document, context: ExportContext, report: dict[str, Any]) -> None:
    titles = {
        "analisis_estrategico": "Análisis estratégico",
        "conclusiones": "Conclusiones",
        "recomendaciones": "Recomendaciones",
    }
    for key in POST_KEYS:
        section = _section(report, key)
        if not section or not str(section.get("content") or "").strip():
            continue
        _docx_heading(document, context, 1, titles[key])
        for block in re.split(r"\n\s*\n", str(section["content"])):
            if block.strip():
                _docx_body(document, block.strip())


def _additional_images(report: dict[str, Any]) -> list[dict[str, Any]]:
    return report_structure._additional_images(report)


def build_docx(report_id: int) -> Path:
    report_structure.ensure_structure_schema()
    base.EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    report = _report_data(report_id)
    output = base.EXPORT_DIR / f"informtit_{report_id}.docx"
    document = Document()
    _configure_docx(document)
    base.setup_header(document, report)
    base.cover_docx(document, report)
    _docx_toc(document)
    context = ExportContext.create()

    _docx_heading(document, context, 1, "Introducción")
    for paragraph in report_structure.introduction(report, report_id):
        _docx_body(document, paragraph)
    _docx_legal(document, context, report)
    _docx_regulation(document, context, report)
    _docx_methodology(document, context, report)
    _docx_requirements(document, context, report_id)
    _docx_schedules(document, context, report_id)
    _docx_nucleus_results(document, context, report_id)
    _docx_complexive(document, context, report)
    _docx_projects(document, context, report_id)
    _docx_post_sections(document, context, report)

    images = _additional_images(report)
    if images:
        _docx_heading(document, context, 1, "Anexos")
        for image in images:
            title = str(image.get("title") or image.get("original_name") or "Evidencia")
            _docx_caption(document, context.figure_caption(title))
            base.add_docx_image(document, image, 6.2, False)
            source = str(image.get("source") or "Fuente institucional cargada en Informtit")
            _docx_caption(document, f"Nota. {source}.")

    document.save(output)
    return output


# ---------------------------------------------------------------------------
# PDF


def _pdf_styles() -> Any:
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            "BodyJustified",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            leading=13.5,
            alignment=TA_JUSTIFY,
            firstLineIndent=1.25 * cm,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            "BulletIndented",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            leading=13,
            leftIndent=1.25 * cm,
            firstLineIndent=-0.63 * cm,
            bulletIndent=0.62 * cm,
            spaceAfter=3,
        )
    )
    styles.add(
        ParagraphStyle(
            "Heading4",
            parent=styles["Heading3"],
            fontName="Helvetica-Bold",
            fontSize=10.5,
            leading=13,
            spaceBefore=7,
            spaceAfter=4,
            keepWithNext=True,
        )
    )
    styles.add(
        ParagraphStyle(
            "FigureCaption",
            parent=styles["BodyText"],
            fontName="Helvetica-Oblique",
            fontSize=8.5,
            leading=10.5,
            alignment=TA_CENTER,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            "TableCell",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=6.8,
            leading=8,
            alignment=TA_LEFT,
        )
    )
    for level in (1, 2, 3):
        style = styles[f"Heading{level}"]
        style.fontName = "Helvetica-Bold"
        style.keepWithNext = True
        style.spaceBefore = 9 if level == 1 else 6
        style.spaceAfter = 5
    return styles


def _pdf_heading(story: list[Any], context: ExportContext, styles: Any, level: int, title: str, page_break: bool = False) -> None:
    if level == 1:
        if context.major_started:
            story.append(PageBreak())
        context.major_started = True
    elif page_break:
        story.append(PageBreak())
    story.append(Paragraph(html.escape(context.heading(level, title)), styles[f"Heading{level}"]))


def _pdf_body(story: list[Any], styles: Any, text: str) -> None:
    story.append(Paragraph(html.escape(_sanitize_analysis(text)), styles["BodyJustified"]))


def _pdf_bullet(story: list[Any], styles: Any, text: str) -> None:
    story.append(Paragraph("• " + html.escape(str(text).strip()), styles["BulletIndented"]))


def _pdf_caption(story: list[Any], styles: Any, text: str) -> None:
    story.append(Paragraph(html.escape(text), styles["FigureCaption"]))


def _pdf_table(headers: list[Any], rows: list[list[Any]], widths: list[float]) -> Table:
    table = Table([headers] + rows, repeatRows=1, colWidths=widths)
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#244a73")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (1, 1), (-1, -1), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    return table


def _pdf_legal(story: list[Any], context: ExportContext, styles: Any, report: dict[str, Any]) -> None:
    section = _section(report, "marco_legal")
    if not section or not str(section.get("content") or "").strip():
        return
    _pdf_heading(story, context, styles, 1, "Marco legal y normativo")
    blocks = [block.strip() for block in re.split(r"\n\s*\n", str(section["content"])) if block.strip()]
    for block in blocks:
        if block in LEGAL_SUBHEADINGS:
            _pdf_heading(story, context, styles, 2, block)
        elif block.startswith("UTET-PRO-"):
            for line in block.splitlines():
                if line.strip():
                    _pdf_bullet(story, styles, line.strip())
        else:
            _pdf_body(story, styles, block)


def _pdf_regulation(story: list[Any], context: ExportContext, styles: Any, report: dict[str, Any]) -> None:
    section = _section(report, "reglamento")
    if not section or not str(section.get("content") or "").strip():
        return
    _pdf_heading(story, context, styles, 1, "Reglamento del Examen Complexivo")
    for block in re.split(r"\n\s*\n", str(section["content"])):
        if block.strip():
            _pdf_body(story, styles, block.strip())


def _pdf_methodology(story: list[Any], context: ExportContext, styles: Any, report: dict[str, Any], temp_paths: list[Path]) -> None:
    _pdf_heading(story, context, styles, 1, "Metodología de núcleos estructurantes")
    for title, items in METHODOLOGY.items():
        _pdf_heading(story, context, styles, 2, title)
        if title in {"Funciones de la unidad", "Opciones de titulación", "Principios metodológicos", "Implementación"}:
            for item in items:
                _pdf_bullet(story, styles, item)
        else:
            for item in items:
                _pdf_body(story, styles, item)

    infographic = base.image_path(base.image_for(report, base.INFOGRAPHIC))
    if infographic:
        _pdf_heading(story, context, styles, 2, "Infografía del proceso de Examen Complexivo")
        _pdf_caption(story, styles, context.figure_caption("Proceso general de núcleos estructurantes y Examen Complexivo"))
        story += [base.fit_image(infographic, 16.5 * cm, 19 * cm), Spacer(1, 0.12 * cm)]
        _pdf_caption(story, styles, "Nota. Fuente institucional cargada en Informtit.")

    catalogs = catalogs_for_report(report)
    if not catalogs:
        return
    _pdf_heading(story, context, styles, 2, "Contenido académico de los núcleos")
    for catalog in catalogs:
        _pdf_heading(story, context, styles, 3, catalog["career"], page_break=True)
        _pdf_body(story, styles, "La carrera organiza su preparación académica en cuatro núcleos estructurantes. Cada núcleo corresponde a una guía de integración curricular vinculada con el perfil de egreso.")
        diagram = _diagram_path(int(report["id"]), catalog["career"])
        create_cycle_diagram(catalog, diagram)
        temp_paths.append(diagram)
        _pdf_caption(story, styles, context.figure_caption(f"Estructura de los núcleos de la carrera de {catalog['career']}"))
        story += [base.fit_image(diagram, 16.2 * cm, 11.2 * cm), Spacer(1, 0.1 * cm)]
        _pdf_caption(story, styles, "Nota. Elaboración propia con base en las guías de integración curricular de la carrera.")
        for nucleus in catalog.get("nuclei", []):
            story.append(Paragraph(html.escape(context.heading(4, f"Núcleo {nucleus['number']}: {nucleus['guide']}")), styles["Heading4"]))
            subjects = nucleus.get("subjects", [])
            if subjects:
                _pdf_body(story, styles, "Esta guía articula las siguientes asignaturas:")
                for subject in subjects:
                    _pdf_bullet(story, styles, subject)
            else:
                _pdf_body(story, styles, "Este núcleo corresponde a la guía institucional de integración curricular registrada para la carrera. Sus contenidos específicos se mantienen en el documento académico de origen.")


def _pdf_requirements(story: list[Any], context: ExportContext, styles: Any, report_id: int) -> None:
    data = report_structure.requirement_analysis(report_id)
    if not data:
        return
    _pdf_heading(story, context, styles, 1, "Análisis del cumplimiento de requisitos de titulación")
    _pdf_heading(story, context, styles, 2, "Resumen general")
    _pdf_body(story, styles, f"Se analizaron {data['total']} estudiantes. {data['complete']} presentaron cumplimiento integral, {data['pending']} registraron al menos un requisito pendiente y {data['incomplete']} presentaron información incompleta. El cumplimiento integral fue del {_pct(data['percentage'])}.")
    _pdf_heading(story, context, styles, 2, "Cumplimiento por requisito")
    _pdf_caption(story, styles, context.table_caption("Cumplimiento de los requisitos de titulación"))
    story += [
        _pdf_table(
            ["Requisito", "Cumple", "No cumple", "Sin información", "Cumplimiento"],
            [[Paragraph(html.escape(row["label"]), styles["TableCell"]), row["complies"], row["does_not_comply"], row["blank"], _pct(row["percentage"])] for row in data["requirements"]],
            [6.4 * cm, 2.1 * cm, 2.2 * cm, 2.4 * cm, 2.5 * cm],
        ),
        Spacer(1, 0.2 * cm),
    ]
    _pdf_heading(story, context, styles, 2, "Cumplimiento por carrera")
    _pdf_caption(story, styles, context.table_caption("Cumplimiento integral por carrera"))
    story += [
        _pdf_table(
            ["Carrera", "Registrados", "Completos", "Pendientes", "Sin información", "Cumplimiento"],
            [[Paragraph(html.escape(row["career"]), styles["TableCell"]), row["registered"], row["complete"], row["pending"], row["incomplete"], _pct(row["percentage"])] for row in data["careers"]],
            [5.7 * cm, 2 * cm, 2 * cm, 2 * cm, 2.3 * cm, 2.5 * cm],
        ),
        Spacer(1, 0.2 * cm),
    ]
    _pdf_heading(story, context, styles, 2, "Análisis de resultados")
    _pdf_body(story, styles, data["narrative"])


def _pdf_schedule_table(rows: list[dict[str, Any]], show_phase: bool, styles: Any) -> Table:
    headers = ["Fase", "Actividad", "Fecha de inicio", "Fecha de fin"] if show_phase else ["Actividad", "Fecha de inicio", "Fecha de fin"]
    values: list[list[Any]] = []
    for row in rows:
        current = [row.get("activity") or "—", row.get("start_date") or "—", row.get("end_date") or "—"]
        if show_phase:
            current.insert(0, row.get("phase") or "—")
        values.append([Paragraph(html.escape(str(value)), styles["TableCell"]) for value in current])
    widths = [4.1 * cm, 7 * cm, 2.7 * cm, 2.7 * cm] if show_phase else [10.5 * cm, 3 * cm, 3 * cm]
    return _pdf_table(headers, values, widths)


def _pdf_schedules(story: list[Any], context: ExportContext, styles: Any, report_id: int) -> None:
    schedules = get_schedules(report_id)
    has_complexive = is_present(report_id, "schedule_complexive") and bool(schedules.get("complexive"))
    has_thesis = is_present(report_id, "schedule_thesis") and bool(schedules.get("thesis"))
    if not has_complexive and not has_thesis:
        return
    _pdf_heading(story, context, styles, 1, "Cronogramas")
    if has_complexive:
        _pdf_heading(story, context, styles, 2, "Cronograma de Núcleos y Examen Complexivo")
        _pdf_caption(story, styles, context.table_caption("Actividades del cronograma de Núcleos y Examen Complexivo"))
        story += [_pdf_schedule_table(schedules["complexive"], False, styles), Spacer(1, 0.25 * cm)]
    if has_thesis:
        _pdf_heading(story, context, styles, 2, "Cronograma del Trabajo de Titulación")
        _pdf_caption(story, styles, context.table_caption("Actividades del cronograma del Trabajo de Titulación"))
        story += [_pdf_schedule_table(schedules["thesis"], True, styles), Spacer(1, 0.25 * cm)]


def _pdf_nucleus_results(story: list[Any], context: ExportContext, styles: Any, report_id: int) -> None:
    courses = get_nuclei(report_id).get("courses", [])
    if not courses:
        return
    _pdf_heading(story, context, styles, 1, "Resultados de los núcleos estructurantes")
    for index, course in enumerate(courses):
        _pdf_heading(story, context, styles, 2, f"{course['career_name']} – Núcleo {course['nucleus_number']}", page_break=index > 0)
        _pdf_body(story, styles, f"El curso fue impartido por {course.get('teacher_name') or 'docente pendiente de confirmar'} y contó con el seguimiento de {course.get('coordinator_name') or 'la coordinación de carrera'}. Se registraron {course.get('graded_students', 0)} estudiantes con calificaciones; {course.get('approved_count', 0)} aprobaron y {course.get('failed_count', 0)} reprobaron. El promedio general fue {_fmt(course.get('course_average'))}.")
        _pdf_caption(story, styles, context.table_caption(f"Calificaciones del Núcleo {course['nucleus_number']} de {course['career_name']}"))
        story += [nuclei_export._pdf_score_table(course, styles), Spacer(1, 0.2 * cm)]
        _pdf_heading(story, context, styles, 3, "Promedios por actividad")
        _pdf_caption(story, styles, context.table_caption(f"Promedios de las actividades del Núcleo {course['nucleus_number']}"))
        story += [nuclei_export._pdf_averages(course), Spacer(1, 0.25 * cm)]


def _pdf_phase_table(data: dict[str, Any], phase: str, styles: Any) -> Table:
    rows = data["rows"]
    if phase == "ordinario":
        headers = ["Estudiante", "Teórico", "Práctico", "Final", "Estado"]
        values = [[Paragraph(html.escape(row["full_name"]), styles["TableCell"]), _fmt(row["ordinary_theory"]), _fmt(row["ordinary_practical"]), _fmt(row["ordinary_final"]), row["ordinary_status"]] for row in rows]
        widths = [7 * cm, 2 * cm, 2 * cm, 2 * cm, 3.2 * cm]
    elif phase == "supletorio":
        headers = ["Estudiante", "Componente", "Teórico sup.", "Práctico sup.", "Final", "Estado"]
        values = [[Paragraph(html.escape(row["full_name"]), styles["TableCell"]), row["supplementary_component"], _fmt(row["supplementary_theory"]), _fmt(row["supplementary_practical"]), _fmt(row["supplementary_final"]), row["final_status"]] for row in rows]
        widths = [5.6 * cm, 2.8 * cm, 2.1 * cm, 2.1 * cm, 1.8 * cm, 2.2 * cm]
    else:
        headers = ["Estudiante", "Final", "Estado", "Supletorio"]
        values = [[Paragraph(html.escape(row["full_name"]), styles["TableCell"]), _fmt(row["final_grade"]), row["final_status"], "Sí" if row["supplementary_participant"] else "No"] for row in rows]
        widths = [9.4 * cm, 2.3 * cm, 3 * cm, 2.3 * cm]
    return _pdf_table(headers, values, widths)


def _pdf_complexive(story: list[Any], context: ExportContext, styles: Any, report: dict[str, Any], temp_paths: list[Path]) -> None:
    careers = [career for career in report.get("careers", []) if _has_notes(career)]
    if not careers:
        return
    _pdf_heading(story, context, styles, 1, "Resultados del Examen Complexivo")
    for index, career in enumerate(careers):
        _pdf_heading(story, context, styles, 2, str(career["name"]), page_break=index > 0)
        warnings = report.get("duplicate_warnings", {}).get(int(career["id"]), [])
        if warnings:
            _pdf_body(story, styles, "Control de calidad: se consolidaron posibles registros duplicados identificados por variaciones en el orden del nombre. Revise los siguientes casos: " + "; ".join(warnings) + ".")
        for phase, heading in (("ordinario", "Resultados de la evaluación ordinaria"), ("supletorio", "Resultados de la evaluación supletoria"), ("consolidado", "Resultado consolidado")):
            data = summary(career["students"], phase)
            if phase == "supletorio" and not data["rows"]:
                continue
            _pdf_heading(story, context, styles, 3, heading)
            analysis = career.get("analyses", {}).get(phase, {})
            before = _sanitize_analysis(analysis.get("text_before") or "") or _phase_before(career["name"], phase, data)
            after = _sanitize_analysis(analysis.get("text_after") or "") or _phase_after(data)
            _pdf_body(story, styles, before)
            _pdf_caption(story, styles, context.table_caption(f"{heading} de {career['name']}"))
            story += [_pdf_phase_table(data, phase, styles), Spacer(1, 0.15 * cm)]
            _pdf_body(story, styles, after)
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temporary:
            chart = Path(temporary.name)
        base.legacy.create_chart(career, chart)
        temp_paths.append(chart)
        _pdf_caption(story, styles, context.figure_caption(f"Resultados consolidados de {career['name']}"))
        story += [base.fit_image(chart, 16 * cm, 9 * cm), Spacer(1, 0.1 * cm)]
        _pdf_caption(story, styles, "Nota. Elaboración propia con base en las calificaciones procesadas por Informtit.")


def _pdf_projects(story: list[Any], context: ExportContext, styles: Any, report_id: int) -> None:
    data = get_projects(report_id)
    projects = data.get("projects", [])
    if not projects:
        return
    _pdf_heading(story, context, styles, 1, "Resultados del Trabajo de Titulación")
    summary_data = data["summary"]
    _pdf_body(story, styles, f"Se registraron {summary_data['total']} {_plural(int(summary_data['total']), 'estudiante')}; {summary_data['approved']} aprobaron y {summary_data['failed']} reprobaron. El promedio final fue {_fmt(summary_data['average_final'])}.")
    for index, project in enumerate(projects):
        _pdf_heading(story, context, styles, 2, project["full_name"], page_break=index > 0)
        _pdf_caption(story, styles, context.table_caption(f"Información del Trabajo de Titulación de {project['full_name']}"))
        info = [
            ["Cédula", project.get("identification") or "—", "Carrera", project.get("career_name") or "—"],
            ["Acta", project.get("act_number") or "—", "Fecha", project.get("act_date") or "—"],
            ["Trabajo escrito", _fmt(project.get("written_average")), "Defensa oral", _fmt(project.get("oral_average"))],
            ["Calificación final", _fmt(project.get("final_grade")), "Estado", "Aprobado" if (project.get("final_grade") or 0) >= 7 else "Reprobado"],
        ]
        table = Table(info, colWidths=[3 * cm, 4.1 * cm, 3 * cm, 6 * cm])
        table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.4, colors.grey), ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"), ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 7.5), ("VALIGN", (0, 0), (-1, -1), "TOP")]))
        story += [table, Spacer(1, 0.15 * cm)]
        _pdf_body(story, styles, f"Vocales: {project.get('vocal_1') or '—'}; {project.get('vocal_2') or '—'}; {project.get('vocal_3') or '—'}.")
        for evaluation_type, title in (("practical", "Evaluación práctica"), ("defense", "Evaluación de la defensa")):
            scores = [row for row in project.get("scores", []) if row["evaluation_type"] == evaluation_type]
            if not scores:
                continue
            _pdf_heading(story, context, styles, 3, title)
            _pdf_caption(story, styles, context.table_caption(f"{title} de {project['full_name']}"))
            rows = [[Paragraph(html.escape(row["criterion"]), styles["TableCell"]), _fmt(row["max_score"]), _fmt(row["vocal_1"]), _fmt(row["vocal_2"]), _fmt(row["vocal_3"])] for row in scores]
            story += [_pdf_table(["Criterio", "Máximo", "Vocal 1", "Vocal 2", "Vocal 3"], rows, [7.2 * cm, 2 * cm, 2.4 * cm, 2.4 * cm, 2.4 * cm]), Spacer(1, 0.2 * cm)]


def _pdf_post_sections(story: list[Any], context: ExportContext, styles: Any, report: dict[str, Any]) -> None:
    titles = {"analisis_estrategico": "Análisis estratégico", "conclusiones": "Conclusiones", "recomendaciones": "Recomendaciones"}
    for key in POST_KEYS:
        section = _section(report, key)
        if not section or not str(section.get("content") or "").strip():
            continue
        _pdf_heading(story, context, styles, 1, titles[key])
        for block in re.split(r"\n\s*\n", str(section["content"])):
            if block.strip():
                _pdf_body(story, styles, block.strip())


def build_pdf(report_id: int) -> Path:
    report_structure.ensure_structure_schema()
    base.EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    report = _report_data(report_id)
    output = base.EXPORT_DIR / f"informtit_{report_id}.pdf"
    styles = _pdf_styles()
    story: list[Any] = base.cover_pdf(report, styles)
    story.append(Paragraph("ÍNDICE", styles["Title"]))
    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle("TOC1", fontName="Helvetica-Bold", fontSize=10, leading=14, spaceBefore=4),
        ParagraphStyle("TOC2", fontName="Helvetica", fontSize=9, leading=12, leftIndent=14),
        ParagraphStyle("TOC3", fontName="Helvetica", fontSize=8, leading=11, leftIndent=28),
    ]
    story += [toc, PageBreak()]
    context = ExportContext.create()
    temp_paths: list[Path] = []

    _pdf_heading(story, context, styles, 1, "Introducción")
    for paragraph in report_structure.introduction(report, report_id):
        _pdf_body(story, styles, paragraph)
    _pdf_legal(story, context, styles, report)
    _pdf_regulation(story, context, styles, report)
    _pdf_methodology(story, context, styles, report, temp_paths)
    _pdf_requirements(story, context, styles, report_id)
    _pdf_schedules(story, context, styles, report_id)
    _pdf_nucleus_results(story, context, styles, report_id)
    _pdf_complexive(story, context, styles, report, temp_paths)
    _pdf_projects(story, context, styles, report_id)
    _pdf_post_sections(story, context, styles, report)

    images = _additional_images(report)
    if images:
        _pdf_heading(story, context, styles, 1, "Anexos")
        for image in images:
            path = base.image_path(image)
            if not path:
                continue
            title = str(image.get("title") or image.get("original_name") or "Evidencia")
            _pdf_caption(story, styles, context.figure_caption(title))
            story += [base.fit_image(path, 16.5 * cm, 20 * cm), Spacer(1, 0.1 * cm)]
            _pdf_caption(story, styles, f"Nota. {image.get('source') or 'Fuente institucional cargada en Informtit'}.")

    document = report_structure.TocDocTemplate(
        str(output),
        pagesize=A4,
        rightMargin=1.45 * cm,
        leftMargin=1.45 * cm,
        topMargin=3.4 * cm,
        bottomMargin=1.35 * cm,
        title=report["name"],
    )
    try:
        document.multiBuild(
            story,
            canvasmaker=lambda *args, **kwargs: base.NumberedCanvas(*args, report=report, **kwargs),
        )
    finally:
        for path in temp_paths:
            path.unlink(missing_ok=True)
    return output


def install() -> None:
    core.build_docx = build_docx
    core.build_pdf = build_pdf
    base.legacy._default_before = _phase_before
    base.legacy._default_after = _phase_after
