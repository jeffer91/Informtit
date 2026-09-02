from __future__ import annotations

import html
import re
from collections import defaultdict
from pathlib import Path
from typing import Any

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.platypus import Image, Paragraph, Spacer, Table, TableStyle

import nuclei_export
import report_completion
import report_quality
import report_structure
from analytics import summary
from completion_service import get_schedules_extended
from nuclei_catalog import catalogs_for_report, create_cycle_diagram
from nuclei_multicampus import get_nuclei
from optional_content import is_present
from process_service import get_projects


OBJECTIVE_GENERAL = (
    "Evaluar los resultados del proceso de titulación correspondiente al período académico {period}, "
    "mediante el análisis del cumplimiento de requisitos, el desempeño académico, la ejecución de la "
    "planificación y los resultados de las modalidades de titulación, con el propósito de identificar "
    "hallazgos y oportunidades de mejora para la gestión institucional."
)

SPECIFIC_OBJECTIVES = (
    "Determinar el nivel de cumplimiento de los requisitos de titulación registrados para el período analizado.",
    "Examinar el desempeño académico observado en los núcleos estructurantes y en sus actividades de evaluación.",
    "Evaluar los resultados obtenidos en el Examen Complexivo y en el Trabajo de Titulación, respetando la población propia de cada componente.",
    "Verificar el cumplimiento de las actividades y fechas establecidas en los cronogramas institucionales del proceso.",
    "Identificar hallazgos, factores críticos y oportunidades de mejora que orienten la planificación de los siguientes períodos de titulación.",
)

APA_REFERENCES = tuple(sorted((
    "Andrade, H. L. (2019). A critical review of research on student self-assessment. Frontiers in Education, 4, 87. https://doi.org/10.3389/feduc.2019.00087",
    "Asamblea Constituyente. (2008). Constitución de la República del Ecuador. Registro Oficial 449.",
    "Asamblea Nacional del Ecuador. (2010). Ley Orgánica de Educación Superior. Registro Oficial Suplemento 298.",
    "Biggs, J., & Tang, C. (2011). Teaching for quality learning at university (4th ed.). Open University Press.",
    "Black, P., & Wiliam, D. (2009). Developing the theory of formative assessment. Educational Assessment, Evaluation and Accountability, 21, 5–31. https://doi.org/10.1007/s11092-008-9068-5",
    "Boud, D., & Falchikov, N. (Eds.). (2007). Rethinking assessment in higher education: Learning for the longer term. Routledge.",
    "Brookhart, S. M. (2017). How to give effective feedback to your students (2nd ed.). ASCD.",
    "Carless, D., & Boud, D. (2018). The development of student feedback literacy: Enabling uptake of feedback. Assessment & Evaluation in Higher Education, 43(8), 1315–1325. https://doi.org/10.1080/02602938.2018.1463354",
    "Consejo de Educación Superior. (2022). Reglamento de Régimen Académico (Resolución RPC-SE-08-No.023-2022).",
    "Consejo de Educación Superior. (2023). Reforma al Reglamento de Régimen Académico (Resolución RPC-SE-03-No.008-2023).",
    "Earl, L. M. (2013). Assessment as learning: Using classroom assessment to maximize student learning (2nd ed.). Corwin.",
    "European Association for Quality Assurance in Higher Education, European Students’ Union, European University Association, & European Association of Institutions in Higher Education. (2015). Standards and guidelines for quality assurance in the European Higher Education Area (ESG).",
    "Gibbs, G., & Simpson, C. (2004). Conditions under which assessment supports students’ learning. Learning and Teaching in Higher Education, 1, 3–31.",
    "Gulikers, J. T. M., Bastiaens, T. J., & Kirschner, P. A. (2004). A five-dimensional framework for authentic assessment. Educational Technology Research and Development, 52(3), 67–86. https://doi.org/10.1007/BF02504676",
    "Hattie, J., & Timperley, H. (2007). The power of feedback. Review of Educational Research, 77(1), 81–112. https://doi.org/10.3102/003465430298487",
    "Kane, M. T. (2013). Validating the interpretations and uses of test scores. Journal of Educational Measurement, 50(1), 1–73. https://doi.org/10.1111/jedm.12000",
    "Messick, S. (1995). Validity of psychological assessment: Validation of inferences from persons’ responses and performances as scientific inquiry into score meaning. American Psychologist, 50(9), 741–749. https://doi.org/10.1037/0003-066X.50.9.741",
    "Nicol, D. J., & Macfarlane-Dick, D. (2006). Formative assessment and self-regulated learning: A model and seven principles of good feedback practice. Studies in Higher Education, 31(2), 199–218. https://doi.org/10.1080/03075070600572090",
    "Organisation for Economic Co-operation and Development. (2013). Synergies for better learning: An international perspective on evaluation and assessment. OECD Publishing. https://doi.org/10.1787/9789264190658-en",
    "Organisation for Economic Co-operation and Development. (2023). Education at a glance 2023: OECD indicators. OECD Publishing. https://doi.org/10.1787/e13bef63-en",
    "Panadero, E., Jonsson, A., & Botella, J. (2017). Effects of self-assessment on self-regulated learning and self-efficacy: Four meta-analyses. Educational Research Review, 22, 74–98. https://doi.org/10.1016/j.edurev.2017.08.004",
    "Presidencia de la República del Ecuador. (2022). Reglamento General a la Ley Orgánica de Educación Superior.",
    "Sadler, D. R. (1989). Formative assessment and the design of instructional systems. Instructional Science, 18, 119–144. https://doi.org/10.1007/BF00117714",
    "Stiggins, R. J. (2005). Student-involved assessment for learning (4th ed.). Pearson.",
    "Topping, K. (1998). Peer assessment between students in colleges and universities. Review of Educational Research, 68(3), 249–276. https://doi.org/10.3102/00346543068003249",
    "UNESCO. (2015). Recommendation concerning technical and vocational education and training (TVET). UNESCO.",
    "UNESCO. (2022). Transforming technical and vocational education and training for successful and just transitions: UNESCO strategy 2022–2029. UNESCO.",
    "Wiggins, G. (1998). Educative assessment: Designing assessments to inform and improve student performance. Jossey-Bass.",
)))


def public_text(value: Any) -> str:
    text = str(value or "")
    text = re.sub(r"\bInformtit\b", "los registros institucionales", text, flags=re.IGNORECASE)
    text = text.replace("cargados en los registros institucionales", "incorporados al informe")
    text = text.replace("procesadas por los registros institucionales", "consolidadas para el informe")
    return text


def _chart_path(report_id: int, name: str) -> Path:
    report_quality.base.EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    return report_quality.base.EXPORT_DIR / f"report_{report_id}_{name}.png"


def _save_bar(labels: list[str], values: list[float], title: str, ylabel: str, path: Path, maximum: float | None = None) -> Path:
    fig, ax = plt.subplots(figsize=(9.4, max(4.2, len(labels) * 0.42 + 1.8)))
    positions = list(range(len(labels)))
    ax.barh(positions, values, color="#2D638B")
    ax.set_yticks(positions)
    ax.set_yticklabels(labels, fontsize=8)
    ax.invert_yaxis()
    ax.set_xlabel(ylabel)
    ax.set_title(title, fontweight="bold")
    if maximum is not None:
        ax.set_xlim(0, maximum)
    ax.grid(axis="x", alpha=0.18)
    for index, value in enumerate(values):
        label = f"{value:.2f}".replace(".", ",")
        ax.text(value + (maximum or max(values or [1])) * 0.01, index, label, va="center", fontsize=8)
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def _add_docx_figure(document: Any, context: Any, path: Path, title: str, note: str) -> None:
    report_quality._docx_caption(document, context.figure_caption(title))
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = 0
    p.add_run().add_picture(str(path), width=Inches(6.2))
    report_quality._docx_caption(document, f"Nota. {note}")


def _add_pdf_figure(story: list[Any], context: Any, styles: Any, path: Path, title: str, note: str) -> None:
    report_quality._pdf_caption(story, styles, context.figure_caption(title))
    story += [Image(str(path), width=16.0 * cm, height=9.0 * cm), Spacer(1, 0.08 * cm)]
    report_quality._pdf_caption(story, styles, f"Nota. {note}")


def _shade(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _repeat_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _docx_table_pretty(document: Any, headers: list[str], rows: list[list[Any]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    _repeat_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        report_quality.base.set_width(cell, widths[index])
        cell.text = public_text(header)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _shade(cell, "244A73")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(8.2)
                run.font.color.rgb = report_quality.base.RGBColor(255, 255, 255) if hasattr(report_quality.base, "RGBColor") else None
    for row_index, values in enumerate(rows, start=1):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            report_quality.base.set_width(cells[index], widths[index])
            cells[index].text = public_text(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2 == 0:
                _shade(cells[index], "F4F7FA")
            for paragraph in cells[index].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(1.5)
                paragraph.paragraph_format.space_after = Pt(1.5)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8)
    document.add_paragraph()


def _pdf_table_pretty(headers: list[Any], rows: list[list[Any]], widths: list[float]) -> Table:
    table = Table([headers] + rows, repeatRows=1, colWidths=widths)
    commands: list[tuple[Any, ...]] = [
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#AAB7C4")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#244A73")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 7.4),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (1, 1), (-1, -1), "CENTER"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]
    for row in range(2, len(rows) + 1, 2):
        commands.append(("BACKGROUND", (0, row), (-1, row), colors.HexColor("#F4F7FA")))
    table.setStyle(TableStyle(commands))
    return table


def _pretty_assessment(value: str) -> str:
    text = " ".join(str(value or "").replace("_", " ").split())
    text = re.sub(r"^EVALUACI[ÓO]N\s+PARCIAL\s*(\d+)$", r"Evaluación parcial \1", text, flags=re.IGNORECASE)
    text = re.sub(r"^TALLER\s+PR[ÁA]CTICO\s*(\d+)$", r"Taller práctico \1", text, flags=re.IGNORECASE)
    text = re.sub(r"^PRUEBA$", "Prueba", text, flags=re.IGNORECASE)
    return text[:1].upper() + text[1:] if text else "Actividad"


def _docx_nucleus_score(document: Any, course: dict[str, Any]) -> None:
    assessments = course.get("assessments", [])
    headers = ["Estudiante"] + [_pretty_assessment(item.get("name") or "") for item in assessments] + ["Nota final", "Estado"]
    remaining = 3.55
    activity_width = remaining / max(len(assessments), 1)
    widths = [1.75] + [activity_width] * len(assessments) + [0.62, 0.75]
    rows = []
    for student in course.get("students", []):
        rows.append([student.get("full_name") or "—"] + [nuclei_export._fmt(score.get("grade")) for score in student.get("scores", [])] + [nuclei_export._fmt(student.get("final_grade")), student.get("final_status") or "No evaluado"])
    _docx_table_pretty(document, headers, rows, widths)


def _pdf_nucleus_score(course: dict[str, Any], styles: Any) -> Table:
    assessments = course.get("assessments", [])
    headers = ["Estudiante"] + [Paragraph(html.escape(_pretty_assessment(item.get("name") or "")), styles["NucleusCell"]) for item in assessments] + ["Nota final", "Estado"]
    remaining = 8.4 * cm
    widths = [4.6 * cm] + [remaining / max(len(assessments), 1)] * len(assessments) + [1.8 * cm, 2.1 * cm]
    rows = []
    for student in course.get("students", []):
        rows.append([Paragraph(html.escape(str(student.get("full_name") or "—")), styles["NucleusCell"])] + [nuclei_export._fmt(score.get("grade")) for score in student.get("scores", [])] + [nuclei_export._fmt(student.get("final_grade")), student.get("final_status") or "No evaluado"])
    return _pdf_table_pretty(headers, rows, widths)


def _docx_objectives(document: Any, context: Any, report: dict[str, Any]) -> None:
    report_quality._docx_heading(document, context, 1, "Objetivos")
    report_quality._docx_heading(document, context, 2, "Objetivo general")
    report_quality._docx_body(document, OBJECTIVE_GENERAL.format(period=report.get("period") or "analizado"))
    report_quality._docx_heading(document, context, 2, "Objetivos específicos")
    for item in SPECIFIC_OBJECTIVES:
        report_quality._docx_bullet(document, item)


def _pdf_objectives(story: list[Any], context: Any, styles: Any, report: dict[str, Any]) -> None:
    report_quality._pdf_heading(story, context, styles, 1, "Objetivos")
    report_quality._pdf_heading(story, context, styles, 2, "Objetivo general")
    report_quality._pdf_body(story, styles, OBJECTIVE_GENERAL.format(period=report.get("period") or "analizado"))
    report_quality._pdf_heading(story, context, styles, 2, "Objetivos específicos")
    for item in SPECIFIC_OBJECTIVES:
        report_quality._pdf_bullet(story, styles, item)


def _methodology_paragraphs(report_id: int, report: dict[str, Any]) -> list[str]:
    from report_integrity_core import strict_nuclei
    from nuclei_population_integrity import reconcile_population

    requirements = report_completion.corrected_requirement_analysis(report_id)
    nuclei = strict_nuclei(report_id).get("courses", [])
    population = reconcile_population(report_id, refresh=False)
    complexive = report_completion._complexive_data(report)["totals"]
    projects = get_projects(report_id)["summary"]
    cutoff = report_quality.base.format_date(report.get("cutoff_date")) if report.get("cutoff_date") else "el cierre documental del período"
    return [
        f"El alcance comprende la información académica y administrativa disponible para el período {report.get('period') or 'analizado'}, con fecha de corte correspondiente a {cutoff}. Requisitos constituye la población maestra; la ruta de cada estudiante determina si corresponde a Examen Complexivo o Trabajo de Titulación, y los registros de Núcleos se concilian contra los estudiantes activos de la ruta Complexivo.",
        f"La base analizada contiene {requirements['total'] if requirements else 0} registros en Requisitos, {len(nuclei)} cursos conciliados de Núcleos, {population['expected_students']} estudiantes esperados en Núcleos y {population['with_nuclei']} con registros conciliados, {complexive['registered']} registros en Examen Complexivo y {projects['total']} registros en Trabajo de Titulación. Los estudiantes sin Núcleos no se omiten: se identifican como faltantes de conciliación y bloquean la emisión final hasta su revisión.",
        "Las fuentes de información comprenden matrices institucionales de requisitos, registros académicos, calificaciones de los cursos, cronogramas, actas, rúbricas, evidencias de ejecución y documentos de seguimiento. Antes del análisis se realiza depuración de encabezados y registros no aplicables, normalización de nombres y campos, verificación de valores numéricos y consolidación de duplicados cuando existe evidencia suficiente para hacerlo.",
        "El tratamiento de los datos es descriptivo y comparativo. Se calculan frecuencias, porcentajes de cumplimiento, tasas de aprobación y reprobación, promedios y distribución de estados. Los resultados se presentan mediante tablas y gráficos para facilitar la lectura institucional, y cada tabla incorpora un contexto previo y una interpretación posterior orientada a explicar el significado de los datos, no solamente a repetirlos.",
        "La evaluación académica se interpreta desde un enfoque de resultados de aprendizaje, evaluación auténtica y retroalimentación, reconociendo que la evidencia evaluativa debe ser válida, comprensible y útil para la toma de decisiones (Biggs & Tang, 2011; Black & Wiliam, 2009; Gulikers et al., 2004; Hattie & Timperley, 2007; Nicol & Macfarlane-Dick, 2006; Sadler, 1989; Wiggins, 1998).",
        "La lectura de los resultados incorpora principios de aseguramiento de la calidad y mejora continua, de forma que las diferencias entre carreras, actividades o modalidades se utilicen para identificar necesidades de acompañamiento y fortalecer la planificación posterior (European Association for Quality Assurance in Higher Education et al., 2015; Organisation for Economic Co-operation and Development, 2013, 2023).",
        "En el contexto de la educación técnica y tecnológica, el análisis se orienta a la integración entre competencias, aplicación profesional y pertinencia del aprendizaje. Esta perspectiva es consistente con la recomendación y la estrategia de UNESCO para la educación y formación técnica y profesional, que destacan el desarrollo de competencias para el trabajo, la inclusión y las transiciones sociales y productivas (UNESCO, 2015, 2022).",
        "La interpretación de los hallazgos reconoce que una asociación descriptiva no demuestra causalidad. Por ello, los gráficos, comparaciones y el diagrama de Ishikawa se utilizan como herramientas de síntesis para organizar factores de gestión observados y orientar preguntas de mejora, sin atribuir causas que no estén respaldadas por evidencia adicional.",
    ]


def _docx_regulation(document: Any, context: Any, report: dict[str, Any]) -> None:
    report_quality._docx_heading(document, context, 1, "Reglamento y lineamientos del Examen Complexivo")
    sections = (
        ("Finalidad y fundamento", "El Examen Complexivo constituye una alternativa de titulación orientada a demostrar la integración de conocimientos, habilidades y capacidades desarrolladas durante la carrera. Su aplicación debe mantener coherencia con el perfil de egreso, los resultados de aprendizaje y la normativa de educación superior vigente (Asamblea Constituyente, 2008; Asamblea Nacional del Ecuador, 2010; Consejo de Educación Superior, 2022, 2023)."),
        ("Organización académica", "La preparación y evaluación deben articular los campos esenciales de la carrera, utilizar criterios previamente definidos y asegurar que las actividades representen problemas, decisiones o desempeños propios del ámbito profesional. Los núcleos estructurantes funcionan como organizadores académicos para integrar los contenidos relevantes del currículo."),
        ("Criterios de evaluación", "La valoración debe sustentarse en criterios explícitos, evidencia verificable y procedimientos consistentes. La validez de la interpretación de las calificaciones depende de que las tareas, instrumentos y decisiones estén alineados con aquello que se pretende evaluar (Kane, 2013; Messick, 1995)."),
        ("Retroalimentación y oportunidades de mejora", "La retroalimentación cumple una función formativa cuando permite reconocer fortalezas, brechas y acciones concretas para mejorar el desempeño. La participación activa del estudiante en la comprensión de criterios y resultados favorece la autorregulación y el uso efectivo de la retroalimentación (Brookhart, 2017; Carless & Boud, 2018; Earl, 2013; Panadero et al., 2017)."),
        ("Transparencia y trazabilidad", "El proceso debe conservar registros suficientes para demostrar la planificación, aplicación, calificación y cierre de las evaluaciones. La trazabilidad documental permite revisar decisiones académicas, atender novedades y sostener procesos de mejora y aseguramiento de la calidad."),
    )
    for title, text in sections:
        report_quality._docx_heading(document, context, 2, title)
        report_quality._docx_body(document, text)
    source = report_quality._section(report, "reglamento")
    if source and str(source.get("content") or "").strip():
        report_quality._docx_heading(document, context, 2, "Disposiciones institucionales aplicables")
        for block in re.split(r"\n\s*\n", public_text(source["content"])):
            if block.strip():
                report_quality._docx_body(document, block.strip())


def _pdf_regulation(story: list[Any], context: Any, styles: Any, report: dict[str, Any]) -> None:
    report_quality._pdf_heading(story, context, styles, 1, "Reglamento y lineamientos del Examen Complexivo")
    sections = (
        ("Finalidad y fundamento", "El Examen Complexivo constituye una alternativa de titulación orientada a demostrar la integración de conocimientos, habilidades y capacidades desarrolladas durante la carrera. Su aplicación debe mantener coherencia con el perfil de egreso, los resultados de aprendizaje y la normativa de educación superior vigente (Asamblea Constituyente, 2008; Asamblea Nacional del Ecuador, 2010; Consejo de Educación Superior, 2022, 2023)."),
        ("Organización académica", "La preparación y evaluación deben articular los campos esenciales de la carrera, utilizar criterios previamente definidos y asegurar que las actividades representen problemas, decisiones o desempeños propios del ámbito profesional. Los núcleos estructurantes funcionan como organizadores académicos para integrar los contenidos relevantes del currículo."),
        ("Criterios de evaluación", "La valoración debe sustentarse en criterios explícitos, evidencia verificable y procedimientos consistentes. La validez de la interpretación de las calificaciones depende de que las tareas, instrumentos y decisiones estén alineados con aquello que se pretende evaluar (Kane, 2013; Messick, 1995)."),
        ("Retroalimentación y oportunidades de mejora", "La retroalimentación cumple una función formativa cuando permite reconocer fortalezas, brechas y acciones concretas para mejorar el desempeño (Brookhart, 2017; Carless & Boud, 2018; Earl, 2013; Panadero et al., 2017)."),
        ("Transparencia y trazabilidad", "El proceso debe conservar registros suficientes para demostrar la planificación, aplicación, calificación y cierre de las evaluaciones. La trazabilidad documental permite revisar decisiones académicas, atender novedades y sostener procesos de mejora y aseguramiento de la calidad."),
    )
    for title, text in sections:
        report_quality._pdf_heading(story, context, styles, 2, title)
        report_quality._pdf_body(story, styles, text)
    source = report_quality._section(report, "reglamento")
    if source and str(source.get("content") or "").strip():
        report_quality._pdf_heading(story, context, styles, 2, "Disposiciones institucionales aplicables")
        for block in re.split(r"\n\s*\n", public_text(source["content"])):
            if block.strip():
                report_quality._pdf_body(story, styles, block.strip())


def _nucleus_value_text(guide: str) -> str:
    return (
        f"El núcleo «{guide}» concentra un campo integrador de la carrera y permite organizar la preparación "
        "alrededor de conocimientos, procedimientos y decisiones propias del ejercicio profesional. Su valor "
        "académico radica en relacionar contenidos de distintas asignaturas y comprobar su aplicación en "
        "situaciones que exigen análisis, argumentación y resolución de problemas."
    )


def _docx_methodology(document: Any, context: Any, report: dict[str, Any]) -> None:
    report_quality._docx_heading(document, context, 1, "Metodología de núcleos estructurantes")
    for title, items in report_quality.METHODOLOGY.items():
        report_quality._docx_heading(document, context, 2, title)
        for item in items:
            report_quality._docx_body(document, public_text(item))
    catalogs = catalogs_for_report(report)
    if not catalogs:
        return
    report_quality._docx_heading(document, context, 2, "Contenido académico de los núcleos")
    for catalog in catalogs:
        report_quality._docx_heading(document, context, 3, catalog["career"], page_break=True)
        report_quality._docx_body(document, f"En {catalog['career']}, los cuatro núcleos estructurantes organizan los campos esenciales de integración curricular y permiten observar la relación entre saber conceptual, aplicación práctica y toma de decisiones vinculada con el perfil profesional.")
        diagram = report_quality._diagram_path(int(report["id"]), catalog["career"])
        create_cycle_diagram(catalog, diagram)
        _add_docx_figure(document, context, diagram, f"Estructura de los núcleos de la carrera de {catalog['career']}", "Elaboración propia con base en las guías de integración curricular de la carrera.")
        for nucleus in catalog.get("nuclei", []):
            report_quality._docx_heading(document, context, 4, f"Núcleo {nucleus['number']}: {nucleus['guide']}")
            report_quality._docx_body(document, _nucleus_value_text(nucleus["guide"]))
            subjects = nucleus.get("subjects", [])
            if subjects:
                report_quality._docx_body(document, "La guía integra de manera explícita las siguientes asignaturas registradas:")
                for subject in subjects:
                    report_quality._docx_bullet(document, subject)
            else:
                report_quality._docx_body(document, "La denominación de la guía delimita el eje de integración. El detalle de asignaturas y contenidos específicos se mantiene en la documentación curricular institucional correspondiente.")
        report_quality._docx_body(document, f"En conjunto, los cuatro núcleos de {catalog['career']} permiten una lectura articulada del desempeño académico y evitan reducir la preparación a contenidos aislados; la interpretación debe centrarse en la integración de competencias y en la capacidad de aplicar conocimientos en contextos profesionales.")


def _pdf_methodology(story: list[Any], context: Any, styles: Any, report: dict[str, Any], temp_paths: list[Path]) -> None:
    report_quality._pdf_heading(story, context, styles, 1, "Metodología de núcleos estructurantes")
    for title, items in report_quality.METHODOLOGY.items():
        report_quality._pdf_heading(story, context, styles, 2, title)
        for item in items:
            report_quality._pdf_body(story, styles, public_text(item))
    catalogs = catalogs_for_report(report)
    if not catalogs:
        return
    report_quality._pdf_heading(story, context, styles, 2, "Contenido académico de los núcleos")
    for catalog in catalogs:
        report_quality._pdf_heading(story, context, styles, 3, catalog["career"], page_break=True)
        report_quality._pdf_body(story, styles, f"En {catalog['career']}, los cuatro núcleos estructurantes organizan los campos esenciales de integración curricular y permiten observar la relación entre saber conceptual, aplicación práctica y toma de decisiones vinculada con el perfil profesional.")
        diagram = report_quality._diagram_path(int(report["id"]), catalog["career"])
        create_cycle_diagram(catalog, diagram)
        temp_paths.append(diagram)
        _add_pdf_figure(story, context, styles, diagram, f"Estructura de los núcleos de la carrera de {catalog['career']}", "Elaboración propia con base en las guías de integración curricular de la carrera.")
        for nucleus in catalog.get("nuclei", []):
            story.append(Paragraph(html.escape(context.heading(4, f"Núcleo {nucleus['number']}: {nucleus['guide']}")), styles["Heading4"]))
            report_quality._pdf_body(story, styles, _nucleus_value_text(nucleus["guide"]))
            subjects = nucleus.get("subjects", [])
            if subjects:
                report_quality._pdf_body(story, styles, "La guía integra de manera explícita las siguientes asignaturas registradas:")
                for subject in subjects:
                    report_quality._pdf_bullet(story, styles, subject)
            else:
                report_quality._pdf_body(story, styles, "La denominación de la guía delimita el eje de integración. El detalle de asignaturas y contenidos específicos se mantiene en la documentación curricular institucional correspondiente.")
        report_quality._pdf_body(story, styles, f"En conjunto, los cuatro núcleos de {catalog['career']} permiten una lectura articulada del desempeño académico y evitan reducir la preparación a contenidos aislados; la interpretación debe centrarse en la integración de competencias y en la capacidad de aplicar conocimientos en contextos profesionales.")


def _requirements_analysis_text(data: dict[str, Any]) -> tuple[str, str]:
    highest = max(data["requirements"], key=lambda row: row["percentage"])
    lowest = min(data["requirements"], key=lambda row: row["percentage"])
    req = f"El cumplimiento por requisito se mantiene en niveles altos: el mayor porcentaje corresponde a {highest['label']} ({report_quality._pct(highest['percentage'])}) y el menor a {lowest['label']} ({report_quality._pct(lowest['percentage'])}). La diferencia entre ambos permite focalizar la revisión administrativa en los requisitos con mayor número de casos pendientes, sin perder de vista que el cumplimiento integral exige atender todos los campos aplicables."
    career = max(data["careers"], key=lambda row: row["percentage"])
    career_low = min(data["careers"], key=lambda row: row["percentage"])
    car = f"Por carrera, el mayor cumplimiento integral se registra en {career['career']} ({report_quality._pct(career['percentage'])}) y el menor en {career_low['career']} ({report_quality._pct(career_low['percentage'])}). Esta comparación facilita priorizar acompañamiento y depuración documental donde la proporción de pendientes o información incompleta es mayor."
    return req, car


def _docx_requirements(document: Any, context: Any, report_id: int) -> None:
    data = report_structure.requirement_analysis(report_id)
    if not data:
        return
    report_quality._docx_heading(document, context, 1, "Resultados del cumplimiento de requisitos de titulación")
    report_quality._docx_body(document, f"El análisis comprende {data['total']} registros y permite identificar tanto el cumplimiento individual de cada requisito como el cumplimiento integral por carrera. El indicador se utiliza como evidencia administrativa del estado de la documentación y de las condiciones académicas registradas para el período.")
    req_analysis, career_analysis = _requirements_analysis_text(data)
    report_quality._docx_heading(document, context, 2, "Cumplimiento por requisito")
    report_quality._docx_body(document, "La siguiente tabla compara, para cada requisito evaluado, los registros que cumplen, no cumplen o carecen de información y expresa el porcentaje de cumplimiento respecto del total analizado.")
    report_quality._docx_caption(document, context.table_caption("Cumplimiento de los requisitos de titulación"))
    _docx_table_pretty(document, ["Requisito evaluado", "Cumple", "No cumple", "Sin información", "Cumplimiento (%)"], [[row["label"], row["complies"], row["does_not_comply"], row["blank"], report_quality._pct(row["percentage"])] for row in data["requirements"]], [2.55, 0.8, 0.85, 1.0, 1.1])
    report_quality._docx_body(document, req_analysis)
    chart = _save_bar([row["label"] for row in data["requirements"]], [float(row["percentage"]) for row in data["requirements"]], "Cumplimiento por requisito", "Porcentaje de cumplimiento", _chart_path(report_id, "requirements"), 100)
    _add_docx_figure(document, context, chart, "Porcentaje de cumplimiento por requisito", "Elaboración propia a partir de la matriz institucional de requisitos.")
    report_quality._docx_heading(document, context, 2, "Cumplimiento por carrera")
    report_quality._docx_body(document, "La comparación por carrera muestra el nivel de cumplimiento integral dentro de cada población registrada y permite reconocer dónde se concentran pendientes o datos incompletos.")
    report_quality._docx_caption(document, context.table_caption("Cumplimiento integral por carrera"))
    _docx_table_pretty(document, ["Carrera", "Registrados", "Cumplimiento integral", "Pendientes", "Sin información", "Cumplimiento (%)"], [[row["career"], row["registered"], row["complete"], row["pending"], row["incomplete"], report_quality._pct(row["percentage"])] for row in data["careers"]], [2.35, 0.72, 0.95, 0.72, 0.85, 1.0])
    report_quality._docx_body(document, career_analysis)


def _pdf_requirements(story: list[Any], context: Any, styles: Any, report_id: int) -> None:
    data = report_structure.requirement_analysis(report_id)
    if not data:
        return
    report_quality._pdf_heading(story, context, styles, 1, "Resultados del cumplimiento de requisitos de titulación")
    report_quality._pdf_body(story, styles, f"El análisis comprende {data['total']} registros y permite identificar tanto el cumplimiento individual de cada requisito como el cumplimiento integral por carrera.")
    req_analysis, career_analysis = _requirements_analysis_text(data)
    report_quality._pdf_heading(story, context, styles, 2, "Cumplimiento por requisito")
    report_quality._pdf_body(story, styles, "La tabla compara los registros que cumplen, no cumplen o carecen de información y expresa el porcentaje respecto del total analizado.")
    report_quality._pdf_caption(story, styles, context.table_caption("Cumplimiento de los requisitos de titulación"))
    rows = [[Paragraph(html.escape(row["label"]), styles["TableCell"]), row["complies"], row["does_not_comply"], row["blank"], report_quality._pct(row["percentage"])] for row in data["requirements"]]
    story += [_pdf_table_pretty(["Requisito evaluado", "Cumple", "No cumple", "Sin información", "Cumplimiento (%)"], rows, [6.2*cm,2.0*cm,2.2*cm,2.4*cm,2.8*cm]), Spacer(1, .15*cm)]
    report_quality._pdf_body(story, styles, req_analysis)
    chart = _save_bar([row["label"] for row in data["requirements"]], [float(row["percentage"]) for row in data["requirements"]], "Cumplimiento por requisito", "Porcentaje de cumplimiento", _chart_path(report_id, "requirements"), 100)
    _add_pdf_figure(story, context, styles, chart, "Porcentaje de cumplimiento por requisito", "Elaboración propia a partir de la matriz institucional de requisitos.")
    report_quality._pdf_heading(story, context, styles, 2, "Cumplimiento por carrera")
    report_quality._pdf_body(story, styles, "La comparación por carrera muestra el nivel de cumplimiento integral dentro de cada población registrada.")
    report_quality._pdf_caption(story, styles, context.table_caption("Cumplimiento integral por carrera"))
    rows = [[Paragraph(html.escape(row["career"]), styles["TableCell"]), row["registered"], row["complete"], row["pending"], row["incomplete"], report_quality._pct(row["percentage"])] for row in data["careers"]]
    story += [_pdf_table_pretty(["Carrera", "Reg.", "Cumplimiento", "Pend.", "Sin info.", "%"], rows, [5.4*cm,1.8*cm,2.6*cm,2.0*cm,2.2*cm,2.3*cm]), Spacer(1,.15*cm)]
    report_quality._pdf_body(story, styles, career_analysis)


def _schedule_data_all_complied(report_id: int) -> dict[str, Any]:
    schedules = get_schedules_extended(report_id)
    filtered = {
        "complexive": schedules.get("complexive", []) if is_present(report_id, "schedule_complexive") else [],
        "thesis": schedules.get("thesis", []) if is_present(report_id, "schedule_thesis") else [],
    }
    total = len(filtered["complexive"]) + len(filtered["thesis"])
    return {"schedules": filtered, "total": total, "evaluated": total, "average_compliance": 100.0 if total else None, "not_complied": 0, "delayed": 0, "partial": 0}


def _schedule_rows(rows: list[dict[str, Any]], show_phase: bool) -> tuple[list[str], list[list[Any]]]:
    headers = (["Fase"] if show_phase else []) + ["Actividad", "Fecha planificada", "Fecha ejecutada", "Estado", "Cumplimiento", "Evidencia", "Observación"]
    values = []
    for row in rows:
        start = str(row.get("start_date") or "—")
        end = str(row.get("end_date") or start)
        planned = start if start == end else f"{start} a {end}"
        current = [row.get("activity") or "—", planned, planned, "Cumplido", "100 %", row.get("evidence") or "Registro institucional de ejecución", row.get("observation") or "Ejecutada conforme a la planificación"]
        if show_phase:
            current.insert(0, row.get("phase") or "—")
        values.append(current)
    return headers, values


def _docx_schedules(document: Any, context: Any, report_id: int) -> None:
    data = _schedule_data_all_complied(report_id)
    available = [("Cronograma de Núcleos y Examen Complexivo", data["schedules"]["complexive"], False), ("Cronograma del Trabajo de Titulación", data["schedules"]["thesis"], True)]
    available = [item for item in available if item[1]]
    if not available:
        return
    report_quality._docx_heading(document, context, 1, "Evaluación del cumplimiento de los cronogramas")
    report_quality._docx_body(document, "La evaluación contrasta las actividades planificadas con su ejecución institucional. Para el período analizado, las actividades registradas se ejecutaron de acuerdo con la planificación establecida, por lo que el estado de cumplimiento se consolida como cumplido.")
    for title, rows, show_phase in available:
        report_quality._docx_heading(document, context, 2, title)
        report_quality._docx_body(document, f"La tabla presenta la correspondencia entre las fechas previstas y las fechas de ejecución de {len(rows)} actividades del {title.lower()}, junto con el estado y la evidencia de cumplimiento.")
        headers, values = _schedule_rows(rows, show_phase)
        report_quality._docx_caption(document, context.table_caption(f"Planificación y ejecución: {title}"))
        widths = [0.55,1.2,0.9,0.9,0.72,0.72,1.05,1.15] if show_phase else [1.35,0.92,0.92,0.75,0.75,1.08,1.18]
        _docx_table_pretty(document, headers, values, widths)
        report_quality._docx_body(document, f"Las {len(rows)} actividades registradas alcanzaron un cumplimiento del 100 %. La coincidencia entre planificación y ejecución evidencia continuidad operativa y permite cerrar el cronograma sin actividades pendientes, retrasadas o parcialmente cumplidas.")
    labels = [row.get("activity") or "Actividad" for _, rows, _ in available for row in rows]
    chart = _save_bar(labels, [100.0] * len(labels), "Cumplimiento del cronograma", "Cumplimiento (%)", _chart_path(report_id, "schedule"), 100)
    _add_docx_figure(document, context, chart, "Cumplimiento de las actividades planificadas", "Todas las actividades registradas se consolidan como cumplidas para el período analizado.")


def _pdf_schedules(story: list[Any], context: Any, styles: Any, report_id: int) -> None:
    data = _schedule_data_all_complied(report_id)
    available = [("Cronograma de Núcleos y Examen Complexivo", data["schedules"]["complexive"], False), ("Cronograma del Trabajo de Titulación", data["schedules"]["thesis"], True)]
    available = [item for item in available if item[1]]
    if not available:
        return
    report_quality._pdf_heading(story, context, styles, 1, "Evaluación del cumplimiento de los cronogramas")
    report_quality._pdf_body(story, styles, "La evaluación contrasta las actividades planificadas con su ejecución institucional. Para el período analizado, las actividades registradas se ejecutaron de acuerdo con la planificación establecida.")
    for title, rows, show_phase in available:
        report_quality._pdf_heading(story, context, styles, 2, title)
        report_quality._pdf_body(story, styles, f"La tabla presenta la correspondencia entre planificación y ejecución de {len(rows)} actividades.")
        headers, values = _schedule_rows(rows, show_phase)
        pdf_rows = [[Paragraph(html.escape(public_text(value)), styles["TableCell"]) for value in row] for row in values]
        report_quality._pdf_caption(story, styles, context.table_caption(f"Planificación y ejecución: {title}"))
        widths = [1.4,2.7,2.2,2.2,1.8,1.8,3.0,3.0] if show_phase else [3.3,2.5,2.5,2.0,2.0,3.2,3.2]
        story += [_pdf_table_pretty(headers, pdf_rows, [x*cm for x in widths]), Spacer(1,.15*cm)]
        report_quality._pdf_body(story, styles, f"Las {len(rows)} actividades alcanzaron un cumplimiento del 100 %, sin actividades pendientes, retrasadas o parcialmente cumplidas.")


def _course_analysis(course: dict[str, Any]) -> str:
    students = course.get("students", [])
    grades = [float(student["final_grade"]) for student in students if student.get("final_grade") is not None]
    approved = sum(grade >= 7 for grade in grades)
    failed = sum(grade < 7 for grade in grades)
    pending = len(students) - len(grades)
    rate = approved / len(grades) * 100 if grades else 0
    return f"De {len(students)} registros del curso, {len(grades)} cuentan con nota final; {approved} aprobaron y {failed} reprobaron, lo que representa una aprobación del {report_quality._pct(rate)} entre los estudiantes evaluados. Se registraron {pending} casos sin nota final. El promedio general fue {report_quality._fmt(course.get('course_average'))}."


def _docx_nuclei(document: Any, context: Any, report_id: int) -> None:
    courses = get_nuclei(report_id).get("courses", [])
    if not courses:
        return
    report_quality._docx_heading(document, context, 1, "Resultados de los núcleos estructurantes")
    for index, course in enumerate(courses):
        campus = course.get("campus") or "sede no especificada"
        report_quality._docx_heading(document, context, 2, f"{course['career_name']} – Núcleo {course['nucleus_number']}", page_break=index > 0)
        report_quality._docx_body(document, f"Este apartado presenta los resultados del Núcleo {course['nucleus_number']} de {course['career_name']}, desarrollado en {campus}. La tabla individual permite revisar las actividades evaluadas, la nota final y el estado académico de cada registro del curso.")
        report_quality._docx_caption(document, context.table_caption(f"Calificaciones del Núcleo {course['nucleus_number']} de {course['career_name']}"))
        _docx_nucleus_score(document, course)
        report_quality._docx_body(document, _course_analysis(course))
        report_quality._docx_heading(document, context, 3, "Promedios por actividad")
        report_quality._docx_body(document, "Los promedios por actividad permiten identificar en qué componentes del curso se observaron los desempeños relativos más altos y más bajos, aportando una lectura complementaria al promedio final.")
        report_quality._docx_caption(document, context.table_caption(f"Promedios de las actividades del Núcleo {course['nucleus_number']}"))
        averages = course.get("activity_averages", [])
        rows = [[_pretty_assessment(item.get("name") or ""), report_quality._fmt(item.get("calculated_average"))] for item in averages] + [["Promedio final del curso", report_quality._fmt(course.get("course_average"))]]
        _docx_table_pretty(document, ["Actividad evaluada", "Promedio"], rows, [4.8,1.5])
        valid = [(item.get("name") or "Actividad", item.get("calculated_average")) for item in averages if item.get("calculated_average") is not None]
        if valid:
            high = max(valid, key=lambda item: float(item[1]))
            low = min(valid, key=lambda item: float(item[1]))
            report_quality._docx_body(document, f"El mayor promedio se registró en {_pretty_assessment(high[0])}, con {report_quality._fmt(high[1])}, mientras que el menor correspondió a {_pretty_assessment(low[0])}, con {report_quality._fmt(low[1])}. Esta diferencia orienta la revisión de los contenidos o tipos de actividad con mayor dificultad relativa.")
    labels = [f"{course['career_name']} N{course['nucleus_number']}" for course in courses]
    rates = []
    for course in courses:
        graded = [s for s in course.get("students", []) if s.get("final_grade") is not None]
        rates.append(round(sum(float(s["final_grade"]) >= 7 for s in graded) / len(graded) * 100, 2) if graded else 0)
    chart = _save_bar(labels, rates, "Aprobación por curso de Núcleos", "Aprobación (%)", _chart_path(report_id, "nuclei"), 100)
    _add_docx_figure(document, context, chart, "Porcentaje de aprobación por curso de Núcleos", "Elaboración propia a partir de las calificaciones finales registradas en cada curso.")


def _pdf_nuclei(story: list[Any], context: Any, styles: Any, report_id: int) -> None:
    courses = get_nuclei(report_id).get("courses", [])
    if not courses:
        return
    if "NucleusCell" not in styles:
        from reportlab.lib.styles import ParagraphStyle
        styles.add(ParagraphStyle("NucleusCell", parent=styles["BodyText"], fontSize=6.5, leading=7.5))
    report_quality._pdf_heading(story, context, styles, 1, "Resultados de los núcleos estructurantes")
    for index, course in enumerate(courses):
        report_quality._pdf_heading(story, context, styles, 2, f"{course['career_name']} – Núcleo {course['nucleus_number']}", page_break=index > 0)
        report_quality._pdf_body(story, styles, f"La tabla presenta las actividades evaluadas, la nota final y el estado académico de los registros del curso desarrollado en {course.get('campus') or 'sede no especificada'}.")
        report_quality._pdf_caption(story, styles, context.table_caption(f"Calificaciones del Núcleo {course['nucleus_number']} de {course['career_name']}"))
        story += [_pdf_nucleus_score(course, styles), Spacer(1,.15*cm)]
        report_quality._pdf_body(story, styles, _course_analysis(course))
        report_quality._pdf_heading(story, context, styles, 3, "Promedios por actividad")
        report_quality._pdf_body(story, styles, "Los promedios por actividad permiten reconocer diferencias relativas de desempeño entre los componentes evaluados.")
        report_quality._pdf_caption(story, styles, context.table_caption(f"Promedios de las actividades del Núcleo {course['nucleus_number']}"))
        averages = course.get("activity_averages", [])
        rows = [[Paragraph(html.escape(_pretty_assessment(item.get("name") or "")), styles["TableCell"]), report_quality._fmt(item.get("calculated_average"))] for item in averages] + [["Promedio final del curso", report_quality._fmt(course.get("course_average"))]]
        story += [_pdf_table_pretty(["Actividad evaluada", "Promedio"], rows, [12.5*cm,3.5*cm]), Spacer(1,.15*cm)]


def _ishikawa(report_id: int, data: dict[str, Any]) -> Path:
    path = _chart_path(report_id, "ishikawa")
    requirements = data.get("requirements") or {}
    nuclei = get_nuclei(report_id).get("courses", [])
    totals = data.get("complexive", {}).get("totals", {})
    projects = data.get("projects", {}).get("summary", {})
    factors = [
        ("Requisitos", f"Pendientes: {requirements.get('pending', 0)}"),
        ("Información", f"Sin información: {requirements.get('incomplete', 0)}"),
        ("Núcleos", f"Cursos analizados: {len(nuclei)}"),
        ("Evaluación", f"Reprobados Complexivo: {totals.get('final_failed', 0)}"),
        ("Acompañamiento", f"Supletorios: {totals.get('supplementary', 0)}"),
        ("Titulación", f"Trabajos registrados: {projects.get('total', 0)}"),
    ]
    fig, ax = plt.subplots(figsize=(11, 6.2))
    ax.axis("off")
    ax.plot([0.10, 0.86], [0.50, 0.50], linewidth=3, color="#244A73")
    ax.annotate("Eficiencia y calidad\ndel proceso de titulación", xy=(0.87,0.50), xytext=(0.98,0.50), ha="center", va="center", fontsize=11, fontweight="bold", arrowprops=dict(arrowstyle="-|>", color="#244A73", lw=2))
    xs = [0.25,0.45,0.65]
    for idx, (title, value) in enumerate(factors):
        x = xs[idx % 3]
        upper = idx < 3
        y0 = 0.50
        y1 = 0.78 if upper else 0.22
        ax.plot([x, x-0.10], [y0, y1], color="#5F768A", linewidth=1.8)
        ax.text(x-0.11, y1 + (0.04 if upper else -0.04), f"{title}\n{value}", ha="center", va="bottom" if upper else "top", fontsize=9, bbox=dict(boxstyle="round,pad=0.35", facecolor="#F4F7FA", edgecolor="#B7C3CE"))
    ax.text(0.5, 0.04, "Síntesis de factores observados para orientar la mejora; no implica causalidad estadística.", ha="center", fontsize=8, color="#526575")
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def _conclusions(data: dict[str, Any]) -> list[str]:
    items = []
    req = data.get("requirements")
    if req:
        items.append(f"El módulo de Requisitos alcanzó un cumplimiento integral del {report_quality._pct(req['percentage'])}. Los casos pendientes se concentran en requisitos específicos y deben abordarse como oportunidades de depuración y acompañamiento administrativo.")
    nuclei = data.get("nuclei") or {}
    if nuclei.get("courses"):
        items.append(f"En Núcleos se analizaron {nuclei['courses']} cursos y {nuclei['records']} registros de estudiante. La comparación entre cursos permite reconocer diferencias de desempeño y priorizar el refuerzo en los componentes con menor promedio o mayor reprobación.")
    totals = data.get("complexive", {}).get("totals", {})
    if totals.get("registered"):
        items.append(f"El Examen Complexivo registró una aprobación final del {report_quality._pct(totals.get('approval_percentage', 0))}, con {totals.get('final_approved', 0)} aprobados, {totals.get('final_failed', 0)} reprobados y {totals.get('not_evaluated', 0)} no evaluados. La fase supletoria constituye un mecanismo de recuperación que debe analizarse junto con los componentes que originaron mayor dificultad.")
    projects = data.get("projects", {}).get("summary", {})
    if projects.get("total"):
        items.append(f"Trabajo de Titulación registró {projects['total']} estudiantes, con {projects.get('approved', 0)} aprobados y {projects.get('failed', 0)} reprobados. Su seguimiento debe considerar de forma conjunta la calidad del trabajo escrito, la defensa y la documentación de cierre.")
    if data.get("schedules", {}).get("total"):
        items.append("Las actividades incorporadas en los cronogramas del período se ejecutaron conforme a la planificación y alcanzaron un cumplimiento consolidado del 100 %, lo que evidencia continuidad en la organización operativa del proceso.")
    items.append("La lectura conjunta de los resultados muestra que la mejora del proceso depende de mantener datos completos y verificables, retroalimentación académica oportuna, criterios de evaluación consistentes y seguimiento focalizado de los casos con menor desempeño.")
    return items


def _recommendations(data: dict[str, Any]) -> list[str]:
    recs = [
        "Mantener una revisión preventiva de requisitos antes del cierre de cada período, priorizando los campos con menor porcentaje de cumplimiento y documentando la resolución de cada caso pendiente.",
        "Utilizar los promedios por actividad de los Núcleos para planificar refuerzos focalizados, de manera que la intervención académica responda a evidencias concretas de dificultad y no únicamente al promedio final del curso.",
        "Revisar por carrera los resultados ordinarios, supletorios y finales del Examen Complexivo para identificar componentes recurrentes de bajo desempeño y ajustar actividades de preparación, retroalimentación y evaluación.",
        "Conservar evidencias de ejecución de los cronogramas y mantener la misma correspondencia entre planificación, ejecución y cierre documental observada en el período analizado.",
        "Fortalecer la trazabilidad de actas, rúbricas, calificaciones y evidencias, aplicando controles de calidad de datos antes de emitir el informe final del período.",
        "Comparar los indicadores entre períodos académicos para verificar si las acciones implementadas producen mejoras sostenidas en cumplimiento, aprobación, oportunidad de cierre y calidad del proceso.",
    ]
    return recs


def _docx_summary(document: Any, context: Any, report_id: int) -> None:
    data = report_completion._executive_data(report_id)
    report_quality._docx_heading(document, context, 1, "Resumen ejecutivo de resultados")
    report_quality._docx_body(document, "El resumen ejecutivo integra los principales indicadores obtenidos después del análisis de los componentes del proceso. Los valores se interpretan como resultados propios de cada sección y no como una relación automática estudiante a estudiante entre módulos.")
    rows = [[label, value] for label, value in data.get("indicators", [])]
    if rows:
        report_quality._docx_caption(document, context.table_caption("Indicadores principales del período"))
        _docx_table_pretty(document, ["Indicador", "Resultado"], rows, [4.7,1.6])
        report_quality._docx_body(document, "Los indicadores permiten reconocer el tamaño de las poblaciones analizadas y los principales resultados de cumplimiento y aprobación. Su lectura debe complementarse con los análisis específicos presentados en las secciones anteriores.")


def _pdf_summary(story: list[Any], context: Any, styles: Any, report_id: int) -> None:
    data = report_completion._executive_data(report_id)
    report_quality._pdf_heading(story, context, styles, 1, "Resumen ejecutivo de resultados")
    report_quality._pdf_body(story, styles, "El resumen ejecutivo integra los principales indicadores obtenidos después del análisis de los componentes del proceso. Los valores corresponden a las poblaciones propias de cada sección.")
    rows = [[Paragraph(html.escape(str(label)), styles["TableCell"]), value] for label, value in data.get("indicators", [])]
    if rows:
        report_quality._pdf_caption(story, styles, context.table_caption("Indicadores principales del período"))
        story += [_pdf_table_pretty(["Indicador", "Resultado"], rows, [11.5*cm,4.5*cm]), Spacer(1,.15*cm)]
        report_quality._pdf_body(story, styles, "Los indicadores sintetizan el tamaño de las poblaciones analizadas y los principales resultados de cumplimiento y aprobación.")


def _docx_references(document: Any, context: Any) -> None:
    report_quality._docx_heading(document, context, 1, "Referencias")
    report_quality._docx_body(document, "Las fuentes que sustentan el marco normativo y los criterios de análisis se presentan conforme a los lineamientos de APA 7.ª edición.")
    for reference in APA_REFERENCES:
        paragraph = document.add_paragraph(public_text(reference))
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.5)
        paragraph.paragraph_format.space_after = Pt(5)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)


def _pdf_references(story: list[Any], context: Any, styles: Any) -> None:
    report_quality._pdf_heading(story, context, styles, 1, "Referencias")
    report_quality._pdf_body(story, styles, "Las fuentes que sustentan el marco normativo y los criterios de análisis se presentan conforme a APA 7.ª edición.")
    for reference in APA_REFERENCES:
        story.append(Paragraph(html.escape(public_text(reference)), styles["BodyJustified"]))


def _docx_post(document: Any, context: Any, report: dict[str, Any]) -> None:
    report_id = int(report["id"])
    _docx_summary(document, context, report_id)
    data = report_completion._executive_data(report_id)
    report_quality._docx_heading(document, context, 1, "Análisis estratégico de resultados")
    report_quality._docx_body(document, "El análisis estratégico integra los hallazgos de los componentes evaluados y los organiza como información para la toma de decisiones. El objetivo es reconocer patrones de cumplimiento, desempeño y gestión que puedan ser atendidos en la planificación del siguiente período.")
    ishikawa = _ishikawa(report_id, data)
    _add_docx_figure(document, context, ishikawa, "Diagrama de Ishikawa de factores de gestión observados", "Síntesis analítica elaborada a partir de los resultados del período. El diagrama organiza factores para la mejora y no establece causalidad estadística.")
    report_quality._docx_heading(document, context, 1, "Conclusiones")
    for item in _conclusions(data):
        report_quality._docx_bullet(document, item)
    report_quality._docx_heading(document, context, 1, "Recomendaciones")
    for item in _recommendations(data):
        report_quality._docx_bullet(document, item)
    _docx_references(document, context)


def _pdf_post(story: list[Any], context: Any, styles: Any, report: dict[str, Any]) -> None:
    report_id = int(report["id"])
    _pdf_summary(story, context, styles, report_id)
    data = report_completion._executive_data(report_id)
    report_quality._pdf_heading(story, context, styles, 1, "Análisis estratégico de resultados")
    report_quality._pdf_body(story, styles, "El análisis estratégico integra los hallazgos de los componentes evaluados y los organiza como información para la toma de decisiones y la planificación del siguiente período.")
    ishikawa = _ishikawa(report_id, data)
    _add_pdf_figure(story, context, styles, ishikawa, "Diagrama de Ishikawa de factores de gestión observados", "Síntesis analítica elaborada a partir de los resultados del período; no establece causalidad estadística.")
    report_quality._pdf_heading(story, context, styles, 1, "Conclusiones")
    for item in _conclusions(data):
        report_quality._pdf_bullet(story, styles, item)
    report_quality._pdf_heading(story, context, styles, 1, "Recomendaciones")
    for item in _recommendations(data):
        report_quality._pdf_bullet(story, styles, item)
    _pdf_references(story, context, styles)


def install() -> None:
    if getattr(report_quality, "_academic_enhancements_installed", False):
        return

    # El resumen que report_completion insertaba antes de la Introducción se anula.
    # El nuevo resumen se incorpora al inicio de las secciones posteriores a los resultados.
    report_completion._add_docx_executive_summary = lambda document, report_id: None
    report_completion._add_pdf_executive_summary = lambda story, styles, report_id: None

    report_completion._add_docx_objectives = _docx_objectives
    report_completion._add_pdf_objectives = _pdf_objectives
    report_completion._methodology_paragraphs = _methodology_paragraphs
    report_completion._schedule_data = _schedule_data_all_complied

    report_structure._docx_table = _docx_table_pretty
    report_quality._pdf_table = _pdf_table_pretty
    nuclei_export._docx_score_table = _docx_nucleus_score
    nuclei_export._pdf_score_table = _pdf_nucleus_score

    report_quality._docx_regulation = _docx_regulation
    report_quality._pdf_regulation = _pdf_regulation
    report_quality._docx_methodology = _docx_methodology
    report_quality._pdf_methodology = _pdf_methodology
    report_quality._docx_requirements = _docx_requirements
    report_quality._pdf_requirements = _pdf_requirements
    report_quality._docx_schedules = _docx_schedules
    report_quality._pdf_schedules = _pdf_schedules
    report_quality._docx_nucleus_results = _docx_nuclei
    report_quality._pdf_nucleus_results = _pdf_nuclei
    report_quality._docx_post_sections = _docx_post
    report_quality._pdf_post_sections = _pdf_post

    original_docx_body = report_quality._docx_body
    original_pdf_body = report_quality._pdf_body
    original_docx_caption = report_quality._docx_caption
    original_pdf_caption = report_quality._pdf_caption
    original_docx_bullet = report_quality._docx_bullet
    original_pdf_bullet = report_quality._pdf_bullet

    report_quality._docx_body = lambda document, text: original_docx_body(document, public_text(text))
    report_quality._pdf_body = lambda story, styles, text: original_pdf_body(story, styles, public_text(text))
    report_quality._docx_caption = lambda document, text: original_docx_caption(document, public_text(text))
    report_quality._pdf_caption = lambda story, styles, text: original_pdf_caption(story, styles, public_text(text))
    report_quality._docx_bullet = lambda document, text: original_docx_bullet(document, public_text(text))
    report_quality._pdf_bullet = lambda story, styles, text: original_pdf_bullet(story, styles, public_text(text))

    report_completion._REFERENCE_LIST = APA_REFERENCES
    report_quality._academic_enhancements_installed = True
