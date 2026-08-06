from __future__ import annotations

import html
from pathlib import Path
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from reportlab.lib.units import cm
from reportlab.platypus import Image, Paragraph, Spacer

import institutional_export as base
import report_structure
from nuclei_catalog import catalogs_for_report, create_cycle_diagram


def _diagram_path(report_id: int, career_name: str) -> Path:
    base.EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    safe_name = "".join(character.lower() if character.isalnum() else "_" for character in career_name)
    return base.EXPORT_DIR / f"nuclei_catalog_{report_id}_{safe_name}.png"


def add_docx_catalogs(document: Any, report: dict[str, Any]) -> None:
    catalogs = catalogs_for_report(report)
    if not catalogs:
        return

    document.add_heading("CONTENIDO DE LOS NÚCLEOS", level=1)
    for catalog in catalogs:
        document.add_heading(catalog["career"], level=2)
        document.add_paragraph(
            "La carrera organiza su preparación académica en cuatro núcleos estructurantes. "
            "Cada núcleo corresponde a una guía que integra las asignaturas relacionadas a continuación."
        )

        diagram = _diagram_path(int(report["id"]), catalog["career"])
        create_cycle_diagram(catalog, diagram)
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(diagram), width=Inches(6.45))

        for nucleus in catalog.get("nuclei", []):
            document.add_heading(
                f"Núcleo {nucleus['number']}: {nucleus['guide']}", level=3
            )
            document.add_paragraph(
                "Esta guía articula las siguientes asignaturas:"
            )
            for subject in nucleus.get("subjects", []):
                document.add_paragraph(subject, style="List Bullet")


def add_pdf_catalogs(story: list[Any], report: dict[str, Any], styles: Any) -> None:
    catalogs = catalogs_for_report(report)
    if not catalogs:
        return

    story.append(Paragraph("CONTENIDO DE LOS NÚCLEOS", styles["Heading1"]))
    for catalog in catalogs:
        story.append(Paragraph(html.escape(catalog["career"]), styles["Heading2"]))
        story.append(
            Paragraph(
                "La carrera organiza su preparación académica en cuatro núcleos estructurantes. "
                "Cada núcleo corresponde a una guía que integra las asignaturas relacionadas a continuación.",
                styles["BodyJustified"],
            )
        )
        diagram = _diagram_path(int(report["id"]), catalog["career"])
        create_cycle_diagram(catalog, diagram)
        image = Image(str(diagram))
        image._restrictSize(16.2 * cm, 11.0 * cm)
        story += [image, Spacer(1, 0.25 * cm)]

        for nucleus in catalog.get("nuclei", []):
            story.append(
                Paragraph(
                    f"Núcleo {nucleus['number']}: {html.escape(nucleus['guide'])}",
                    styles["Heading3"],
                )
            )
            story.append(
                Paragraph(
                    "Esta guía articula las siguientes asignaturas:",
                    styles["BodyJustified"],
                )
            )
            for subject in nucleus.get("subjects", []):
                story.append(
                    Paragraph(f"• {html.escape(subject)}", styles["BodyJustified"])
                )


def install() -> None:
    if getattr(report_structure, "_nuclei_catalog_export_installed", False):
        return

    original_docx_sections = report_structure._add_docx_sections
    original_pdf_sections = report_structure._add_pdf_sections

    def docx_sections_with_catalog(
        document: Any,
        report: dict[str, Any],
        post: bool,
    ) -> None:
        original_docx_sections(document, report, post)
        if not post:
            add_docx_catalogs(document, report)

    def pdf_sections_with_catalog(
        story: list[Any],
        report: dict[str, Any],
        styles: Any,
        post: bool,
    ) -> None:
        original_pdf_sections(story, report, styles, post)
        if not post:
            add_pdf_catalogs(story, report, styles)

    report_structure._add_docx_sections = docx_sections_with_catalog
    report_structure._add_pdf_sections = pdf_sections_with_catalog
    report_structure._nuclei_catalog_export_installed = True
