from __future__ import annotations

import html
from collections import defaultdict
from statistics import mean
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, Spacer, Table, TableStyle

import app as core
import report_quality
import report_structure
from analytics import summary
from completion_service import get_completion_data, get_schedules_extended
from eligibility_service import get_eligibility
from parser import canonical_name_key, clean_moodle_name
from process_service import get_projects
from roster_service import REQUIREMENTS, get_report_roster


CURRENT_REPORT_ID: int | None = None

REQUIREMENT_DEFINITIONS = {
    "academic_status": "Verifica que el estudiante haya aprobado las asignaturas y créditos establecidos en su plan de estudios.",
    "documentation_status": "Confirma la entrega y validez de la documentación requerida para continuar el proceso.",
    "financial_status": "Comprueba que el estudiante no mantenga obligaciones financieras que impidan la titulación.",
    "titulation_status": "Registra el estado general del estudiante dentro del proceso de titulación.",
    "practices_linkage_status": "Verifica conjuntamente el cumplimiento de prácticas preprofesionales y vinculación con la sociedad.",
    "linkage_status": "Confirma el cumplimiento específico de las actividades de vinculación con la sociedad.",
    "graduate_followup_status": "Verifica el registro solicitado por el sistema institucional de seguimiento a graduados.",
    "english_status": "Confirma el cumplimiento del requisito institucional de suficiencia o formación en inglés.",
    "data_update_status": "Comprueba que los datos personales y académicos del estudiante se encuentren actualizados.",
    "titulation_approval": "Registra la aprobación administrativa y académica para continuar el proceso de titulación.",
    "complexive_approval": "Identifica la habilitación institucional para la opción de Examen Complexivo o Trabajo de Titulación.",
}


def _normalize(value: Any) -> str:
    return " ".join(str(value or "").strip().split())


def _dedupe_roster(students: list[dict[str, Any]]) -> list[dict[str, Any]]:
    groups: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for student in students:
        identification = _normalize(student.get("identification"))
        email = _normalize(student.get("email")).casefold()
        name = canonical_name_key(clean_moodle_name(str(student.get("full_name") or "")))
        career = canonical_name_key(str(student.get("career_name") or ""))
        key = (
            f"id:{identification}"
            if identification
            else f"email:{email}"
            if email
            else f"name:{career}|{name}"
        )
        groups[key].append(dict(student))

    result: list[dict[str, Any]] = []
    requirement_keys = [key for key, _ in REQUIREMENTS]
    for items in groups.values():
        selected = max(
            items,
            key=lambda item: sum(bool(_normalize(item.get(key))) for key in requirement_keys),
        )
        merged = dict(selected)
        for key in requirement_keys:
            values = [_normalize(item.get(key)).upper() for item in items]
            if "NO CUMPLE" in values:
                merged[key] = "NO CUMPLE"
            elif "CUMPLE" in values:
                merged[key] = "CUMPLE"
            else:
                merged[key] = next((value for value in values if value), "")
        result.append(merged)
    return result


def corrected_requirement_analysis(report_id: int) -> dict[str, Any] | None:
    students = _dedupe_roster(get_report_roster(report_id).get("students", []))
    if not students:
        return None

    active = [
        (key, label)
        for key, label in REQUIREMENTS
        if any(_normalize(student.get(key)) for student in students)
    ]
    if not active:
        return None

    def classify(student: dict[str, Any]) -> str:
        values = [_normalize(student.get(key)).upper() for key, _ in active]
        if any(value == "NO CUMPLE" for value in values):
            return "pending"
        if any(not value for value in values):
            return "incomplete"
        return "complete" if all(value == "CUMPLE" for value in values) else "incomplete"

    total = len(students)
    states = [classify(student) for student in students]
    complete = states.count("complete")
    pending = states.count("pending")
    incomplete = states.count("incomplete")

    requirement_rows = []
    for key, label in active:
        values = [_normalize(student.get(key)).upper() for student in students]
        complies = values.count("CUMPLE")
        requirement_rows.append(
            {
                "key": key,
                "label": label,
                "complies": complies,
                "does_not_comply": values.count("NO CUMPLE"),
                "blank": values.count(""),
                "percentage": round(complies / total * 100, 2),
            }
        )

    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for student in students:
        grouped[str(student.get("career_name") or "Sin carrera")].append(student)
    career_rows = []
    for career_name, career_students in sorted(grouped.items()):
        career_states = [classify(student) for student in career_students]
        registered = len(career_students)
        career_complete = career_states.count("complete")
        career_rows.append(
            {
                "career": career_name,
                "registered": registered,
                "complete": career_complete,
                "pending": career_states.count("pending"),
                "incomplete": career_states.count("incomplete"),
                "percentage": round(career_complete / registered * 100, 2),
            }
        )

    lowest_percentage = min(row["percentage"] for row in requirement_rows)
    lowest_requirements = [row for row in requirement_rows if row["percentage"] == lowest_percentage]
    highest_issue_count = max(row["pending"] + row["incomplete"] for row in career_rows)
    highest_issue_careers = [
        row for row in career_rows
        if row["pending"] + row["incomplete"] == highest_issue_count
    ]
    requirement_names = ", ".join(row["label"] for row in lowest_requirements)
    career_names = " y ".join(row["career"] for row in highest_issue_careers)
    narrative = (
        f"De los {total} estudiantes únicos registrados, {complete} cumplieron integralmente los requisitos, "
        f"equivalente al {report_quality._pct(round(complete / total * 100, 2))}. Se identificaron {pending} "
        f"estudiantes con al menos un requisito marcado como NO CUMPLE y {incomplete} con información incompleta. "
        f"El menor nivel de cumplimiento se registró en {requirement_names}, con {report_quality._pct(lowest_percentage)}. "
        f"La mayor cantidad de casos pendientes o incompletos se presentó en {career_names}, con {highest_issue_count} casos por carrera."
    )
    return {
        "total": total,
        "complete": complete,
        "pending": pending,
        "incomplete": incomplete,
        "percentage": round(complete / total * 100, 2),
        "requirements": requirement_rows,
        "careers": career_rows,
        "narrative": narrative,
    }


def _complexive_data(report: dict[str, Any]) -> dict[str, Any]:
    rows = []
    totals = {
        "registered": 0,
        "ordinary_approved": 0,
        "supplementary": 0,
        "recovered": 0,
        "final_approved": 0,
        "final_failed": 0,
        "not_evaluated": 0,
    }
    for career in report.get("careers", []):
        if not report_quality._has_notes(career):
            continue
        ordinary = summary(career["students"], "ordinario")
        supplementary = summary(career["students"], "supletorio")
        final = summary(career["students"], "consolidado")
        recovered = sum(
            student["supplementary_participant"]
            and student["ordinary_status"] != "Aprobado"
            and student["final_status"] == "Aprobado"
            for student in final["rows"]
        )
        row = {
            "career": career["name"],
            "registered": final["total"],
            "ordinary_approved": ordinary["approved"],
            "supplementary": supplementary["total"],
            "recovered": recovered,
            "final_approved": final["approved"],
            "final_failed": final["failed"],
            "not_evaluated": final["not_evaluated"],
            "approval_percentage": final["approved_pct"],
        }
        rows.append(row)
        for key in totals:
            totals[key] += int(row[key])
    totals["approval_percentage"] = round(
        totals["final_approved"] / totals["registered"] * 100, 2
    ) if totals["registered"] else 0.0
    totals["failure_percentage"] = round(
        totals["final_failed"] / totals["registered"] * 100, 2
    ) if totals["registered"] else 0.0
    totals["not_evaluated_percentage"] = round(
        totals["not_evaluated"] / totals["registered"] * 100, 2
    ) if totals["registered"] else 0.0
    return {"rows": rows, "totals": totals}


def _schedule_data(report_id: int) -> dict[str, Any]:
    schedules = get_schedules_extended(report_id)
    all_rows = schedules.get("complexive", []) + schedules.get("thesis", [])
    evaluated = [row for row in all_rows if row.get("execution_status") or row.get("compliance_percentage") is not None]
    percentages = [float(row["compliance_percentage"]) for row in evaluated if row.get("compliance_percentage") is not None]
    return {
        "schedules": schedules,
        "total": len(all_rows),
        "evaluated": len(evaluated),
        "average_compliance": round(mean(percentages), 2) if percentages else None,
        "not_complied": sum(row.get("execution_status") == "No cumplido" for row in evaluated),
        "delayed": sum(row.get("execution_status") == "Cumplido con retraso" for row in evaluated),
        "partial": sum(row.get("execution_status") == "Cumplido parcialmente" for row in evaluated),
    }


def _executive_data(report_id: int) -> dict[str, Any]:
    report = report_quality._report_data(report_id)
    requirements = corrected_requirement_analysis(report_id)
    eligibility = get_eligibility(report_id)
    complexive = _complexive_data(report)
    projects = get_projects(report_id)
    schedules = _schedule_data(report_id)
    completion = get_completion_data(report_id)
    totals = complexive["totals"]
    evaluated = totals["registered"] - totals["not_evaluated"] + projects["summary"]["total"]
    return {
        "report": report,
        "requirements": requirements,
        "eligibility": eligibility,
        "complexive": complexive,
        "projects": projects,
        "schedules": schedules,
        "completion": completion,
        "indicators": [
            ("Estudiantes registrados", requirements["total"] if requirements else eligibility["summary"]["registered"]),
            ("Cumplieron requisitos", requirements["complete"] if requirements else 0),
            ("Habilitados por los cuatro núcleos", eligibility["summary"]["habilitated"]),
            ("Total evaluado", evaluated),
            ("Aprobados en ordinario", totals["ordinary_approved"]),
            ("Enviados a supletorio", totals["supplementary"]),
            ("Recuperados mediante supletorio", totals["recovered"]),
            ("Aprobados finales en Complexivo", totals["final_approved"]),
            ("Reprobados finales en Complexivo", totals["final_failed"]),
            ("No evaluados en Complexivo", totals["not_evaluated"]),
            ("Estudiantes en Trabajo de Titulación", projects["summary"]["total"]),
            ("Aprobados en Trabajo de Titulación", projects["summary"]["approved"]),
        ],
    }


def _automatic_incidents(data: dict[str, Any]) -> list[dict[str, str]]:
    incidents: list[dict[str, str]] = []
    requirements = data["requirements"]
    eligibility = data["eligibility"]["summary"]
    complexive = data["complexive"]["totals"]
    schedules = data["schedules"]
    duplicate_count = sum(len(items) for items in data["report"].get("duplicate_warnings", {}).values())

    if requirements and (requirements["pending"] or requirements["incomplete"]):
        incidents.append({
            "category": "Requisitos",
            "description": f"Se registraron {requirements['pending']} estudiantes con requisitos pendientes y {requirements['incomplete']} con información incompleta.",
            "responsible": "Áreas responsables de los campos de requisitos",
            "treatment": "Completar y validar los registros antes de habilitar la evaluación.",
            "status": "En seguimiento",
            "evidence": "Matriz de cumplimiento de requisitos",
        })
    if eligibility["not_habilitated"] or eligibility["pending"]:
        incidents.append({
            "category": "Núcleos",
            "description": f"Se identificaron {eligibility['not_habilitated']} estudiantes no habilitados y {eligibility['pending']} pendientes de completar los cuatro núcleos.",
            "responsible": "Coordinaciones de carrera y Unidad de Titulación",
            "treatment": "Verificar notas, ejecutar recuperación cuando corresponda y habilitar únicamente a quienes aprueben los cuatro núcleos.",
            "status": "En seguimiento",
            "evidence": "Matriz de habilitación por núcleos",
        })
    if eligibility["unmatched_nucleus_records"]:
        incidents.append({
            "category": "Correspondencia de datos",
            "description": f"Existen {eligibility['unmatched_nucleus_records']} registros de calificaciones de núcleos sin coincidencia en la base de estudiantes.",
            "responsible": "Coordinación de Titulación",
            "treatment": "Validar correo, cédula, nombre y carrera antes de consolidar el resultado.",
            "status": "Abierto",
            "evidence": "Listado de registros sin coincidencia",
        })
    if complexive["not_evaluated"]:
        incidents.append({
            "category": "Evaluación",
            "description": f"Se registraron {complexive['not_evaluated']} estudiantes sin evaluación completa en el Examen Complexivo.",
            "responsible": "Coordinaciones de carrera",
            "treatment": "Determinar la causa de la ausencia o falta de nota y documentar el estado final del caso.",
            "status": "En seguimiento",
            "evidence": "Consolidado del Examen Complexivo",
        })
    if duplicate_count:
        incidents.append({
            "category": "Calidad de datos",
            "description": f"Se consolidaron {duplicate_count} posibles registros duplicados por variaciones en el orden del nombre.",
            "responsible": "Coordinación de Titulación y Secretaría",
            "treatment": "Verificar la identidad con cédula y correo institucional y corregir la fuente original.",
            "status": "En seguimiento",
            "evidence": "Control de duplicados de Informtit",
        })
    if schedules["total"] and schedules["evaluated"] < schedules["total"]:
        incidents.append({
            "category": "Cronograma",
            "description": f"La ejecución fue evaluada en {schedules['evaluated']} de {schedules['total']} actividades planificadas.",
            "responsible": "Responsables de cada actividad",
            "treatment": "Completar fecha ejecutada, estado, porcentaje, evidencia y observación.",
            "status": "Abierto",
            "evidence": "Matriz de evaluación de cronogramas",
        })
    return incidents


def _automatic_actions(data: dict[str, Any]) -> list[dict[str, str]]:
    actions: list[dict[str, str]] = []
    for incident in _automatic_incidents(data):
        actions.append({
            "finding": incident["description"],
            "action": incident["treatment"],
            "responsible": incident["responsible"],
            "due_date": "",
            "indicator": "Porcentaje de casos corregidos o actividades completadas",
            "evidence": incident["evidence"],
            "status": "Pendiente",
        })
    rows = data["complexive"]["rows"]
    if rows:
        lowest = min(row["approval_percentage"] for row in rows)
        critical = [row["career"] for row in rows if row["approval_percentage"] == lowest]
        if lowest < 80:
            actions.append({
                "finding": f"Las carreras con menor aprobación final fueron {' y '.join(critical)}, con {report_quality._pct(lowest)}.",
                "action": "Implementar refuerzo académico focalizado y seguimiento temprano de los componentes con mayor dificultad.",
                "responsible": "Coordinaciones de carrera y docentes responsables",
                "due_date": "",
                "indicator": "Incremento del porcentaje de aprobación en el siguiente período",
                "evidence": "Plan de refuerzo y reporte comparativo",
                "status": "Pendiente",
            })
    return actions


def _docx_simple_table(document: Any, rows: list[tuple[str, Any]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = report_quality._fmt(value)
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(9)
        for run in cells[1].paragraphs[0].runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)
    document.add_paragraph()


def _add_docx_executive_summary(document: Any, report_id: int) -> None:
    data = _executive_data(report_id)
    heading = document.add_heading("RESUMEN EJECUTIVO", level=1)
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.first_line_indent = Cm(0)
    _docx_simple_table(document, data["indicators"])
    issues = _automatic_incidents(data)
    if issues:
        report_quality._docx_body(
            document,
            "Los principales aspectos que requieren seguimiento son los siguientes:",
        )
        for item in issues[:4]:
            report_quality._docx_bullet(document, item["description"])


def _add_pdf_executive_summary(story: list[Any], styles: Any, report_id: int) -> None:
    data = _executive_data(report_id)
    story.append(Paragraph("RESUMEN EJECUTIVO", styles["Heading1"]))
    table = Table(
        [[Paragraph(html.escape(label), styles["TableCell"]), report_quality._fmt(value)] for label, value in data["indicators"]],
        colWidths=[10.5 * cm, 5.5 * cm],
    )
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story += [table, Spacer(1, 0.25 * cm)]
    issues = _automatic_incidents(data)
    if issues:
        report_quality._pdf_body(story, styles, "Los principales aspectos que requieren seguimiento son los siguientes:")
        for item in issues[:4]:
            report_quality._pdf_bullet(story, styles, item["description"])


def _add_docx_objectives(document: Any, context: Any, report: dict[str, Any]) -> None:
    report_quality._docx_heading(document, context, 1, "Objetivos")
    report_quality._docx_heading(document, context, 2, "Objetivo general")
    report_quality._docx_body(
        document,
        f"Evaluar el desarrollo y los resultados del proceso de titulación correspondiente al período académico {report.get('period') or 'analizado'}, modalidad {report_quality.base.modality(report)}, mediante el análisis del cumplimiento de requisitos, la ejecución de las actividades planificadas y los resultados alcanzados en las diferentes opciones de titulación.",
    )
    report_quality._docx_heading(document, context, 2, "Objetivos específicos")
    for item in (
        "Determinar el nivel de cumplimiento de los requisitos de titulación.",
        "Evaluar el cumplimiento de los cronogramas establecidos.",
        "Analizar la aprobación de los cuatro núcleos y la habilitación para el Examen Complexivo.",
        "Analizar los resultados del Examen Complexivo ordinario y supletorio.",
        "Analizar los resultados del Trabajo de Titulación cuando existan registros.",
        "Identificar dificultades, novedades y oportunidades de mejora.",
        "Proponer acciones verificables para optimizar los siguientes períodos.",
    ):
        report_quality._docx_bullet(document, item)


def _add_pdf_objectives(story: list[Any], context: Any, styles: Any, report: dict[str, Any]) -> None:
    report_quality._pdf_heading(story, context, styles, 1, "Objetivos")
    report_quality._pdf_heading(story, context, styles, 2, "Objetivo general")
    report_quality._pdf_body(
        story,
        styles,
        f"Evaluar el desarrollo y los resultados del proceso de titulación correspondiente al período académico {report.get('period') or 'analizado'}, modalidad {report_quality.base.modality(report)}, mediante el análisis del cumplimiento de requisitos, la ejecución de las actividades planificadas y los resultados alcanzados en las diferentes opciones de titulación.",
    )
    report_quality._pdf_heading(story, context, styles, 2, "Objetivos específicos")
    for item in (
        "Determinar el nivel de cumplimiento de los requisitos de titulación.",
        "Evaluar el cumplimiento de los cronogramas establecidos.",
        "Analizar la aprobación de los cuatro núcleos y la habilitación para el Examen Complexivo.",
        "Analizar los resultados del Examen Complexivo ordinario y supletorio.",
        "Analizar los resultados del Trabajo de Titulación cuando existan registros.",
        "Identificar dificultades, novedades y oportunidades de mejora.",
        "Proponer acciones verificables para optimizar los siguientes períodos.",
    ):
        report_quality._pdf_bullet(story, styles, item)


def _methodology_paragraphs(report_id: int, report: dict[str, Any]) -> list[str]:
    requirement_data = corrected_requirement_analysis(report_id)
    project_count = get_projects(report_id)["summary"]["total"]
    cutoff = report_quality.base.format_date(report.get("cutoff_date")) if report.get("cutoff_date") else "no registrada"
    return [
        f"La información fue procesada con fecha de corte {cutoff}. La población estuvo conformada por {requirement_data['total'] if requirement_data else 0} estudiantes únicos registrados en la base institucional del período {report.get('period') or 'analizado'}.",
        "Se incluyeron los estudiantes vinculados al informe activo y se excluyeron encabezados, totales generales y registros que no correspondían al rol de estudiante. Las fuentes utilizadas fueron la base institucional importada, las calificaciones copiadas desde Moodle, los cronogramas, las actas y rúbricas del Trabajo de Titulación y los documentos cargados en Informtit, según la información efectivamente disponible.",
        "La depuración priorizó la cédula y el correo institucional; cuando estos datos no estuvieron disponibles, se comparó el nombre normalizado dentro de la misma carrera. Los posibles duplicados se consolidaron conservando el registro con mayor información y completando los campos faltantes sin reemplazar una calificación válida.",
        "Los porcentajes se calcularon dividiendo la cantidad de casos de cada estado para el total de registros aplicables y multiplicando el resultado por cien. Los estudiantes no evaluados se contabilizaron en su categoría, pero se excluyeron del cálculo de promedios de calificaciones.",
        "Para el Examen Complexivo, la nota supletoria reemplazó únicamente el componente rendido. La aprobación exigió una calificación mínima de 70/100 tanto en el componente teórico como en el práctico y una calificación ponderada final igual o superior a 70/100.",
        "Para la habilitación al Examen Complexivo, cada estudiante debía registrar y aprobar los cuatro núcleos con una calificación mínima de 7,00. Una nota inferior o la ausencia de calificación en cualquiera de los núcleos impidió confirmar la habilitación.",
        f"El Trabajo de Titulación se analizó de forma independiente con los {project_count} registros incorporados. Cuando un dato de seguimiento, evidencia o ejecución no fue cargado, no se asumió su cumplimiento y se identificó como información pendiente.",
        f"Aunque el informe corresponde a los estudiantes del período académico {report.get('period') or 'analizado'}, las actividades de evaluación, defensa, supletorio y cierre se desarrollaron conforme a las fechas del cronograma institucional, que pueden extenderse más allá del cierre de clases del período.",
    ]


def _add_docx_report_methodology(document: Any, context: Any, report_id: int, report: dict[str, Any]) -> None:
    report_quality._docx_heading(document, context, 1, "Alcance y metodología de elaboración del informe")
    for paragraph in _methodology_paragraphs(report_id, report):
        report_quality._docx_body(document, paragraph)


def _add_pdf_report_methodology(story: list[Any], context: Any, styles: Any, report_id: int, report: dict[str, Any]) -> None:
    report_quality._pdf_heading(story, context, styles, 1, "Alcance y metodología de elaboración del informe")
    for paragraph in _methodology_paragraphs(report_id, report):
        report_quality._pdf_body(story, styles, paragraph)


def _add_docx_requirement_definitions(document: Any, context: Any, report_id: int) -> None:
    data = corrected_requirement_analysis(report_id)
    if not data:
        return
    report_quality._docx_heading(document, context, 2, "Definición de los requisitos")
    rows = [
        [row["label"], REQUIREMENT_DEFINITIONS.get(row["key"], "Campo definido en la matriz institucional."), "Área institucional que administra el campo"]
        for row in data["requirements"]
    ]
    report_quality._docx_caption(document, context.table_caption("Definición y responsabilidad de los requisitos"))
    report_quality._docx_table(document, ["Requisito", "Definición", "Responsable"], rows, [1.65, 3.55, 1.35])


def _add_pdf_requirement_definitions(story: list[Any], context: Any, styles: Any, report_id: int) -> None:
    data = corrected_requirement_analysis(report_id)
    if not data:
        return
    report_quality._pdf_heading(story, context, styles, 2, "Definición de los requisitos")
    rows = [
        [Paragraph(html.escape(row["label"]), styles["TableCell"]), Paragraph(html.escape(REQUIREMENT_DEFINITIONS.get(row["key"], "Campo definido en la matriz institucional.")), styles["TableCell"]), Paragraph("Área institucional que administra el campo", styles["TableCell"])]
        for row in data["requirements"]
    ]
    report_quality._pdf_caption(story, styles, context.table_caption("Definición y responsabilidad de los requisitos"))
    story += [report_quality._pdf_table(["Requisito", "Definición", "Responsable"], rows, [4.1 * cm, 8.3 * cm, 4.1 * cm]), Spacer(1, 0.2 * cm)]


def _schedule_rows(rows: list[dict[str, Any]], show_phase: bool) -> tuple[list[str], list[list[Any]]]:
    headers = (["Fase"] if show_phase else []) + ["Actividad", "Fecha planificada", "Fecha ejecutada", "Estado", "% cumplimiento", "Evidencia", "Observación"]
    values = []
    for row in rows:
        current = []
        if show_phase:
            current.append(row.get("phase") or "—")
        planned = row.get("start_date") or "—"
        if row.get("end_date") and row.get("end_date") != row.get("start_date"):
            planned += f" a {row['end_date']}"
        current.extend([
            row.get("activity") or "—",
            planned,
            row.get("executed_date") or "—",
            row.get("execution_status") or "Sin evaluar",
            report_quality._pct(row["compliance_percentage"]) if row.get("compliance_percentage") is not None else "—",
            row.get("evidence") or "—",
            row.get("observation") or "—",
        ])
        values.append(current)
    return headers, values


def _add_docx_schedules(document: Any, context: Any, report_id: int) -> None:
    data = _schedule_data(report_id)
    schedules = data["schedules"]
    available = [("Cronograma de Núcleos y Examen Complexivo", schedules.get("complexive", []), False), ("Cronograma del Trabajo de Titulación", schedules.get("thesis", []), True)]
    available = [item for item in available if item[1]]
    if not available:
        return
    report_quality._docx_heading(document, context, 1, "Evaluación del cumplimiento de los cronogramas")
    for title, rows, show_phase in available:
        report_quality._docx_heading(document, context, 2, title)
        headers, values = _schedule_rows(rows, show_phase)
        report_quality._docx_caption(document, context.table_caption(f"Planificación y ejecución: {title}"))
        widths = ([0.8] if show_phase else []) + [1.3, 1.05, 0.85, 0.85, 0.75, 0.95, 1.05]
        report_quality._docx_table(document, headers, values, widths)
    if data["evaluated"]:
        report_quality._docx_body(document, f"Se evaluó la ejecución de {data['evaluated']} de {data['total']} actividades. El cumplimiento promedio registrado fue {report_quality._pct(data['average_compliance'])}. Se identificaron {data['delayed']} actividades cumplidas con retraso, {data['partial']} cumplidas parcialmente y {data['not_complied']} no cumplidas.")
    else:
        report_quality._docx_body(document, "Los cronogramas contienen las fechas planificadas, pero todavía no se ha registrado la información de ejecución necesaria para evaluar su cumplimiento.")


def _add_pdf_schedules(story: list[Any], context: Any, styles: Any, report_id: int) -> None:
    data = _schedule_data(report_id)
    schedules = data["schedules"]
    available = [("Cronograma de Núcleos y Examen Complexivo", schedules.get("complexive", []), False), ("Cronograma del Trabajo de Titulación", schedules.get("thesis", []), True)]
    available = [item for item in available if item[1]]
    if not available:
        return
    report_quality._pdf_heading(story, context, styles, 1, "Evaluación del cumplimiento de los cronogramas")
    for title, rows, show_phase in available:
        report_quality._pdf_heading(story, context, styles, 2, title)
        headers, values = _schedule_rows(rows, show_phase)
        pdf_rows = [[Paragraph(html.escape(str(value)), styles["TableCell"]) for value in row] for row in values]
        report_quality._pdf_caption(story, styles, context.table_caption(f"Planificación y ejecución: {title}"))
        widths = ([2.0 * cm] if show_phase else []) + [3.2 * cm, 2.3 * cm, 2.0 * cm, 2.2 * cm, 1.8 * cm, 2.4 * cm, 2.6 * cm]
        story += [report_quality._pdf_table(headers, pdf_rows, widths), Spacer(1, 0.2 * cm)]
    if data["evaluated"]:
        report_quality._pdf_body(story, styles, f"Se evaluó la ejecución de {data['evaluated']} de {data['total']} actividades. El cumplimiento promedio registrado fue {report_quality._pct(data['average_compliance'])}. Se identificaron {data['delayed']} actividades cumplidas con retraso, {data['partial']} cumplidas parcialmente y {data['not_complied']} no cumplidas.")
    else:
        report_quality._pdf_body(story, styles, "Los cronogramas contienen las fechas planificadas, pero todavía no se ha registrado la información de ejecución necesaria para evaluar su cumplimiento.")


def _add_docx_eligibility(document: Any, context: Any, report_id: int) -> None:
    eligibility = get_eligibility(report_id)
    if not any(getattr(row, "get", lambda *_: None)("nucleus_1") is not None for row in eligibility["rows"]) and not eligibility["unmatched"]:
        return
    report_quality._docx_heading(document, context, 2, "Habilitación para el Examen Complexivo")
    summary_data = eligibility["summary"]
    report_quality._docx_body(document, f"De {summary_data['complexive_candidates']} candidatos al Examen Complexivo, {summary_data['habilitated']} aprobaron los cuatro núcleos y quedaron habilitados, {summary_data['not_habilitated']} registraron uno o más núcleos reprobados y {summary_data['pending']} mantuvieron calificaciones pendientes. El porcentaje de habilitación fue {report_quality._pct(summary_data['habilitation_percentage'])}.")
    report_quality._docx_caption(document, context.table_caption("Habilitación por carrera mediante los cuatro núcleos"))
    report_quality._docx_table(document, ["Carrera", "Candidatos", "Habilitados", "No habilitados", "Pendientes", "% habilitación"], [[row["career_name"], row["total"], row["habilitated"], row["not_habilitated"], row["pending"], report_quality._pct(row["habilitation_percentage"])] for row in eligibility["careers"]], [2.4, 0.75, 0.8, 0.9, 0.75, 0.9])
    for career in eligibility["careers"]:
        rows = [row for row in eligibility["rows"] if row["career_name"] == career["career_name"] and row["option"] == "Examen Complexivo"]
        if not rows:
            continue
        report_quality._docx_heading(document, context, 3, career["career_name"])
        report_quality._docx_table(document, ["Estudiante", "Núcleo 1", "Núcleo 2", "Núcleo 3", "Núcleo 4", "Estado"], [[row["full_name"], report_quality._fmt(row["nucleus_1"]), report_quality._fmt(row["nucleus_2"]), report_quality._fmt(row["nucleus_3"]), report_quality._fmt(row["nucleus_4"]), row["status"]] for row in rows], [2.8, 0.65, 0.65, 0.65, 0.65, 1.0])


def _add_pdf_eligibility(story: list[Any], context: Any, styles: Any, report_id: int) -> None:
    eligibility = get_eligibility(report_id)
    if not any(row.get("nucleus_1") is not None for row in eligibility["rows"]) and not eligibility["unmatched"]:
        return
    report_quality._pdf_heading(story, context, styles, 2, "Habilitación para el Examen Complexivo")
    summary_data = eligibility["summary"]
    report_quality._pdf_body(story, styles, f"De {summary_data['complexive_candidates']} candidatos al Examen Complexivo, {summary_data['habilitated']} aprobaron los cuatro núcleos y quedaron habilitados, {summary_data['not_habilitated']} registraron uno o más núcleos reprobados y {summary_data['pending']} mantuvieron calificaciones pendientes. El porcentaje de habilitación fue {report_quality._pct(summary_data['habilitation_percentage'])}.")
    report_quality._pdf_caption(story, styles, context.table_caption("Habilitación por carrera mediante los cuatro núcleos"))
    rows = [[Paragraph(html.escape(row["career_name"]), styles["TableCell"]), row["total"], row["habilitated"], row["not_habilitated"], row["pending"], report_quality._pct(row["habilitation_percentage"])] for row in eligibility["careers"]]
    story += [report_quality._pdf_table(["Carrera", "Candidatos", "Habilitados", "No habilitados", "Pendientes", "% habilitación"], rows, [5.8 * cm, 2.0 * cm, 2.1 * cm, 2.3 * cm, 2.0 * cm, 2.4 * cm]), Spacer(1, 0.2 * cm)]
    for career in eligibility["careers"]:
        career_rows = [row for row in eligibility["rows"] if row["career_name"] == career["career_name"] and row["option"] == "Examen Complexivo"]
        if not career_rows:
            continue
        report_quality._pdf_heading(story, context, styles, 3, career["career_name"])
        values = [[Paragraph(html.escape(row["full_name"]), styles["TableCell"]), report_quality._fmt(row["nucleus_1"]), report_quality._fmt(row["nucleus_2"]), report_quality._fmt(row["nucleus_3"]), report_quality._fmt(row["nucleus_4"]), row["status"]] for row in career_rows]
        story += [report_quality._pdf_table(["Estudiante", "Núcleo 1", "Núcleo 2", "Núcleo 3", "Núcleo 4", "Estado"], values, [7.2 * cm, 1.7 * cm, 1.7 * cm, 1.7 * cm, 1.7 * cm, 3.0 * cm]), Spacer(1, 0.2 * cm)]


def _add_docx_complexive_consolidated(document: Any, context: Any, report: dict[str, Any]) -> None:
    data = _complexive_data(report)
    if not data["rows"]:
        return
    report_quality._docx_heading(document, context, 2, "Consolidado general del Examen Complexivo")
    report_quality._docx_caption(document, context.table_caption("Consolidado institucional del Examen Complexivo"))
    headers = ["Carrera", "Registrados", "Aprobados ordinarios", "Supletorios", "Recuperados", "Aprobados finales", "Reprobados finales", "No evaluados", "% aprobación final"]
    rows = [[row["career"], row["registered"], row["ordinary_approved"], row["supplementary"], row["recovered"], row["final_approved"], row["final_failed"], row["not_evaluated"], report_quality._pct(row["approval_percentage"])] for row in data["rows"]]
    report_quality._docx_table(document, headers, rows, [1.7, 0.55, 0.75, 0.6, 0.6, 0.7, 0.7, 0.65, 0.8])
    percentages = [row["approval_percentage"] for row in data["rows"]]
    highest = max(percentages)
    lowest = min(percentages)
    highest_names = " y ".join(row["career"] for row in data["rows"] if row["approval_percentage"] == highest)
    lowest_names = " y ".join(row["career"] for row in data["rows"] if row["approval_percentage"] == lowest)
    totals = data["totals"]
    report_quality._docx_body(document, f"El porcentaje global de aprobación final fue {report_quality._pct(totals['approval_percentage'])}; la reprobación representó {report_quality._pct(totals['failure_percentage'])} y los no evaluados {report_quality._pct(totals['not_evaluated_percentage'])}. La mayor aprobación se registró en {highest_names}, con {report_quality._pct(highest)}, mientras que la menor correspondió a {lowest_names}, con {report_quality._pct(lowest)}. El supletorio permitió recuperar a {totals['recovered']} estudiantes.")


def _add_pdf_complexive_consolidated(story: list[Any], context: Any, styles: Any, report: dict[str, Any]) -> None:
    data = _complexive_data(report)
    if not data["rows"]:
        return
    report_quality._pdf_heading(story, context, styles, 2, "Consolidado general del Examen Complexivo")
    report_quality._pdf_caption(story, styles, context.table_caption("Consolidado institucional del Examen Complexivo"))
    headers = ["Carrera", "Reg.", "Ord.", "Sup.", "Recup.", "Aprob.", "Reprob.", "No eval.", "% final"]
    rows = [[Paragraph(html.escape(row["career"]), styles["TableCell"]), row["registered"], row["ordinary_approved"], row["supplementary"], row["recovered"], row["final_approved"], row["final_failed"], row["not_evaluated"], report_quality._pct(row["approval_percentage"])] for row in data["rows"]]
    story += [report_quality._pdf_table(headers, rows, [4.5 * cm, 1.3 * cm, 1.3 * cm, 1.3 * cm, 1.4 * cm, 1.4 * cm, 1.4 * cm, 1.5 * cm, 2.2 * cm]), Spacer(1, 0.2 * cm)]
    percentages = [row["approval_percentage"] for row in data["rows"]]
    highest = max(percentages)
    lowest = min(percentages)
    highest_names = " y ".join(row["career"] for row in data["rows"] if row["approval_percentage"] == highest)
    lowest_names = " y ".join(row["career"] for row in data["rows"] if row["approval_percentage"] == lowest)
    totals = data["totals"]
    report_quality._pdf_body(story, styles, f"El porcentaje global de aprobación final fue {report_quality._pct(totals['approval_percentage'])}; la reprobación representó {report_quality._pct(totals['failure_percentage'])} y los no evaluados {report_quality._pct(totals['not_evaluated_percentage'])}. La mayor aprobación se registró en {highest_names}, con {report_quality._pct(highest)}, mientras que la menor correspondió a {lowest_names}, con {report_quality._pct(lowest)}. El supletorio permitió recuperar a {totals['recovered']} estudiantes.")


def _add_docx_global_process(document: Any, context: Any, report_id: int, report: dict[str, Any]) -> None:
    complexive = _complexive_data(report)["totals"]
    projects = get_projects(report_id)["summary"]
    if not complexive["registered"] and not projects["total"]:
        return
    report_quality._docx_heading(document, context, 1, "Consolidado general del proceso")
    rows = [
        ["Examen Complexivo", complexive["registered"], complexive["registered"] - complexive["not_evaluated"], complexive["final_approved"], complexive["final_failed"], complexive["not_evaluated"], report_quality._pct(complexive["approval_percentage"])],
        ["Trabajo de Titulación", projects["total"], projects["total"], projects["approved"], projects["failed"], 0, report_quality._pct(round(projects["approved"] / projects["total"] * 100, 2) if projects["total"] else 0)],
    ]
    report_quality._docx_caption(document, context.table_caption("Resultados consolidados por opción de titulación"))
    report_quality._docx_table(document, ["Opción", "Registrados", "Evaluados", "Aprobados", "Reprobados", "No evaluados", "% aprobación"], rows, [2.0, 0.8, 0.8, 0.8, 0.8, 0.85, 0.95])


def _add_pdf_global_process(story: list[Any], context: Any, styles: Any, report_id: int, report: dict[str, Any]) -> None:
    complexive = _complexive_data(report)["totals"]
    projects = get_projects(report_id)["summary"]
    if not complexive["registered"] and not projects["total"]:
        return
    report_quality._pdf_heading(story, context, styles, 1, "Consolidado general del proceso")
    rows = [
        ["Examen Complexivo", complexive["registered"], complexive["registered"] - complexive["not_evaluated"], complexive["final_approved"], complexive["final_failed"], complexive["not_evaluated"], report_quality._pct(complexive["approval_percentage"])],
        ["Trabajo de Titulación", projects["total"], projects["total"], projects["approved"], projects["failed"], 0, report_quality._pct(round(projects["approved"] / projects["total"] * 100, 2) if projects["total"] else 0)],
    ]
    report_quality._pdf_caption(story, styles, context.table_caption("Resultados consolidados por opción de titulación"))
    story += [report_quality._pdf_table(["Opción", "Registrados", "Evaluados", "Aprobados", "Reprobados", "No evaluados", "% aprobación"], rows, [4.5 * cm, 2.0 * cm, 2.0 * cm, 2.0 * cm, 2.0 * cm, 2.2 * cm, 2.3 * cm]), Spacer(1, 0.2 * cm)]


def _all_incidents(report_id: int) -> tuple[dict[str, Any], list[dict[str, Any]], list[dict[str, Any]]]:
    data = _executive_data(report_id)
    manual = data["completion"]["incidents"]
    incidents = _automatic_incidents(data) + manual
    actions = _automatic_actions(data) + data["completion"]["actions"]
    return data, incidents, actions


def _add_docx_post_sections(document: Any, context: Any, report: dict[str, Any]) -> None:
    report_id = int(report["id"])
    data, incidents, actions = _all_incidents(report_id)
    if incidents:
        report_quality._docx_heading(document, context, 1, "Novedades e incidencias del proceso")
        report_quality._docx_caption(document, context.table_caption("Novedades e incidencias registradas"))
        report_quality._docx_table(document, ["Categoría", "Descripción", "Responsable", "Tratamiento", "Estado", "Evidencia"], [[item.get("category") or "—", item.get("description") or "—", item.get("responsible") or "—", item.get("treatment") or "—", item.get("status") or "—", item.get("evidence") or "—"] for item in incidents], [0.8, 1.75, 1.2, 1.55, 0.75, 0.95])

    report_quality._docx_heading(document, context, 1, "Análisis comparativo y discusión de resultados")
    complexive = data["complexive"]
    totals = complexive["totals"]
    report_quality._docx_body(document, f"El Examen Complexivo registró una aprobación final global de {report_quality._pct(totals['approval_percentage'])}. Participaron {totals['supplementary']} estudiantes en supletorio y {totals['recovered']} lograron recuperar su condición académica mediante esta oportunidad. La interpretación por carrera debe considerar simultáneamente la aprobación de los cuatro núcleos, el cumplimiento de requisitos y la disponibilidad completa de las calificaciones.")
    if data["schedules"]["average_compliance"] is not None:
        report_quality._docx_body(document, f"El cumplimiento promedio de las actividades del cronograma que cuentan con evaluación fue {report_quality._pct(data['schedules']['average_compliance'])}. Las actividades no evaluadas requieren completar evidencia y observación antes del cierre documental.")

    report_quality._docx_heading(document, context, 1, "Conclusiones")
    conclusions = [
        f"El cumplimiento integral de requisitos alcanzó {report_quality._pct(data['requirements']['percentage'])} sobre {data['requirements']['total']} estudiantes únicos." if data["requirements"] else "No se dispuso de información suficiente para calcular el cumplimiento integral de requisitos.",
        f"La habilitación para el Examen Complexivo fue confirmada para {data['eligibility']['summary']['habilitated']} de {data['eligibility']['summary']['complexive_candidates']} candidatos, una vez verificada la aprobación individual de los cuatro núcleos.",
        f"El Examen Complexivo concluyó con {totals['final_approved']} aprobados, {totals['final_failed']} reprobados y {totals['not_evaluated']} no evaluados.",
        f"En Trabajo de Titulación se registraron {data['projects']['summary']['total']} estudiantes, de los cuales {data['projects']['summary']['approved']} aprobaron y {data['projects']['summary']['failed']} reprobaron.",
        f"La trazabilidad del proceso requiere mantener correspondencia entre la base institucional, las notas de núcleos, las evaluaciones del Complexivo, las actas y las evidencias de ejecución. Se identificaron {len(incidents)} novedades automáticas o registradas para seguimiento.",
    ]
    for conclusion in conclusions:
        report_quality._docx_bullet(document, conclusion)

    report_quality._docx_heading(document, context, 1, "Recomendaciones")
    for action in actions:
        report_quality._docx_bullet(document, action["action"])

    if actions:
        report_quality._docx_heading(document, context, 1, "Plan de mejora")
        report_quality._docx_caption(document, context.table_caption("Plan de mejora del proceso de titulación"))
        report_quality._docx_table(document, ["Hallazgo", "Acción de mejora", "Responsable", "Fecha límite", "Indicador", "Evidencia", "Estado"], [[item.get("finding") or "—", item.get("action") or "—", item.get("responsible") or "—", item.get("due_date") or "Por definir", item.get("indicator") or "—", item.get("evidence") or "—", item.get("status") or "Pendiente"] for item in actions], [1.45, 1.65, 1.1, 0.75, 1.1, 0.9, 0.75])

    report_quality._docx_heading(document, context, 1, "Referencias legales e institucionales")
    for reference in (
        "Constitución de la República del Ecuador.",
        "Ley Orgánica de Educación Superior.",
        "Reglamento de Régimen Académico, Resolución RPC-SE-08-No.023-2022.",
        "Reglamento de la Unidad de Titulación y Eficiencia Terminal, código UTET-REG-25, versión 2.0.",
        "Manual de Procesos de la Unidad de Titulación y Eficiencia Terminal, versión 2.",
        "Guías institucionales de integración curricular y registros académicos cargados en Informtit.",
    ):
        report_quality._docx_bullet(document, reference)


def _add_pdf_post_sections(story: list[Any], context: Any, styles: Any, report: dict[str, Any]) -> None:
    report_id = int(report["id"])
    data, incidents, actions = _all_incidents(report_id)
    if incidents:
        report_quality._pdf_heading(story, context, styles, 1, "Novedades e incidencias del proceso")
        rows = [[Paragraph(html.escape(str(item.get("category") or "—")), styles["TableCell"]), Paragraph(html.escape(str(item.get("description") or "—")), styles["TableCell"]), Paragraph(html.escape(str(item.get("responsible") or "—")), styles["TableCell"]), Paragraph(html.escape(str(item.get("treatment") or "—")), styles["TableCell"]), item.get("status") or "—", Paragraph(html.escape(str(item.get("evidence") or "—")), styles["TableCell"])] for item in incidents]
        report_quality._pdf_caption(story, styles, context.table_caption("Novedades e incidencias registradas"))
        story += [report_quality._pdf_table(["Categoría", "Descripción", "Responsable", "Tratamiento", "Estado", "Evidencia"], rows, [2.0 * cm, 4.2 * cm, 3.0 * cm, 4.0 * cm, 2.0 * cm, 2.6 * cm]), Spacer(1, 0.2 * cm)]

    report_quality._pdf_heading(story, context, styles, 1, "Análisis comparativo y discusión de resultados")
    totals = data["complexive"]["totals"]
    report_quality._pdf_body(story, styles, f"El Examen Complexivo registró una aprobación final global de {report_quality._pct(totals['approval_percentage'])}. Participaron {totals['supplementary']} estudiantes en supletorio y {totals['recovered']} lograron recuperar su condición académica mediante esta oportunidad. La interpretación por carrera debe considerar simultáneamente la aprobación de los cuatro núcleos, el cumplimiento de requisitos y la disponibilidad completa de las calificaciones.")
    if data["schedules"]["average_compliance"] is not None:
        report_quality._pdf_body(story, styles, f"El cumplimiento promedio de las actividades del cronograma que cuentan con evaluación fue {report_quality._pct(data['schedules']['average_compliance'])}. Las actividades no evaluadas requieren completar evidencia y observación antes del cierre documental.")

    report_quality._pdf_heading(story, context, styles, 1, "Conclusiones")
    conclusions = [
        f"El cumplimiento integral de requisitos alcanzó {report_quality._pct(data['requirements']['percentage'])} sobre {data['requirements']['total']} estudiantes únicos." if data["requirements"] else "No se dispuso de información suficiente para calcular el cumplimiento integral de requisitos.",
        f"La habilitación para el Examen Complexivo fue confirmada para {data['eligibility']['summary']['habilitated']} de {data['eligibility']['summary']['complexive_candidates']} candidatos, una vez verificada la aprobación individual de los cuatro núcleos.",
        f"El Examen Complexivo concluyó con {totals['final_approved']} aprobados, {totals['final_failed']} reprobados y {totals['not_evaluated']} no evaluados.",
        f"En Trabajo de Titulación se registraron {data['projects']['summary']['total']} estudiantes, de los cuales {data['projects']['summary']['approved']} aprobaron y {data['projects']['summary']['failed']} reprobaron.",
        f"La trazabilidad del proceso requiere mantener correspondencia entre la base institucional, las notas de núcleos, las evaluaciones del Complexivo, las actas y las evidencias de ejecución. Se identificaron {len(incidents)} novedades automáticas o registradas para seguimiento.",
    ]
    for conclusion in conclusions:
        report_quality._pdf_bullet(story, styles, conclusion)

    report_quality._pdf_heading(story, context, styles, 1, "Recomendaciones")
    for action in actions:
        report_quality._pdf_bullet(story, styles, action["action"])

    if actions:
        report_quality._pdf_heading(story, context, styles, 1, "Plan de mejora")
        rows = [[Paragraph(html.escape(str(item.get("finding") or "—")), styles["TableCell"]), Paragraph(html.escape(str(item.get("action") or "—")), styles["TableCell"]), Paragraph(html.escape(str(item.get("responsible") or "—")), styles["TableCell"]), item.get("due_date") or "Por definir", Paragraph(html.escape(str(item.get("indicator") or "—")), styles["TableCell"]), Paragraph(html.escape(str(item.get("evidence") or "—")), styles["TableCell"]), item.get("status") or "Pendiente"] for item in actions]
        report_quality._pdf_caption(story, styles, context.table_caption("Plan de mejora del proceso de titulación"))
        story += [report_quality._pdf_table(["Hallazgo", "Acción", "Responsable", "Fecha", "Indicador", "Evidencia", "Estado"], rows, [3.2 * cm, 3.6 * cm, 2.7 * cm, 1.8 * cm, 3.0 * cm, 2.4 * cm, 1.8 * cm]), Spacer(1, 0.2 * cm)]

    report_quality._pdf_heading(story, context, styles, 1, "Referencias legales e institucionales")
    for reference in (
        "Constitución de la República del Ecuador.",
        "Ley Orgánica de Educación Superior.",
        "Reglamento de Régimen Académico, Resolución RPC-SE-08-No.023-2022.",
        "Reglamento de la Unidad de Titulación y Eficiencia Terminal, código UTET-REG-25, versión 2.0.",
        "Manual de Procesos de la Unidad de Titulación y Eficiencia Terminal, versión 2.",
        "Guías institucionales de integración curricular y registros académicos cargados en Informtit.",
    ):
        report_quality._pdf_bullet(story, styles, reference)


def install() -> None:
    global CURRENT_REPORT_ID
    if getattr(report_quality, "_completion_report_installed", False):
        return

    report_structure.requirement_analysis = corrected_requirement_analysis
    report_quality.METHODOLOGY["Opciones de titulación"] = [
        "Las opciones institucionales de titulación son el Examen Complexivo y el Trabajo de Titulación. Este último podrá desarrollarse mediante Proyecto de Titulación o Artículo Académico, de conformidad con la normativa institucional."
    ]

    original_build_docx = report_quality.build_docx
    original_build_pdf = report_quality.build_pdf
    original_docx_heading = report_quality._docx_heading
    original_pdf_heading = report_quality._pdf_heading
    original_docx_legal = report_quality._docx_legal
    original_pdf_legal = report_quality._pdf_legal
    original_docx_requirements = report_quality._docx_requirements
    original_pdf_requirements = report_quality._pdf_requirements
    original_docx_nucleus = report_quality._docx_nucleus_results
    original_pdf_nucleus = report_quality._pdf_nucleus_results
    original_docx_complexive = report_quality._docx_complexive
    original_pdf_complexive = report_quality._pdf_complexive
    original_docx_projects = report_quality._docx_projects
    original_pdf_projects = report_quality._pdf_projects

    injected = {"docx": False, "pdf": False}

    def build_docx(report_id: int):
        global CURRENT_REPORT_ID
        CURRENT_REPORT_ID = report_id
        injected["docx"] = False
        try:
            return original_build_docx(report_id)
        finally:
            CURRENT_REPORT_ID = None

    def build_pdf(report_id: int):
        global CURRENT_REPORT_ID
        CURRENT_REPORT_ID = report_id
        injected["pdf"] = False
        try:
            return original_build_pdf(report_id)
        finally:
            CURRENT_REPORT_ID = None

    def docx_heading(document: Any, context: Any, level: int, title: str, page_break: bool = False):
        if CURRENT_REPORT_ID and level == 1 and title == "Introducción" and not injected["docx"]:
            _add_docx_executive_summary(document, CURRENT_REPORT_ID)
            injected["docx"] = True
        return original_docx_heading(document, context, level, title, page_break)

    def pdf_heading(story: list[Any], context: Any, styles: Any, level: int, title: str, page_break: bool = False):
        if CURRENT_REPORT_ID and level == 1 and title == "Introducción" and not injected["pdf"]:
            _add_pdf_executive_summary(story, styles, CURRENT_REPORT_ID)
            injected["pdf"] = True
        return original_pdf_heading(story, context, styles, level, title, page_break)

    def docx_legal(document: Any, context: Any, report: dict[str, Any]):
        _add_docx_objectives(document, context, report)
        _add_docx_report_methodology(document, context, int(report["id"]), report)
        return original_docx_legal(document, context, report)

    def pdf_legal(story: list[Any], context: Any, styles: Any, report: dict[str, Any]):
        _add_pdf_objectives(story, context, styles, report)
        _add_pdf_report_methodology(story, context, styles, int(report["id"]), report)
        return original_pdf_legal(story, context, styles, report)

    def docx_requirements(document: Any, context: Any, report_id: int):
        original_docx_requirements(document, context, report_id)
        _add_docx_requirement_definitions(document, context, report_id)

    def pdf_requirements(story: list[Any], context: Any, styles: Any, report_id: int):
        original_pdf_requirements(story, context, styles, report_id)
        _add_pdf_requirement_definitions(story, context, styles, report_id)

    def docx_nucleus(document: Any, context: Any, report_id: int):
        original_docx_nucleus(document, context, report_id)
        _add_docx_eligibility(document, context, report_id)

    def pdf_nucleus(story: list[Any], context: Any, styles: Any, report_id: int):
        original_pdf_nucleus(story, context, styles, report_id)
        _add_pdf_eligibility(story, context, styles, report_id)

    def docx_complexive(document: Any, context: Any, report: dict[str, Any]):
        original_docx_complexive(document, context, report)
        _add_docx_complexive_consolidated(document, context, report)

    def pdf_complexive(story: list[Any], context: Any, styles: Any, report: dict[str, Any], temp_paths: list[Any]):
        original_pdf_complexive(story, context, styles, report, temp_paths)
        _add_pdf_complexive_consolidated(story, context, styles, report)

    def docx_projects(document: Any, context: Any, report_id: int):
        projects = get_projects(report_id).get("projects", [])
        if projects:
            original_docx_projects(document, context, report_id)
        elif any(get_schedules_extended(report_id).get("thesis", [])):
            report_quality._docx_heading(document, context, 1, "Resultados del Trabajo de Titulación")
            report_quality._docx_body(document, "No se registraron estudiantes en esta opción de titulación durante el período analizado.")
        report = report_quality._report_data(report_id)
        _add_docx_global_process(document, context, report_id, report)

    def pdf_projects(story: list[Any], context: Any, styles: Any, report_id: int):
        projects = get_projects(report_id).get("projects", [])
        if projects:
            original_pdf_projects(story, context, styles, report_id)
        elif any(get_schedules_extended(report_id).get("thesis", [])):
            report_quality._pdf_heading(story, context, styles, 1, "Resultados del Trabajo de Titulación")
            report_quality._pdf_body(story, styles, "No se registraron estudiantes en esta opción de titulación durante el período analizado.")
        report = report_quality._report_data(report_id)
        _add_pdf_global_process(story, context, styles, report_id, report)

    report_quality.build_docx = build_docx
    report_quality.build_pdf = build_pdf
    report_quality._docx_heading = docx_heading
    report_quality._pdf_heading = pdf_heading
    report_quality._docx_legal = docx_legal
    report_quality._pdf_legal = pdf_legal
    report_quality._docx_requirements = docx_requirements
    report_quality._pdf_requirements = pdf_requirements
    report_quality._docx_schedules = _add_docx_schedules
    report_quality._pdf_schedules = _add_pdf_schedules
    report_quality._docx_nucleus_results = docx_nucleus
    report_quality._pdf_nucleus_results = pdf_nucleus
    report_quality._docx_complexive = docx_complexive
    report_quality._pdf_complexive = pdf_complexive
    report_quality._docx_projects = docx_projects
    report_quality._pdf_projects = pdf_projects
    report_quality._docx_post_sections = _add_docx_post_sections
    report_quality._pdf_post_sections = _add_pdf_post_sections
    core.build_docx = build_docx
    core.build_pdf = build_pdf
    report_quality._completion_report_installed = True
