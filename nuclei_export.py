from __future__ import annotations

import html
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, Spacer, Table, TableStyle

import institutional_export as base
import report_structure
from nuclei_service import get_nuclei


def _fmt(value: Any) -> str:
    if value is None or value == "":
        return "—"
    return f"{float(value):.2f}".replace(".", ",")


def _percent(part: int, total: int) -> str:
    return _fmt(part / total * 100 if total else 0) + " %"


def _docx_score_table(document: Any, course: dict[str, Any]) -> None:
    assessments = course.get("assessments", [])
    headers = ["Estudiante"] + [item["name"] for item in assessments] + ["Total", "Estado"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    remaining = 3.25
    activity_width = remaining / max(len(assessments), 1)
    widths = [1.8] + [activity_width] * len(assessments) + [0.55, 0.7]
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        base.set_width(cell, widths[index])
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(6.5)
    for student in course.get("students", []):
        values = [student["full_name"]] + [_fmt(score.get("grade")) for score in student.get("scores", [])] + [_fmt(student.get("final_grade")), student.get("final_status") or "No evaluado"]
        cells = table.add_row().cells
        for index, value in enumerate(values):
            base.set_width(cells[index], widths[index])
            cells[index].text = str(value)
            for paragraph in cells[index].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(6.5)
    document.add_paragraph()


def _docx_averages(document: Any, course: dict[str, Any]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Actividad"
    table.rows[0].cells[1].text = "Promedio"
    for cell in table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for item in course.get("activity_averages", []):
        row = table.add_row().cells
        row[0].text = str(item.get("name") or "")
        row[1].text = _fmt(item.get("calculated_average"))
    row = table.add_row().cells
    row[0].text = "Total del curso"
    row[1].text = _fmt(course.get("course_average"))
    document.add_paragraph()


def add_docx_nuclei(document: Any, report: dict[str, Any]) -> None:
    courses = get_nuclei(int(report["id"])).get("courses", [])
    if not courses:
        return
    document.add_heading("RESULTADOS DE LOS NÚCLEOS ESTRUCTURANTES", level=1)
    for course in courses:
        document.add_heading(
            f"{course['career_name']} – Núcleo {course['nucleus_number']}", level=2
        )
        responsible = (
            f"El curso fue impartido por {course.get('teacher_name') or 'docente pendiente de confirmar'}"
            f" y contó con el seguimiento de {course.get('coordinator_name') or 'la coordinación de carrera'}."
        )
        document.add_paragraph(responsible)
        document.add_paragraph(
            f"Se registraron {course.get('graded_students', 0)} estudiantes con calificaciones. "
            f"El promedio general fue {_fmt(course.get('course_average'))}; "
            f"{course.get('approved_count', 0)} estudiantes aprobaron "
            f"({_percent(int(course.get('approved_count') or 0), int(course.get('graded_students') or 0))}) y "
            f"{course.get('failed_count', 0)} reprobaron "
            f"({_percent(int(course.get('failed_count') or 0), int(course.get('graded_students') or 0))})."
        )
        _docx_score_table(document, course)
        document.add_heading("Promedios por actividad", level=3)
        _docx_averages(document, course)
        if int(course.get("missing_grades") or 0) or int(course.get("extra_grades") or 0):
            document.add_paragraph(
                f"Control de correspondencia: {course.get('matched_students', 0)} coincidencias, "
                f"{course.get('missing_grades', 0)} participantes sin calificación y "
                f"{course.get('extra_grades', 0)} calificaciones sin participante identificado."
            )


def _pdf_score_table(course: dict[str, Any], styles: Any) -> Table:
    assessments = course.get("assessments", [])
    remaining = 8.5 * cm
    activity_width = remaining / max(len(assessments), 1)
    widths = [4.5 * cm] + [activity_width] * len(assessments) + [1.5 * cm, 2.0 * cm]
    data: list[list[Any]] = [
        ["Estudiante"] + [Paragraph(html.escape(item["name"]), styles["NucleusCell"]) for item in assessments] + ["Total", "Estado"]
    ]
    for student in course.get("students", []):
        data.append(
            [Paragraph(html.escape(student["full_name"]), styles["NucleusCell"])]
            + [_fmt(score.get("grade")) for score in student.get("scores", [])]
            + [_fmt(student.get("final_grade")), student.get("final_status") or "No evaluado"]
        )
    table = Table(data, repeatRows=1, colWidths=widths)
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.35, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#244a73")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 6),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (1, 1), (-1, -1), "CENTER"),
            ]
        )
    )
    return table


def _pdf_averages(course: dict[str, Any]) -> Table:
    data = [["Actividad", "Promedio"]]
    for item in course.get("activity_averages", []):
        data.append([item.get("name") or "", _fmt(item.get("calculated_average"))])
    data.append(["Total del curso", _fmt(course.get("course_average"))])
    table = Table(data, colWidths=[12.5 * cm, 3.5 * cm], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#244a73")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("ALIGN", (1, 1), (1, -1), "CENTER"),
            ]
        )
    )
    return table


def add_pdf_nuclei(story: list[Any], report: dict[str, Any], styles: Any) -> None:
    courses = get_nuclei(int(report["id"])).get("courses", [])
    if not courses:
        return
    if "NucleusCell" not in styles:
        from reportlab.lib.styles import ParagraphStyle
        styles.add(ParagraphStyle("NucleusCell", parent=styles["BodyText"], fontSize=6, leading=7))
    story.append(Paragraph("RESULTADOS DE LOS NÚCLEOS ESTRUCTURANTES", styles["Heading1"]))
    for course in courses:
        story.append(Paragraph(f"{course['career_name']} – Núcleo {course['nucleus_number']}", styles["Heading2"]))
        story.append(
            Paragraph(
                f"El curso fue impartido por {html.escape(course.get('teacher_name') or 'docente pendiente de confirmar')} "
                f"y contó con el seguimiento de {html.escape(course.get('coordinator_name') or 'la coordinación de carrera')}. "
                f"Se registraron {course.get('graded_students', 0)} estudiantes con calificaciones. "
                f"El promedio general fue {_fmt(course.get('course_average'))}; "
                f"{course.get('approved_count', 0)} aprobaron y {course.get('failed_count', 0)} reprobaron.",
                styles["BodyJustified"],
            )
        )
        story += [_pdf_score_table(course, styles), Spacer(1, 0.25 * cm), Paragraph("Promedios por actividad", styles["Heading3"]), _pdf_averages(course), Spacer(1, 0.35 * cm)]


def install() -> None:
    original_docx = report_structure._add_docx_complexive
    original_pdf = report_structure._add_pdf_complexive

    def docx_with_nuclei(document: Any, report: dict[str, Any]) -> None:
        add_docx_nuclei(document, report)
        original_docx(document, report)

    def pdf_with_nuclei(story: list[Any], report: dict[str, Any], styles: Any, temp_paths: list[Any]) -> None:
        add_pdf_nuclei(story, report, styles)
        original_pdf(story, report, styles, temp_paths)

    report_structure._add_docx_complexive = docx_with_nuclei
    report_structure._add_pdf_complexive = pdf_with_nuclei
