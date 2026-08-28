from __future__ import annotations

import tempfile
from pathlib import Path
from typing import Any

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
    Image as RLImage,
)

from analytics import summary
from db import BASE_DIR, connection, rows_to_dicts

EXPORT_DIR = BASE_DIR / "exports"
UPLOAD_DIR = BASE_DIR / "uploads"


def _fmt(value: Any) -> str:
    if value is None:
        return "—"
    if isinstance(value, float):
        return f"{value:.2f}".replace(".", ",")
    return str(value)


def load_report_data(report_id: int) -> dict[str, Any]:
    with connection() as conn:
        report_row = conn.execute("SELECT * FROM reports WHERE id = ?", (report_id,)).fetchone()
        if not report_row:
            raise ValueError("El informe no existe.")
        report = dict(report_row)
        report["sections"] = rows_to_dicts(
            conn.execute(
                "SELECT * FROM institutional_sections WHERE report_id = ? AND visible = 1 ORDER BY sort_order, id",
                (report_id,),
            ).fetchall()
        )
        careers = rows_to_dicts(
            conn.execute(
                "SELECT * FROM careers WHERE report_id = ? ORDER BY sort_order, name", (report_id,)
            ).fetchall()
        )
        students_by_career: dict[int, list[dict[str, Any]]] = {
            int(career["id"]): [] for career in careers
        }
        analyses_by_career: dict[int, dict[str, dict[str, Any]]] = {
            int(career["id"]): {} for career in careers
        }
        images_by_career: dict[int, list[dict[str, Any]]] = {
            int(career["id"]): [] for career in careers
        }
        if careers:
            career_ids = [int(career["id"]) for career in careers]
            for start in range(0, len(career_ids), 400):
                chunk = career_ids[start:start + 400]
                placeholders = ",".join("?" for _ in chunk)

                student_rows = rows_to_dicts(conn.execute(
                    f"""
                    SELECT * FROM students
                    WHERE career_id IN ({placeholders})
                    ORDER BY career_id, full_name, id
                    """,
                    tuple(chunk),
                ).fetchall())
                for student in student_rows:
                    students_by_career.setdefault(int(student["career_id"]), []).append(student)

                analysis_rows = rows_to_dicts(conn.execute(
                    f"""
                    SELECT * FROM analyses
                    WHERE career_id IN ({placeholders})
                    ORDER BY career_id, id
                    """,
                    tuple(chunk),
                ).fetchall())
                for analysis in analysis_rows:
                    analyses_by_career.setdefault(int(analysis["career_id"]), {})[
                        str(analysis["section"])
                    ] = dict(analysis)

                image_rows = rows_to_dicts(conn.execute(
                    f"""
                    SELECT * FROM images
                    WHERE career_id IN ({placeholders})
                    ORDER BY career_id, sort_order, id
                    """,
                    tuple(chunk),
                ).fetchall())
                for image in image_rows:
                    images_by_career.setdefault(int(image["career_id"]), []).append(image)

        for career in careers:
            career_id = int(career["id"])
            career["students"] = students_by_career.get(career_id, [])
            career["analyses"] = analyses_by_career.get(career_id, {})
            career["images"] = images_by_career.get(career_id, [])
        report["careers"] = careers
        report["general_images"] = rows_to_dicts(
            conn.execute(
                "SELECT * FROM images WHERE report_id = ? AND career_id IS NULL ORDER BY sort_order, id",
                (report_id,),
            ).fetchall()
        )
        return report


def create_chart(career: dict[str, Any], output_path: Path) -> None:
    ordinary = summary(career["students"], "ordinario")
    final = summary(career["students"], "consolidado")
    labels = ["Aprobados ordinario", "Requieren mejora", "Aprobados finales", "Reprobados finales"]
    values = [ordinary["approved"], ordinary["failed"] + ordinary["not_evaluated"], final["approved"], final["failed"]]
    figure = plt.figure(figsize=(8, 4.6))
    axis = figure.add_subplot(111)
    axis.bar(labels, values)
    axis.set_title(f"Resultados de {career['name']}")
    axis.set_ylabel("Número de estudiantes")
    axis.tick_params(axis="x", rotation=18)
    for index, value in enumerate(values):
        axis.text(index, value + 0.1, str(value), ha="center")
    figure.tight_layout()
    figure.savefig(output_path, dpi=160, bbox_inches="tight")
    plt.close(figure)


def build_docx(report_id: int) -> Path:
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    report = load_report_data(report_id)
    output = EXPORT_DIR / f"informtit_{report_id}.docx"
    document = Document()

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("INFORME FINAL DEL PROCESO DE TITULACIÓN")
    run.bold = True
    run.font.size = Pt(20)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"{report['period']} – Modalidad {report['modality'].replace('_', ' ').title()}").bold = True
    document.add_paragraph()

    metadata = document.add_table(rows=0, cols=2)
    metadata.style = "Table Grid"
    for label, value in [
        ("Código", report.get("code", "")),
        ("Versión", report.get("version", "")),
        ("Fecha de elaboración", report.get("elaboration_date", "")),
        ("Elaborado por", f"{report.get('prepared_by', '')} – {report.get('prepared_role', '')}"),
        ("Revisado por", f"{report.get('reviewed_by', '')} – {report.get('reviewed_role', '')}"),
        ("Aprobado por", f"{report.get('approved_by', '')} – {report.get('approved_role', '')}"),
    ]:
        cells = metadata.add_row().cells
        cells[0].text = label
        cells[1].text = value
    document.add_page_break()

    document.add_heading("Contenido", level=1)
    for section in report["sections"]:
        document.add_paragraph(section["title"], style="List Number")
    for career in report["careers"]:
        document.add_paragraph(f"Resultados: {career['name']}", style="List Number")
    document.add_page_break()

    for section in report["sections"]:
        document.add_heading(section["title"], level=1)
        for paragraph in section["content"].split("\n"):
            if paragraph.strip():
                document.add_paragraph(paragraph.strip())

    for image in report["general_images"]:
        path = UPLOAD_DIR / image["filename"]
        if path.exists():
            document.add_picture(str(path), width=Inches(6.2))
            caption = document.add_paragraph(image.get("title") or image.get("original_name"))
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for career in report["careers"]:
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_heading(career["name"], level=1)
        _add_docx_phase(document, career, "ordinario", "Resultados de la evaluación ordinaria")
        _add_docx_phase(document, career, "supletorio", "Resultados de la evaluación supletoria")
        _add_docx_phase(document, career, "consolidado", "Resultado consolidado")

        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp:
            chart_path = Path(temp.name)
        create_chart(career, chart_path)
        document.add_picture(str(chart_path), width=Inches(6.4))
        chart_path.unlink(missing_ok=True)

        for image in career["images"]:
            path = UPLOAD_DIR / image["filename"]
            if path.exists():
                document.add_picture(str(path), width=Inches(6.2))
                caption = document.add_paragraph(image.get("title") or image.get("original_name"))
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.save(output)
    return output


def _add_docx_phase(document: Document, career: dict[str, Any], phase: str, heading: str) -> None:
    data = summary(career["students"], phase)
    analysis = career["analyses"].get(phase, {})
    document.add_heading(heading, level=2)
    before = analysis.get("text_before") or _default_before(career["name"], phase, data)
    document.add_paragraph(before)

    rows = data["rows"]
    if phase == "ordinario":
        headers = ["Estudiante", "Teórico", "Práctico", "Final", "Estado"]
        values = [
            [row["full_name"], _fmt(row["ordinary_theory"]), _fmt(row["ordinary_practical"]), _fmt(row["ordinary_final"]), row["ordinary_status"]]
            for row in rows
        ]
    elif phase == "supletorio":
        headers = ["Estudiante", "Componente", "Teórico sup.", "Práctico sup.", "Final", "Estado"]
        values = [
            [row["full_name"], row["supplementary_component"], _fmt(row["supplementary_theory"]), _fmt(row["supplementary_practical"]), _fmt(row["supplementary_final"]), row["final_status"]]
            for row in rows
        ]
    else:
        headers = ["Estudiante", "Final", "Estado", "Rindió supletorio"]
        values = [
            [row["full_name"], _fmt(row["final_grade"]), row["final_status"], "Sí" if row["supplementary_participant"] else "No"]
            for row in rows
        ]

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for value_row in values:
        cells = table.add_row().cells
        for index, value in enumerate(value_row):
            cells[index].text = str(value)

    after = analysis.get("text_after") or _default_after(data)
    document.add_paragraph(after)


def _default_before(career_name: str, phase: str, data: dict[str, Any]) -> str:
    labels = {
        "ordinario": "evaluación ordinaria",
        "supletorio": "evaluación supletoria",
        "consolidado": "resultado consolidado",
    }
    return (
        f"A continuación, se presentan los resultados de la {labels[phase]} de la carrera de {career_name}. "
        f"La tabla contiene información de {data['total']} estudiantes y resume el desempeño alcanzado."
    )


def _default_after(data: dict[str, Any]) -> str:
    return (
        f"De los {data['total']} registros analizados, {data['approved']} alcanzaron la aprobación "
        f"({data['approved_pct']:.2f} %) y {data['failed']} no alcanzaron la calificación mínima. "
        f"El promedio registrado fue {_fmt(data['average_final'])}."
    )


def build_pdf(report_id: int) -> Path:
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    report = load_report_data(report_id)
    output = EXPORT_DIR / f"informtit_{report_id}.pdf"
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="BodyJustified", parent=styles["BodyText"], alignment=TA_JUSTIFY, leading=15))
    styles.add(ParagraphStyle(name="CenterTitle", parent=styles["Title"], alignment=TA_CENTER, leading=24))
    story: list[Any] = []
    temp_paths: list[Path] = []

    story.append(Spacer(1, 5 * cm))
    story.append(Paragraph("INFORME FINAL DEL PROCESO DE TITULACIÓN", styles["CenterTitle"]))
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph(f"{report['period']} – Modalidad {report['modality'].replace('_', ' ').title()}", styles["Heading2"]))
    story.append(Spacer(1, 2 * cm))
    metadata = [
        ["Código", report.get("code", "")],
        ["Versión", report.get("version", "")],
        ["Fecha", report.get("elaboration_date", "")],
        ["Elaborado por", report.get("prepared_by", "")],
        ["Revisado por", report.get("reviewed_by", "")],
        ["Aprobado por", report.get("approved_by", "")],
    ]
    meta_table = Table(metadata, colWidths=[4 * cm, 11 * cm])
    meta_table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.grey), ("BACKGROUND", (0, 0), (0, -1), colors.lightgrey), ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story.append(meta_table)
    story.append(PageBreak())

    for section in report["sections"]:
        story.append(Paragraph(section["title"], styles["Heading1"]))
        for paragraph in section["content"].split("\n"):
            if paragraph.strip():
                story.append(Paragraph(paragraph.strip(), styles["BodyJustified"]))
                story.append(Spacer(1, 0.25 * cm))

    for career in report["careers"]:
        story.append(PageBreak())
        story.append(Paragraph(career["name"], styles["Heading1"]))
        for phase, heading in [
            ("ordinario", "Resultados de la evaluación ordinaria"),
            ("supletorio", "Resultados de la evaluación supletoria"),
            ("consolidado", "Resultado consolidado"),
        ]:
            data = summary(career["students"], phase)
            analysis = career["analyses"].get(phase, {})
            story.append(Paragraph(heading, styles["Heading2"]))
            story.append(Paragraph(analysis.get("text_before") or _default_before(career["name"], phase, data), styles["BodyJustified"]))
            story.append(Spacer(1, 0.2 * cm))
            table_data = _pdf_table_data(data, phase)
            table = Table(table_data, repeatRows=1, colWidths=_pdf_col_widths(phase))
            table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.35, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#244a73")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(table)
            story.append(Spacer(1, 0.25 * cm))
            story.append(Paragraph(analysis.get("text_after") or _default_after(data), styles["BodyJustified"]))
            story.append(Spacer(1, 0.4 * cm))

        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp:
            chart_path = Path(temp.name)
        create_chart(career, chart_path)
        story.append(RLImage(str(chart_path), width=16 * cm, height=9 * cm))
        temp_paths.append(chart_path)

    doc = SimpleDocTemplate(str(output), pagesize=A4, rightMargin=1.5 * cm, leftMargin=1.5 * cm, topMargin=1.5 * cm, bottomMargin=1.5 * cm, title=report["name"])
    try:
        doc.build(story)
    finally:
        for temp_path in temp_paths:
            temp_path.unlink(missing_ok=True)
    return output


def _pdf_col_widths(phase: str) -> list[float]:
    if phase == "ordinario":
        return [7 * cm, 2 * cm, 2 * cm, 2 * cm, 2.5 * cm]
    if phase == "supletorio":
        return [5.5 * cm, 3 * cm, 2 * cm, 2 * cm, 2 * cm, 2.5 * cm]
    return [8 * cm, 2.5 * cm, 2.5 * cm, 3 * cm]


def _pdf_table_data(data: dict[str, Any], phase: str) -> list[list[Any]]:
    rows = data["rows"]
    if phase == "ordinario":
        result = [["Estudiante", "Teórico", "Práctico", "Final", "Estado"]]
        result.extend([[row["full_name"], _fmt(row["ordinary_theory"]), _fmt(row["ordinary_practical"]), _fmt(row["ordinary_final"]), row["ordinary_status"]] for row in rows])
        return result
    if phase == "supletorio":
        result = [["Estudiante", "Componente", "Teórico sup.", "Práctico sup.", "Final", "Estado"]]
        result.extend([[row["full_name"], row["supplementary_component"], _fmt(row["supplementary_theory"]), _fmt(row["supplementary_practical"]), _fmt(row["supplementary_final"]), row["final_status"]] for row in rows])
        return result
    result = [["Estudiante", "Final", "Estado", "Supletorio"]]
    result.extend([[row["full_name"], _fmt(row["final_grade"]), row["final_status"], "Sí" if row["supplementary_participant"] else "No"] for row in rows])
    return result
