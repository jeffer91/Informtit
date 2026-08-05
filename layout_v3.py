from __future__ import annotations

from typing import Any

from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import PageBreak, Paragraph, Spacer, Table, TableStyle

import institutional_export as base
import layout_v2
from institutional_defaults import value


def _font(run: Any, size: float, bold: bool = False) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold


def _row_height(row: Any, inches: float) -> None:
    row.height = Inches(inches)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


def _signature_top(cell: Any, report: dict[str, Any], label: str, section_name: str) -> None:
    p = base.clear_cell(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    _font(p.add_run(label), 8.2)

    image = base.image_path(base.image_for(report, section_name))
    if image:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(str(image), width=Inches(1.58))


def _label_value(cell: Any, label: str, text: str, line_break: bool = False) -> None:
    p = base.clear_cell(cell)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    _font(p.add_run(label), 8.0, True)
    if line_break:
        p.add_run("\n")
    _font(p.add_run(text), 8.0)


def cover_docx(document: Any, report: dict[str, Any]) -> None:
    # El bloque central se desplaza hacia abajo para quedar centrado visualmente
    # entre el encabezado y el pie de página.
    p = document.add_paragraph()
    p.paragraph_format.space_before = Inches(2.08)
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(p.add_run("Informe Final Del Proceso De Titulación."), 18, True)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(p.add_run(str(report.get("period", ""))), 18, True)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Inches(3.62)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(p.add_run(f"Modalidad {base.modality(report)}"), 18, True)

    table = document.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            base.set_width(cell, 2.38)

    _row_height(table.rows[0], 1.72)
    _row_height(table.rows[1], 0.32)
    _row_height(table.rows[2], 0.62)

    people = [
        ("ELABORADO POR:", base.SIG_PREPARED, value(report, "prepared_by"), value(report, "prepared_role")),
        ("REVISADO POR:", base.SIG_REVIEWED, value(report, "reviewed_by"), value(report, "reviewed_role")),
        ("APROBADO POR:", base.SIG_APPROVED, value(report, "approved_by"), value(report, "approved_role")),
    ]
    for index, (label, section_name, name, role) in enumerate(people):
        _signature_top(table.cell(0, index), report, label, section_name)
        _label_value(table.cell(1, index), "NOMBRE: ", name)
        _label_value(table.cell(2, index), "CARGO:", role, True)

    document.add_page_break()


def _signature_top_flowables(report: dict[str, Any], label: str, section_name: str, styles: Any) -> list[Any]:
    style = ParagraphStyle(
        f"SignatureTopV3{section_name}",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8.5,
        leading=10,
        alignment=TA_CENTER,
        spaceAfter=4,
    )
    items: list[Any] = [Paragraph(label, style), Spacer(1, 0.18 * cm)]
    image = base.image_path(base.image_for(report, section_name))
    items.append(base.fit_image(image, 4.65 * cm, 1.9 * cm) if image else Spacer(1, 1.9 * cm))
    return items


def _name(text: str, styles: Any) -> Paragraph:
    style = ParagraphStyle(
        "SignatureNameV3",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8.3,
        leading=9.3,
        alignment=TA_LEFT,
    )
    return Paragraph(f"<b>NOMBRE:</b> {text}", style)


def _role(text: str, styles: Any) -> Paragraph:
    style = ParagraphStyle(
        "SignatureRoleV3",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8.3,
        leading=9.3,
        alignment=TA_LEFT,
    )
    return Paragraph(f"<b>CARGO:</b><br/>{text}", style)


def cover_pdf(report: dict[str, Any], styles: Any) -> list[Any]:
    title = ParagraphStyle(
        "CoverTitleV3",
        parent=styles["Title"],
        alignment=TA_CENTER,
        fontName="Helvetica-Bold",
        fontSize=17,
        leading=20,
        spaceAfter=2,
    )
    period = ParagraphStyle("CoverPeriodV3", parent=title, fontSize=16, leading=19, spaceAfter=2)
    mode = ParagraphStyle("CoverModeV3", parent=title, fontSize=16, leading=19, spaceAfter=0)

    people = [
        ("ELABORADO POR:", base.SIG_PREPARED, value(report, "prepared_by"), value(report, "prepared_role")),
        ("REVISADO POR:", base.SIG_REVIEWED, value(report, "reviewed_by"), value(report, "reviewed_role")),
        ("APROBADO POR:", base.SIG_APPROVED, value(report, "approved_by"), value(report, "approved_role")),
    ]
    data = [
        [_signature_top_flowables(report, label, section_name, styles) for label, section_name, _, _ in people],
        [_name(name, styles) for _, _, name, _ in people],
        [_role(role, styles) for _, _, _, role in people],
    ]
    table = Table(data, colWidths=[5.55 * cm] * 3, rowHeights=[3.45 * cm, 0.68 * cm, 1.22 * cm])
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.75, colors.black),
        ("VALIGN", (0, 0), (-1, 0), "TOP"),
        ("VALIGN", (0, 1), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, 0), 4),
        ("RIGHTPADDING", (0, 0), (-1, 0), 4),
        ("TOPPADDING", (0, 0), (-1, 0), 4),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 4),
        ("LEFTPADDING", (0, 1), (-1, -1), 5),
        ("RIGHTPADDING", (0, 1), (-1, -1), 4),
        ("TOPPADDING", (0, 1), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 1), (-1, -1), 2),
    ]))

    return [
        Spacer(1, 6.0 * cm),
        Paragraph("Informe Final Del Proceso De Titulación.", title),
        Paragraph(str(report.get("period", "")), period),
        Paragraph(f"Modalidad {base.modality(report)}", mode),
        Spacer(1, 8.85 * cm),
        table,
        PageBreak(),
    ]


def install() -> None:
    layout_v2.install()
    base.cover_docx = cover_docx
    base.cover_pdf = cover_pdf
