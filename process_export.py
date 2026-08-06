from __future__ import annotations

import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

import app as core
import institutional_export as base
from analytics import summary
from process_service import get_projects, get_schedules


def _docx_header_cell(cell: Any, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(8)


def _docx_body_cell(cell: Any, text: Any, center: bool = False) -> None:
    cell.text = "—" if text is None or text == "" else str(text)
    for paragraph in cell.paragraphs:
        if center:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8)


def _add_docx_schedule(document: Document, title: str, rows: list[dict[str, Any]], show_phase: bool) -> None:
    document.add_heading(title, level=1)
    columns = ["Fase", "Actividad", "Fecha de inicio", "Fecha de fin"] if show_phase else ["Actividad", "Fecha de inicio", "Fecha de fin"]
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for index, label in enumerate(columns):
        _docx_header_cell(table.rows[0].cells[index], label)
    for row in rows:
        cells = table.add_row().cells
        offset = 0
        if show_phase:
            _docx_body_cell(cells[0], row.get("phase"))
            offset = 1
        _docx_body_cell(cells[offset], row.get("activity"))
        _docx_body_cell(cells[offset + 1], row.get("start_date"), True)
        _docx_body_cell(cells[offset + 2], row.get("end_date"), True)
    document.add_paragraph()


def _add_docx_projects(document: Document, report_id: int) -> None:
    data = get_projects(report_id)
    projects = data["projects"]
    document.add_heading("Resultados del Trabajo de Titulación", level=1)
    if not projects:
        document.add_paragraph("No se registraron estudiantes en la opción Trabajo de Titulación.")
        return
    summary_data = data["summary"]
    document.add_paragraph(
        f"Se registraron {summary_data['total']} estudiantes. "
        f"Aprobados: {summary_data['approved']}; reprobados: {summary_data['failed']}; "
        f"promedio final: {summary_data['average_final'] if summary_data['average_final'] is not None else '—'}."
    )
    for project in projects:
        document.add_heading(project["full_name"], level=2)
        info = document.add_table(rows=4, cols=4)
        info.style = "Table Grid"
        values = [
            ("Cédula", project.get("identification"), "Carrera", project.get("career_name")),
            ("Código de carrera", project.get("career_code"), "Acta de grado", project.get("act_number")),
            ("Fecha del acta", project.get("act_date"), "Trabajo escrito", project.get("written_average")),
            ("Defensa oral", project.get("oral_average"), "Calificación final", project.get("final_grade")),
        ]
        for row_index, row_values in enumerate(values):
            for column_index, value in enumerate(row_values):
                if column_index % 2 == 0:
                    _docx_header_cell(info.rows[row_index].cells[column_index], str(value))
                else:
                    _docx_body_cell(info.rows[row_index].cells[column_index], value, True)
        document.add_paragraph(
            f"Vocales: {project.get('vocal_1') or '—'}; {project.get('vocal_2') or '—'}; {project.get('vocal_3') or '—'}."
        )
        for evaluation_type, title in (("practical", "Evaluación práctica"), ("defense", "Evaluación de la defensa")):
            document.add_heading(title, level=3)
            scores = [row for row in project["scores"] if row["evaluation_type"] == evaluation_type]
            table = document.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            headers = ["Criterio", "Máximo", "Primer vocal", "Segundo vocal", "Tercer vocal"]
            for index, label in enumerate(headers):
                _docx_header_cell(table.rows[0].cells[index], label)
            for score in scores:
                cells = table.add_row().cells
                for index, value in enumerate((score["criterion"], score["max_score"], score["vocal_1"], score["vocal_2"], score["vocal_3"])):
                    _docx_body_cell(cells[index], value, index > 0)
        document.add_paragraph()


def build_docx(report_id: int) -> Path:
    output = base.build_docx(report_id)
    document = Document(output)
    schedules = get_schedules(report_id)
    document.add_page_break()
    _add_docx_schedule(document, "Cronograma de Núcleos y Examen Complexivo", schedules["complexive"], False)
    _add_docx_schedule(document, "Cronograma del Trabajo de Titulación", schedules["thesis"], True)
    _add_docx_projects(document, report_id)
    document.save(output)
    return output


def _pdf_schedule(title: str, rows: list[dict[str, Any]], show_phase: bool, styles: Any) -> list[Any]:
    headers = ["Fase", "Actividad", "Fecha de inicio", "Fecha de fin"] if show_phase else ["Actividad", "Fecha de inicio", "Fecha de fin"]
    data: list[list[Any]] = [headers]
    for row in rows:
        values = [row.get("activity"), row.get("start_date"), row.get("end_date")]
        if show_phase:
            values.insert(0, row.get("phase"))
        data.append([Paragraph(str(value or "—"), styles["BodyText"]) for value in values])
    widths = [4.1 * cm, 7.0 * cm, 2.7 * cm, 2.7 * cm] if show_phase else [10.5 * cm, 3.0 * cm, 3.0 * cm]
    table = Table(data, repeatRows=1, colWidths=widths)
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#244a73")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ALIGN", (-2, 1), (-1, -1), "CENTER"),
    ]))
    return [Paragraph(title, styles["Heading1"]), table, Spacer(1, 0.45 * cm)]


def _pdf_projects(report_id: int, styles: Any) -> list[Any]:
    data = get_projects(report_id)
    projects = data["projects"]
    story: list[Any] = [Paragraph("Resultados del Trabajo de Titulación", styles["Heading1"])]
    if not projects:
        story.append(Paragraph("No se registraron estudiantes en la opción Trabajo de Titulación.", styles["BodyText"]))
        return story
    summary_data = data["summary"]
    story.append(Paragraph(
        f"Se registraron {summary_data['total']} estudiantes; {summary_data['approved']} aprobaron y "
        f"{summary_data['failed']} reprobaron. El promedio final fue "
        f"{summary_data['average_final'] if summary_data['average_final'] is not None else '—'}.",
        styles["BodyJustified"],
    ))
    for project in projects:
        story += [Spacer(1, 0.25 * cm), Paragraph(project["full_name"], styles["Heading2"])]
        info = [
            ["Cédula", project.get("identification") or "—", "Carrera", project.get("career_name") or "—"],
            ["Acta", project.get("act_number") or "—", "Fecha", project.get("act_date") or "—"],
            ["Trabajo escrito", project.get("written_average"), "Defensa oral", project.get("oral_average")],
            ["Calificación final", project.get("final_grade"), "Estado", "Aprobado" if (project.get("final_grade") or 0) >= 7 else "Reprobado"],
        ]
        info_table = Table(info, colWidths=[3.0 * cm, 4.1 * cm, 3.0 * cm, 6.0 * cm])
        info_table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.4, colors.grey), ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"), ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 7.5), ("VALIGN", (0, 0), (-1, -1), "TOP")]))
        story += [info_table, Spacer(1, 0.18 * cm), Paragraph(
            f"Vocales: {project.get('vocal_1') or '—'}; {project.get('vocal_2') or '—'}; {project.get('vocal_3') or '—'}.",
            styles["BodyText"],
        )]
        for evaluation_type, title in (("practical", "Evaluación práctica"), ("defense", "Evaluación de la defensa")):
            rows = [row for row in project["scores"] if row["evaluation_type"] == evaluation_type]
            score_data = [["Criterio", "Máximo", "Vocal 1", "Vocal 2", "Vocal 3"]] + [
                [row["criterion"], row["max_score"], row["vocal_1"], row["vocal_2"], row["vocal_3"]] for row in rows
            ]
            table = Table(score_data, repeatRows=1, colWidths=[7.2 * cm, 2.0 * cm, 2.4 * cm, 2.4 * cm, 2.4 * cm])
            table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.35, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#244a73")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 7), ("ALIGN", (1, 1), (-1, -1), "CENTER")]))
            story += [Paragraph(title, styles["Heading3"]), table, Spacer(1, 0.22 * cm)]
    return story


def build_pdf(report_id: int) -> Path:
    base.EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    report = base.legacy.load_report_data(report_id)
    output = base.EXPORT_DIR / f"informtit_{report_id}.pdf"
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("BodyJustified", parent=styles["BodyText"], alignment=TA_JUSTIFY, leading=15))
    story = base.cover_pdf(report, styles)
    temp_paths: list[Path] = []
    inf_added = False
    for section in report["sections"]:
        if section.get("section_key") == "metodologia" and not inf_added:
            path = base.image_path(base.image_for(report, base.INFOGRAPHIC))
            if path:
                story += [Paragraph("Infografía Proceso de Examen Complexivo", styles["Heading1"]), base.fit_image(path, 16.5 * cm, 19.5 * cm), PageBreak()]
            inf_added = True
        story.append(Paragraph(section["title"], styles["Heading1"]))
        for text in str(section.get("content") or "").split("\n"):
            if text.strip():
                story += [Paragraph(text.strip(), styles["BodyJustified"]), Spacer(1, 0.25 * cm)]

    schedules = get_schedules(report_id)
    story += [PageBreak()]
    story += _pdf_schedule("Cronograma de Núcleos y Examen Complexivo", schedules["complexive"], False, styles)
    story += _pdf_schedule("Cronograma del Trabajo de Titulación", schedules["thesis"], True, styles)
    story += _pdf_projects(report_id, styles)

    for image in report["general_images"]:
        path = base.image_path(image)
        if image.get("section") not in base.RESERVED and path:
            story += [base.fit_image(path, 16.5 * cm, 20 * cm), Spacer(1, 0.3 * cm)]
    for career in report["careers"]:
        story += [PageBreak(), Paragraph(career["name"], styles["Heading1"])]
        path = base.image_path(base.image_for(report, base.NUCLEI, int(career["id"])))
        if path:
            story += [base.fit_image(path, 11.5 * cm, 8.5 * cm), Spacer(1, 0.3 * cm)]
        for phase, heading in (("ordinario", "Resultados de la evaluación ordinaria"), ("supletorio", "Resultados de la evaluación supletoria"), ("consolidado", "Resultado consolidado")):
            phase_data = summary(career["students"], phase)
            analysis = career["analyses"].get(phase, {})
            story += [Paragraph(heading, styles["Heading2"]), Paragraph(analysis.get("text_before") or base.legacy._default_before(career["name"], phase, phase_data), styles["BodyJustified"]), Spacer(1, 0.2 * cm)]
            table = Table(base.legacy._pdf_table_data(phase_data, phase), repeatRows=1, colWidths=base.legacy._pdf_col_widths(phase))
            table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.35, colors.grey), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#244a73")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 7), ("VALIGN", (0, 0), (-1, -1), "TOP")]))
            story += [table, Spacer(1, 0.25 * cm), Paragraph(analysis.get("text_after") or base.legacy._default_after(phase_data), styles["BodyJustified"]), Spacer(1, 0.4 * cm)]
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp:
            chart = Path(temp.name)
        base.legacy.create_chart(career, chart)
        story.append(base.fit_image(chart, 16 * cm, 9 * cm))
        temp_paths.append(chart)
        for image in career["images"]:
            path = base.image_path(image)
            if image.get("section") != base.NUCLEI and path:
                story += [base.fit_image(path, 16.5 * cm, 20 * cm), Spacer(1, 0.3 * cm)]
    doc = SimpleDocTemplate(str(output), pagesize=A4, rightMargin=1.45 * cm, leftMargin=1.45 * cm, topMargin=3.4 * cm, bottomMargin=1.35 * cm, title=report["name"])
    try:
        doc.build(story, canvasmaker=lambda *args, **kwargs: base.NumberedCanvas(*args, report=report, **kwargs))
    finally:
        for path in temp_paths:
            path.unlink(missing_ok=True)
    return output


def install() -> None:
    core.build_docx = build_docx
    core.build_pdf = build_pdf
